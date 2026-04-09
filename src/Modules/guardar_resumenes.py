from docx import Document
from urllib.parse import quote
import io, requests
import re
import html as pyhtml
from html.parser import HTMLParser
import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from Modules.logger_config import get_logger
logger = get_logger(__name__)


class _HtmlToTextParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self._chunks = []

    def handle_starttag(self, tag, attrs):
        tag = (tag or "").lower()
        if tag == "br":
            self._chunks.append("\n")
            return
        if tag == "li":
            self._chunks.append("\n- ")
            return
        if tag in {"p", "div", "section", "article", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "table", "tr"}:
            self._chunks.append("\n")

    def handle_endtag(self, tag):
        tag = (tag or "").lower()
        if tag in {"p", "div", "section", "article", "h1", "h2", "h3", "h4", "h5", "h6", "li", "ul", "ol", "table", "tr"}:
            self._chunks.append("\n")

    def handle_data(self, data):
        if data:
            self._chunks.append(data)

    def get_text(self):
        raw = "".join(self._chunks)
        raw = pyhtml.unescape(raw)
        raw = raw.replace("\r\n", "\n").replace("\r", "\n")
        raw = re.sub(r"[ \t]+\n", "\n", raw)
        raw = re.sub(r"\n[ \t]+", "\n", raw)
        raw = re.sub(r"[ \t]{2,}", " ", raw)
        raw = re.sub(r"\n{3,}", "\n\n", raw)
        return raw.strip()


def _summary_to_word_text(contenido_resumen):
    if not contenido_resumen:
        return ""

    text = str(contenido_resumen).strip()
    if not text:
        return ""

    # Si llega como HTML, lo convertimos a texto legible para Word.
    if "<" in text and ">" in text:
        parser = _HtmlToTextParser()
        parser.feed(text)
        parser.close()
        text = parser.get_text()

    return text


def _looks_like_html(text):
    if not text:
        return False

    return bool(re.search(r"<\s*(h[1-6]|p|div|section|article|ul|ol|li|strong|b|em|i|br)\b", text, flags=re.I))


def _looks_like_markdown(text):
    if not text:
        return False

    markdown_patterns = (
        r"^\s{0,3}#{1,6}\s+",
        r"^\s{0,3}[-*+]\s+",
        r"^\s{0,3}\d+[.)]\s+",
        r"\*\*[^*]+\*\*",
        r"_[^_]+_",
        r"`[^`]+`",
    )
    return any(re.search(pattern, text, flags=re.MULTILINE) for pattern in markdown_patterns)


class _DocxHtmlSummaryParser(HTMLParser):
    def __init__(self, document):
        super().__init__(convert_charrefs=False)
        self.document = document
        self.current_paragraph = None
        self.bold_depth = 0
        self.italic_depth = 0
        self.list_stack = []
        self.in_heading_level = None
        self.has_written_content = False

    def _new_paragraph(self, style=None):
        self.current_paragraph = self.document.add_paragraph(style=style) if style else self.document.add_paragraph()
        return self.current_paragraph

    def _new_heading(self, level):
        word_level = min(max(level, 1), 4)
        self.current_paragraph = self.document.add_heading("", level=word_level)
        return self.current_paragraph

    def _ensure_paragraph(self):
        if self.current_paragraph is None:
            self._new_paragraph()
        return self.current_paragraph

    def _add_text(self, text):
        if text is None:
            return

        decoded = pyhtml.unescape(text)
        if not decoded:
            return
        if not decoded.strip() and "\n" in decoded:
            return

        paragraph = self._ensure_paragraph()
        run = paragraph.add_run(decoded)
        if self.bold_depth > 0:
            run.bold = True
        if self.italic_depth > 0:
            run.italic = True
        if decoded.strip():
            self.has_written_content = True

    def handle_starttag(self, tag, attrs):
        tag = (tag or "").lower()
        if tag in ("p", "div", "section", "article"):
            self._new_paragraph()
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.in_heading_level = int(tag[1])
            self._new_heading(self.in_heading_level)
        elif tag == "br":
            paragraph = self._ensure_paragraph()
            paragraph.add_run().add_break()
        elif tag in ("strong", "b"):
            self.bold_depth += 1
        elif tag in ("em", "i"):
            self.italic_depth += 1
        elif tag == "ul":
            self.list_stack.append("ul")
        elif tag == "ol":
            self.list_stack.append("ol")
        elif tag == "li":
            style = "List Number" if self.list_stack and self.list_stack[-1] == "ol" else "List Bullet"
            self._new_paragraph(style=style)

    def handle_endtag(self, tag):
        tag = (tag or "").lower()
        if tag in ("strong", "b"):
            self.bold_depth = max(0, self.bold_depth - 1)
        elif tag in ("em", "i"):
            self.italic_depth = max(0, self.italic_depth - 1)
        elif tag in ("ul", "ol"):
            if self.list_stack:
                self.list_stack.pop()
            self.current_paragraph = None
        elif tag in ("p", "div", "section", "article", "li"):
            self.current_paragraph = None
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.in_heading_level = None
            self.current_paragraph = None

    def handle_data(self, data):
        self._add_text(data)

    def handle_entityref(self, name):
        self._add_text(f"&{name};")

    def handle_charref(self, name):
        self._add_text(f"&#{name};")


def _write_html_to_doc(doc, html_text):
    parser = _DocxHtmlSummaryParser(doc)
    parser.feed(html_text)
    parser.close()
    return parser.has_written_content


def _add_markdown_runs(paragraph, text):
    if not text:
        return

    token_pattern = re.compile(
        r"(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`|(?<!\*)\*[^*\n]+\*(?!\*)|(?<!_)_[^_\n]+_(?!_))"
    )

    last_end = 0
    for match in token_pattern.finditer(text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("__") and token.endswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("_") and token.endswith("_"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)

        last_end = match.end()

    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _flush_markdown_paragraph(doc, paragraph_lines):
    if not paragraph_lines:
        return

    text = " ".join(line.strip() for line in paragraph_lines if line.strip()).strip()
    if not text:
        return

    paragraph = doc.add_paragraph()
    _add_markdown_runs(paragraph, text)


def _write_markdown_to_doc(doc, markdown_text):
    paragraph_lines = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            _flush_markdown_paragraph(doc, paragraph_lines)
            paragraph_lines = []
            continue

        heading_match = re.match(r"^\s{0,3}(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            _flush_markdown_paragraph(doc, paragraph_lines)
            paragraph_lines = []
            level = min(len(heading_match.group(1)), 4)
            paragraph = doc.add_heading(level=level)
            _add_markdown_runs(paragraph, heading_match.group(2).strip())
            continue

        bullet_match = re.match(r"^\s{0,3}[-*+]\s+(.+)$", stripped)
        if bullet_match:
            _flush_markdown_paragraph(doc, paragraph_lines)
            paragraph_lines = []
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_markdown_runs(paragraph, bullet_match.group(1).strip())
            continue

        numbered_match = re.match(r"^\s{0,3}\d+[.)]\s+(.+)$", stripped)
        if numbered_match:
            _flush_markdown_paragraph(doc, paragraph_lines)
            paragraph_lines = []
            paragraph = doc.add_paragraph(style="List Number")
            _add_markdown_runs(paragraph, numbered_match.group(1).strip())
            continue

        quote_match = re.match(r"^\s{0,3}>\s+(.+)$", stripped)
        if quote_match:
            _flush_markdown_paragraph(doc, paragraph_lines)
            paragraph_lines = []
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(quote_match.group(1).strip())
            run.italic = True
            continue

        paragraph_lines.append(stripped)

    _flush_markdown_paragraph(doc, paragraph_lines)

# -----------------------------------------------------
# GUARDAR RESUMEN EN TXT EN LA CARPETA (SUBVENCIÓN)
# -----------------------------------------------------
""" Función para guardar el resumen en un .txt dentro de la subvención """
def guardar_resumen_sharepoint(token, drive_id, path, contenido_resumen, subvencion_prefix):
    """
    Crea un .docx a partir de tu texto formateado y lo sube a SharePoint via Graph API.
    """
    nombre_archivo = f"resumen_{subvencion_prefix}"
    # 1) Crear el documento y añadir el resumen respetando HTML o Markdown básico
    doc = Document()
    raw_text = str(contenido_resumen or "").strip()
    plain_text = _summary_to_word_text(contenido_resumen)

    if raw_text and _looks_like_html(raw_text):
        wrote_content = _write_html_to_doc(doc, raw_text)
        if not wrote_content and plain_text:
            for block in plain_text.split("\n\n"):
                block = block.strip()
                if block:
                    doc.add_paragraph(block)
    elif plain_text and _looks_like_markdown(plain_text):
        _write_markdown_to_doc(doc, plain_text)
    elif plain_text:
        for block in plain_text.split("\n\n"):
            block = block.strip()
            if block:
                doc.add_paragraph(block)
    else:
        doc.add_paragraph("Sin contenido")

    # 2) Guardar el .docx en un buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 3) Subirlo con Graph API
    url = (f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{path}/{nombre_archivo}.docx:/content")
    headers = {
        "Authorization": f"Bearer {token}",
        # Content-Type correcto para .docx
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    response = requests.put(url, headers=headers, data=buffer.read())

    # 5) Comprobar resultado
    if response.status_code in (200, 201):
        logger.info(f"'{nombre_archivo}.docx' subido correctamente.")
        sharepoint_url = response.json().get("webUrl")
        return sharepoint_url
    else:
        logger.error("Error al subir:", response.status_code, response.text)
        return


# ---------------------------------------------------------
# GUARDAR RESUMEN EN BASE DE DATOS
# ---------------------------------------------------------
def guardar_resumen_en_base_datos(cursor, conn, subvencion_prefix, resumen_grantify):
    """
    Función para guardar el resumen en una base
    de datos
    """
    try:
        logger.info("Subiendo resumen a la base de datos...")
        # 3) Query de UPDATE con RETURNING para ver la fila afectada
        update_query = """
            UPDATE public.grants
            SET resumen_completo = %s
            WHERE codigobdns = %s
            RETURNING id, codigobdns, titulo_corto, resumen_completo;
        """

        cursor.execute(update_query, (resumen_grantify, subvencion_prefix))
        # Recuperamos la fila que se ha actualizado (si existe)
        fila_actualizada = cursor.fetchone()

        # 4) Confirmamos cambios
        conn.commit()

        # 5) Mostramos resultado
        if fila_actualizada:
            row_id = fila_actualizada[0]  
            codigo_bdns = fila_actualizada[1] 
            logger.info(f"Row ID: {row_id}")
            logger.info(f"Código BDNS: {codigo_bdns}")
            return row_id
        else:
            logger.error(f"Resumen a SQL: No se encontró ninguna fila con codigobdns = {subvencion_prefix}")
            return None
                    
    except Exception as e:
        logger.error(f"An error occurred subiendo resumen a SQL: {e}")
        return None

def guardar_link_en_base_datos(cursor, conn, subvencion_prefix, link_resumen):
    """
    Función para guardar el link del resumen en la base
    de datos
    """
    try:
        logger.info("Subiendo link del resumen a la base de datos...")
        # 3) Query de UPDATE con RETURNING para ver la fila afectada
        update_query = """
            UPDATE public.grants
            SET link_sharepoint_resumen = %s
            WHERE codigobdns = %s
            RETURNING id, codigobdns, titulo_corto, link_sharepoint_resumen;
        """

        cursor.execute(update_query, (link_resumen, subvencion_prefix))
        # Recuperamos la fila que se ha actualizado (si existe)
        fila_actualizada = cursor.fetchone()

        # 4) Confirmamos cambios
        conn.commit()

        # 5) Mostramos resultado
        if fila_actualizada:
            row_id = fila_actualizada[0]  
            codigo_bdns = fila_actualizada[1] 
            logger.info(f"Row ID: {row_id}")
            logger.info(f"Código BDNS: {codigo_bdns}")
            logger.info(f"Link Sharepoint Resumen: {fila_actualizada[3]}")
            return row_id
        else:
            logger.error(f"Link a SQL: No se encontró ninguna fila con codigobdns = {subvencion_prefix}")
            return None
                    
    except Exception as e:
        logger.error(f"An error occurred subiendo link a SQL: {e}")
        return None

def guardar_resumen_main(token, drive_id, path, resumen_innovalaus, resumen_grantial, cursor, conn, subvencion_prefix):
    
    logger.info("Guardando resumenes...")

    row_id = guardar_resumen_en_base_datos(cursor, conn, subvencion_prefix, resumen_grantial)

    if not row_id: 
        logger.error("Resumen no se ha podido guardar en sql")

    sharepoint_url = guardar_resumen_sharepoint(token, drive_id, path, resumen_innovalaus, subvencion_prefix)

    if not sharepoint_url: 
        logger.error("Resumen no se ha podido guardar en Sharepoint")

    guardar_link_en_base_datos(cursor, conn, subvencion_prefix, sharepoint_url)

    if not row_id:
        logger.error("Link no se ha podido guardar en sql")
    
    return row_id
