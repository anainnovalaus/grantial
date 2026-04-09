from flask import Flask, jsonify, request, send_file, make_response
from flask_cors import CORS
import os
import sys
from dotenv import load_dotenv
from datetime import datetime, timedelta, date
import boto3
import json
import csv
import base64
import requests
from argon2 import PasswordHasher, exceptions as argon2_exceptions
import uuid
from werkzeug.utils import secure_filename
import io
import re
import unicodedata
import zipfile
from html.parser import HTMLParser
import html as pyhtml
import threading
import logging
from urllib.parse import urlencode
from urllib.request import Request, urlopen
from urllib.error import HTTPError, URLError
from openai import OpenAI

from botocore.exceptions import ClientError
from flask import request, jsonify, send_file, make_response
from src.Modules.logger_config import get_logger
logger = get_logger(__name__)

# Importar el modelo de recomendación
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from src.Modules.main import run_main
from functions.get_hoy import subvenciones_hoy
from functions.get_finalidad import sub_finalidad
from functions.get_region import sub_region
from functions.get_beneficiario import sub_beneficiario
from functions.get_match import get_match
from functions.get_resumen import GetResumen
from functions.retreiver_milvus import get_preguntas
from functions.search_grants import GrantSearch
from utils.llm import LLM
from utils.postgreSQL import get_connection
from utils.recommendation_model import (
    ensure_recommendation_tables,
    get_recommended_grants,
    get_user_preferences,
    record_reco_event,
    refresh_reco_entity_state,
    select_user_entity,
    store_user_preference,
)
from utils.auth_routes import auth_bp
from utils.auth_helpers import require_auth
from utils.scrapeEntidad import scrape_main, extract_text
from utils.NewClientMakeMatch import main_match_new_client
from utils.newsletter_routes import newsletter_bp
from utils.chat_limits import check_user_limit, increment_usage

app = Flask(__name__)
CORS(app)

# Load environment variables from .env
load_dotenv()

def get_ssm_param(name, secure=True):
    ssm = boto3.client("ssm", region_name="eu-central-1")
    response = ssm.get_parameter(Name=name, WithDecryption=secure)
    return response["Parameter"]["Value"]

# Register the authentication blueprint
app.register_blueprint(auth_bp, url_prefix='/api/auth')
# Register the newsletter blueprint
app.register_blueprint(newsletter_bp, url_prefix='/api/newsletter')

# Warm recommendation schema/indexes at startup so first user request doesn't pay setup latency.
try:
    ensure_recommendation_tables()
except Exception as e:
    logger.warning("No se pudo inicializar esquema de recomendaciones al arrancar: %s", e)

OPENAI_PROMPT_SCRAPE_ENTITY = get_ssm_param("/grantify/openai/chatgpt_prompt_scrape_entity")

# Función para obtener el nombre del bucket de S3 desde SSM o variable de entorno
def get_bucket_name():
    try:
        ssm = boto3.client("ssm", region_name="eu-central-1")
        resp = ssm.get_parameter(Name="/grantify/s3/docsentidad", WithDecryption=False)
        return resp["Parameter"]["Value"]
    except ClientError as e:
        print("No se pudo leer de SSM, usando ENV:", e)
        bucket = os.getenv("GRANTIFY_S3_BUCKET")
        if not bucket:
            raise RuntimeError("No se encontró el bucket en SSM ni en variable de entorno")
        return bucket

S3_BUCKET_NAME = get_bucket_name()
s3_client = boto3.client("s3", region_name="eu-central-1")

# Clase para eliminar etiquetas HTML en los resúmenes
def normalize_summary_html(raw: str) -> str:
    if not raw:
        return ""
    # 1) Desescapar entidades HTML (&lt;h2&gt; -> <h2>)
    html = pyhtml.unescape(raw)

    # 2) Normalizar títulos y limpiar estilos inline
    html = re.sub(r'<h1\b([^>]*)>', r'<h2\1>', html, flags=re.I)
    html = re.sub(r'</h1>', '</h2>', html, flags=re.I)
    html = re.sub(r'\s*(style|class)="[^"]*"', '', html, flags=re.I)

    return html


def _table_has_column(cursor, table_name: str, column_name: str) -> bool:
    cursor.execute("""
        SELECT 1
        FROM information_schema.columns
        WHERE table_schema = 'public'
          AND table_name = %s
          AND column_name = %s
        LIMIT 1
    """, (table_name, column_name))
    return cursor.fetchone() is not None


def _grant_has_column(cursor, column_name: str) -> bool:
    return _table_has_column(cursor, 'grants', column_name)


def _normalize_cif_nif(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r'[^A-Za-z0-9]', '', value).upper()


def _parse_sql_date_like(value):
    if value in (None, ""):
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value).strip()
    if not text:
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            continue
    try:
        # Maneja ISO con timezone o formato extendido
        return datetime.fromisoformat(text.replace("Z", "+00:00")).date()
    except Exception:
        return None


def _format_deadline_for_ui(deadline_value, publication_date_value=None):
    deadline_date = _parse_sql_date_like(deadline_value)
    if deadline_date:
        return deadline_date.strftime("%d/%m/%Y")

    if deadline_value not in (None, "", "No disponible"):
        return str(deadline_value)

    publication_date = _parse_sql_date_like(publication_date_value)
    if publication_date:
        today = date.today()
        if today - timedelta(days=30) <= publication_date <= today:
            return "Proximamente..."

    return "No disponible"


def _safe_float(value):
    if value in (None, "", "null"):
        return None
    try:
        if isinstance(value, str):
            cleaned = value.strip()
            if "," in cleaned and "." in cleaned:
                cleaned = cleaned.replace(".", "").replace(",", ".")
            elif "," in cleaned:
                cleaned = cleaned.replace(",", ".")
            cleaned = re.sub(r"[^\d.\-]", "", cleaned)
            if cleaned == "":
                return None
            return float(cleaned)
        return float(value)
    except (TypeError, ValueError):
        return None


def _normalize_match_ratio(value):
    numeric_value = _safe_float(value)
    if numeric_value is None:
        return None
    if numeric_value > 1:
        numeric_value = numeric_value / 100.0
    return max(0.0, min(1.0, numeric_value))


def _ensure_entity_processing_status_table(cursor):
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS public.entity_processing_status (
            entity_id text PRIMARY KEY,
            user_id text NOT NULL,
            status varchar(32) NOT NULL DEFAULT 'running',
            stage varchar(64) NOT NULL DEFAULT 'queued',
            message text,
            progress numeric(5,2) NOT NULL DEFAULT 0,
            processed_items integer NOT NULL DEFAULT 0,
            total_items integer NOT NULL DEFAULT 0,
            matches_found integer NOT NULL DEFAULT 0,
            best_match_score numeric(8,6),
            first_high_match_found boolean NOT NULL DEFAULT false,
            first_high_match_grant_id text,
            first_high_match_score numeric(8,6),
            started_at timestamp without time zone,
            completed_at timestamp without time zone,
            error text,
            created_at timestamp without time zone NOT NULL DEFAULT NOW(),
            updated_at timestamp without time zone NOT NULL DEFAULT NOW()
        )
    """)
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_entity_processing_status_user
        ON public.entity_processing_status (user_id)
    """)


def _ensure_entity_documents_table(cursor):
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS public.entity_documents (
            id BIGSERIAL PRIMARY KEY,
            entity_id INTEGER NOT NULL,
            original_filename TEXT NOT NULL,
            s3_key TEXT NOT NULL,
            s3_bucket TEXT NOT NULL,
            file_size BIGINT,
            analysis_result TEXT,
            upload_date TIMESTAMP WITHOUT TIME ZONE NOT NULL DEFAULT NOW(),
            status VARCHAR(32) NOT NULL DEFAULT 'uploaded'
        )
    """)
    cursor.execute("""
        ALTER TABLE public.entity_documents
        ADD COLUMN IF NOT EXISTS document_type_code VARCHAR(64)
    """)
    cursor.execute("""
        ALTER TABLE public.entity_documents
        ADD COLUMN IF NOT EXISTS mime_type TEXT
    """)
    cursor.execute("""
        ALTER TABLE public.entity_documents
        ADD COLUMN IF NOT EXISTS uploaded_by TEXT
    """)
    cursor.execute("""
        ALTER TABLE public.entity_documents
        ADD COLUMN IF NOT EXISTS is_current BOOLEAN NOT NULL DEFAULT TRUE
    """)
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_entity_documents_entity_upload
        ON public.entity_documents (entity_id, upload_date DESC)
    """)
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_entity_documents_entity_type_current
        ON public.entity_documents (entity_id, document_type_code, is_current, upload_date DESC)
    """)
    cursor.execute("""
        UPDATE public.entity_documents
        SET is_current = TRUE
        WHERE is_current IS NULL
    """)


CRM_PIPELINE_ALLOWED_STATUSES = {
    "detectada",
    "preparando_documentacion",
    "presentada",
    "requerimiento",
    "subsanado",
    "concedida",
    "denegada",
    "justificacion",
    "terminada",
}

EXTRA_CORPORATE_DOCUMENT_CODE = "documentacion_extra"
EXTRA_CORPORATE_DOCUMENT_LABEL = "Documentacion extra"
EXTRA_CORPORATE_DOCUMENT_DESCRIPTION = (
    "Adjuntos adicionales que no formen parte del listado documental principal."
)

CORPORATE_DOCUMENT_CATALOG = [
    {
        "code": "pick_deck",
        "label": "Pick deck",
        "description": "Presentacion corporativa o deck comercial para entender mejor la propuesta de valor.",
        "display_order": 1,
    },
    {
        "code": "escrituras_empresa",
        "label": "Escrituras de la empresa",
        "description": "Documento de constitucion o escrituras inscritas de la sociedad.",
        "display_order": 2,
    },
    {
        "code": "modelo_200",
        "label": "Modelo 200",
        "description": "Ultimo impuesto de sociedades presentado.",
        "display_order": 3,
    },
    {
        "code": "vida_laboral_empresa",
        "label": "Vida Laboral de la empresa",
        "description": "Informe actualizado de la vida laboral de la empresa.",
        "display_order": 4,
    },
    {
        "code": "idc",
        "label": "IDC",
        "description": "Informe de datos para la cotizacion o IDC actualizado.",
        "display_order": 5,
    },
    {
        "code": "tarjeta_identificacion_fiscal",
        "label": "Tarjeta de identificacion fiscal",
        "description": "Tarjeta acreditativa del NIF de la empresa.",
        "display_order": 6,
    },
    {
        "code": "certificado_situacion_censal",
        "label": "Certificado de situacion censal",
        "description": "Certificado actualizado de situacion censal de la AEAT.",
        "display_order": 7,
    },
    {
        "code": "poderes",
        "label": "Poderes",
        "description": "Poderes o documento acreditativo de la representacion.",
        "display_order": 8,
    },
    {
        "code": "modelo_036",
        "label": "Modelo 036",
        "description": "Modelo 036 de alta o modificaciones censales.",
        "display_order": 9,
    },
]

DOCUMENT_TYPE_INFERENCE_RULES = (
    ("pick_deck", (r"\bpick deck\b", r"\bpitch deck\b", r"presentacion corporativa", r"deck comercial", r"corporate deck")),
    ("escrituras_empresa", (r"escritur", r"constitucion de la sociedad", r"registro mercantil")),
    ("modelo_200", (r"\bmodelo 200\b", r"impuesto de sociedades", r"\bmod200\b")),
    ("vida_laboral_empresa", (r"vida laboral", r"informe de vida laboral")),
    ("idc", (r"\bidc\b", r"informe de datos para la cotizacion", r"datos de cotizacion")),
    ("tarjeta_identificacion_fiscal", (r"tarjeta de identificacion fiscal", r"\btif\b", r"\bnif\b.*empresa")),
    ("certificado_situacion_censal", (r"situacion censal", r"certificado censal", r"censo de empresarios")),
    ("poderes", (r"\bpoderes\b", r"apoderamiento", r"documento acreditativo de la representacion")),
    ("modelo_036", (r"\bmodelo 036\b", r"\b036\b", r"declaracion censal")),
)


def _supports_multiple_current_documents(document_type_code: str | None) -> bool:
    return document_type_code == EXTRA_CORPORATE_DOCUMENT_CODE


def _sanitize_zip_segment(value: str | None, fallback: str) -> str:
    sanitized = secure_filename(value or "")
    return sanitized or fallback


def _build_unique_zip_entry_name(folder_label, original_filename, used_names):
    folder = _sanitize_zip_segment(folder_label, "otros")
    filename = _sanitize_zip_segment(original_filename, "documento")
    base_name, extension = os.path.splitext(filename)
    candidate = f"{folder}/{filename}"
    counter = 2

    while candidate in used_names:
        suffix = f"_{counter}"
        candidate = f"{folder}/{base_name}{suffix}{extension}"
        counter += 1

    used_names.add(candidate)
    return candidate


def _ensure_crm_pipeline_tables(cursor):
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS public.entity_grant_pipeline (
            id BIGSERIAL PRIMARY KEY,
            entity_id INTEGER NOT NULL,
            grant_id INTEGER NOT NULL,
            status VARCHAR(64) NOT NULL DEFAULT 'detectada',
            notes TEXT,
            created_by TEXT,
            created_at TIMESTAMP WITHOUT TIME ZONE NOT NULL DEFAULT NOW(),
            updated_at TIMESTAMP WITHOUT TIME ZONE NOT NULL DEFAULT NOW(),
            UNIQUE (entity_id, grant_id)
        )
    """)
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_entity_grant_pipeline_entity
        ON public.entity_grant_pipeline (entity_id, updated_at DESC)
    """)
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_entity_grant_pipeline_status
        ON public.entity_grant_pipeline (entity_id, status)
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS public.entity_grant_pipeline_history (
            id BIGSERIAL PRIMARY KEY,
            pipeline_id BIGINT NOT NULL,
            from_status VARCHAR(64),
            to_status VARCHAR(64) NOT NULL,
            changed_by TEXT,
            notes TEXT,
            changed_at TIMESTAMP WITHOUT TIME ZONE NOT NULL DEFAULT NOW()
        )
    """)
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_entity_grant_pipeline_history_pipeline
        ON public.entity_grant_pipeline_history (pipeline_id, changed_at DESC)
    """)


def _ensure_corporate_document_tables(cursor):
    _ensure_entity_documents_table(cursor)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS public.corporate_document_catalog (
            code VARCHAR(64) PRIMARY KEY,
            label TEXT NOT NULL,
            description TEXT,
            display_order INTEGER NOT NULL DEFAULT 0,
            is_required BOOLEAN NOT NULL DEFAULT TRUE
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS public.entity_corporate_documents (
            id BIGSERIAL PRIMARY KEY,
            entity_id INTEGER NOT NULL,
            document_type_code VARCHAR(64) NOT NULL,
            original_filename TEXT NOT NULL,
            s3_key TEXT NOT NULL,
            s3_bucket TEXT NOT NULL,
            file_size BIGINT,
            mime_type TEXT,
            status VARCHAR(32) NOT NULL DEFAULT 'uploaded',
            uploaded_by TEXT,
            upload_date TIMESTAMP WITHOUT TIME ZONE NOT NULL DEFAULT NOW(),
            expires_at DATE,
            is_current BOOLEAN NOT NULL DEFAULT TRUE
        )
    """)
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_entity_corporate_documents_entity
        ON public.entity_corporate_documents (entity_id, upload_date DESC)
    """)
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_entity_corporate_documents_entity_type
        ON public.entity_corporate_documents (entity_id, document_type_code, is_current)
    """)

    for document_type in CORPORATE_DOCUMENT_CATALOG:
        cursor.execute(
            """
            INSERT INTO public.corporate_document_catalog (code, label, description, display_order, is_required)
            VALUES (%s, %s, %s, %s, TRUE)
            ON CONFLICT (code) DO UPDATE SET
                label = EXCLUDED.label,
                description = EXCLUDED.description,
                display_order = EXCLUDED.display_order
            """,
            (
                document_type["code"],
                document_type["label"],
                document_type["description"],
                document_type["display_order"],
            ),
        )

    cursor.execute(
        """
        INSERT INTO public.corporate_document_catalog (code, label, description, display_order, is_required)
        VALUES (%s, %s, %s, %s, FALSE)
        ON CONFLICT (code) DO UPDATE SET
            label = EXCLUDED.label,
            description = EXCLUDED.description,
            display_order = EXCLUDED.display_order,
            is_required = EXCLUDED.is_required
        """,
        (
            EXTRA_CORPORATE_DOCUMENT_CODE,
            EXTRA_CORPORATE_DOCUMENT_LABEL,
            EXTRA_CORPORATE_DOCUMENT_DESCRIPTION,
            999,
        ),
    )


def _normalize_document_matching_text(value):
    if value in (None, ""):
        return ""

    normalized = unicodedata.normalize("NFKD", str(value))
    without_accents = "".join(char for char in normalized if not unicodedata.combining(char))
    lowered = without_accents.lower()
    return re.sub(r"\s+", " ", lowered).strip()


def _infer_document_type_code(filename: str | None, analysis_result: str | None = None):
    haystack = _normalize_document_matching_text(f"{filename or ''}\n{(analysis_result or '')[:12000]}")
    if not haystack:
        return None

    for code, patterns in DOCUMENT_TYPE_INFERENCE_RULES:
        if any(re.search(pattern, haystack) for pattern in patterns):
            return code

    return None


def _sync_legacy_corporate_documents(cursor):
    cursor.execute("""
        INSERT INTO public.entity_documents (
            entity_id,
            original_filename,
            s3_key,
            s3_bucket,
            file_size,
            analysis_result,
            upload_date,
            status,
            document_type_code,
            mime_type,
            uploaded_by,
            is_current
        )
        SELECT
            legacy.entity_id,
            legacy.original_filename,
            legacy.s3_key,
            legacy.s3_bucket,
            legacy.file_size,
            NULL,
            legacy.upload_date,
            COALESCE(legacy.status, 'uploaded'),
            legacy.document_type_code,
            legacy.mime_type,
            legacy.uploaded_by,
            COALESCE(legacy.is_current, TRUE)
        FROM public.entity_corporate_documents legacy
        WHERE NOT EXISTS (
            SELECT 1
            FROM public.entity_documents current_docs
            WHERE current_docs.s3_key = legacy.s3_key
        )
    """)


def _backfill_inferred_entity_document_types(cursor, entity_id=None):
    if entity_id not in (None, ""):
        cursor.execute("""
            SELECT id, original_filename, analysis_result
            FROM public.entity_documents
            WHERE entity_id = %s
              AND document_type_code IS NULL
              AND COALESCE(is_current, TRUE) IS TRUE
            ORDER BY upload_date DESC
        """, (entity_id,))
    else:
        cursor.execute("""
            SELECT id, original_filename, analysis_result
            FROM public.entity_documents
            WHERE document_type_code IS NULL
              AND COALESCE(is_current, TRUE) IS TRUE
            ORDER BY upload_date DESC
        """)

    for document_id, original_filename, analysis_result in cursor.fetchall():
        inferred_code = _infer_document_type_code(original_filename, analysis_result)
        if not inferred_code:
            continue

        cursor.execute("""
            UPDATE public.entity_documents
            SET document_type_code = %s
            WHERE id = %s
              AND document_type_code IS NULL
        """, (inferred_code, document_id))


def _extract_document_text(raw_bytes: bytes, file_extension: str, safe_filename: str):
    normalized_extension = (file_extension or "").lower()
    if not raw_bytes:
        return None

    if normalized_extension == ".txt":
        try:
            return raw_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return raw_bytes.decode("latin-1", errors="ignore")

    if normalized_extension not in {".pdf", ".docx", ".pptx"}:
        logger.info("Extraccion de texto no soportada para %s (%s)", safe_filename, normalized_extension or "sin_extension")
        return None

    try:
        extracted_text = extract_text(io.BytesIO(raw_bytes), normalized_extension)
        return extracted_text or None
    except Exception as exc:
        logger.warning("No se pudo extraer texto de %s: %s", safe_filename, exc)
        return None


def _store_entity_document(cursor, entity_id, uploaded_file, uploaded_by=None, document_type_code=None):
    if not uploaded_file or uploaded_file.filename == '':
        raise ValueError("Debes adjuntar un archivo")

    raw_bytes = uploaded_file.read()
    if not raw_bytes:
        raise ValueError("El archivo esta vacio")

    safe_filename = secure_filename(uploaded_file.filename or f"{document_type_code or 'documento'}.bin")
    file_extension = os.path.splitext(safe_filename)[1].lower()
    analysis_result = _extract_document_text(raw_bytes, file_extension, safe_filename)

    if document_type_code:
        unique_filename = f"entity-documents/{entity_id}/{document_type_code}/{uuid.uuid4()}{file_extension}"
    else:
        unique_filename = f"entity-documents/{entity_id}/{uuid.uuid4()}{file_extension}"

    metadata = {
        'entity_id': str(entity_id),
        'original_filename': safe_filename,
        'upload_timestamp': datetime.now().isoformat(),
    }
    if document_type_code:
        metadata['document_type_code'] = document_type_code

    s3_client.upload_fileobj(
        Fileobj=io.BytesIO(raw_bytes),
        Bucket=S3_BUCKET_NAME,
        Key=unique_filename,
        ExtraArgs={'Metadata': metadata},
    )

    if document_type_code and not _supports_multiple_current_documents(document_type_code):
        cursor.execute("""
            UPDATE public.entity_documents
            SET is_current = FALSE
            WHERE entity_id = %s
              AND document_type_code = %s
              AND COALESCE(is_current, TRUE) IS TRUE
        """, (entity_id, document_type_code))

    cursor.execute("""
        INSERT INTO public.entity_documents (
            entity_id,
            original_filename,
            s3_key,
            s3_bucket,
            file_size,
            analysis_result,
            upload_date,
            status,
            document_type_code,
            mime_type,
            uploaded_by,
            is_current
        )
        VALUES (%s, %s, %s, %s, %s, %s, NOW(), 'uploaded', %s, %s, %s, TRUE)
        RETURNING id, upload_date
    """, (
        entity_id,
        safe_filename,
        unique_filename,
        S3_BUCKET_NAME,
        len(raw_bytes),
        analysis_result,
        document_type_code,
        uploaded_file.mimetype,
        uploaded_by,
    ))
    inserted_row = cursor.fetchone()

    return {
        "id": inserted_row[0],
        "filename": safe_filename,
        "s3_key": unique_filename,
        "s3_bucket": S3_BUCKET_NAME,
        "file_size": len(raw_bytes),
        "mime_type": uploaded_file.mimetype,
        "status": "uploaded",
        "upload_date": _serialize_datetimeish(inserted_row[1]),
        "document_type_code": document_type_code,
        "analysis_result": analysis_result,
    }


def _resolve_entity_id_for_user(cursor, user_id, preferred_entity_id=None):
    if preferred_entity_id not in (None, ""):
        cursor.execute("""
            SELECT ue.entity_id
            FROM user_entities ue
            WHERE ue.user_id = %s
              AND ue.entity_id = %s
            LIMIT 1
        """, (user_id, preferred_entity_id))
        row = cursor.fetchone()
        if row:
            return row[0]
        return None

    selected_entity_id = _get_selected_entity_id_for_user(cursor, user_id)
    if selected_entity_id:
        return selected_entity_id

    cursor.execute("""
        SELECT ue.entity_id
        FROM user_entities ue
        WHERE ue.user_id = %s
        ORDER BY ue.is_selected DESC, ue.updated_at DESC NULLS LAST, ue.created_at DESC
        LIMIT 1
    """, (user_id,))
    row = cursor.fetchone()
    return row[0] if row else None


def _record_pipeline_history(cursor, pipeline_id, from_status, to_status, user_id, notes=None):
    cursor.execute(
        """
        INSERT INTO public.entity_grant_pipeline_history (
            pipeline_id,
            from_status,
            to_status,
            changed_by,
            notes,
            changed_at
        )
        VALUES (%s, %s, %s, %s, %s, NOW())
        """,
        (pipeline_id, from_status, to_status, user_id, notes),
    )


def _serialize_datetimeish(value):
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if value in (None, ""):
        return None
    return str(value)


def _format_amount_like_ui(value):
    if value in (None, "", "No especificado"):
        return "No especificado"

    if isinstance(value, (int, float)):
        formatted = f"{float(value):,.2f}"
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
        if formatted.endswith(",00"):
            formatted = formatted[:-3]
        return f"{formatted}€"

    text = str(value).strip()
    if not text:
        return "No especificado"
    if "€" in text or "%" in text:
        return text
    if re.search(r"[A-Za-zÁÉÍÓÚáéíóúÑñ]", text):
        return text
    if re.fullmatch(r"[\d.,\s\-]+", text):
        return f"{text}€"
    return text


def _fetch_crm_pipeline_grants(cursor, entity_id):
    _ensure_crm_pipeline_tables(cursor)

    importe_beneficiario_expr = (
        "g.importe_beneficiario AS importe_beneficiario"
        if _grant_has_column(cursor, "importe_beneficiario")
        else "NULL::text AS importe_beneficiario"
    )
    fecha_inicio_expr = (
        "g.fecha_inicio_solicitud AS fecha_inicio_solicitud"
        if _grant_has_column(cursor, "fecha_inicio_solicitud")
        else "NULL::date AS fecha_inicio_solicitud"
    )
    region_expr = (
        "g.region_impacto AS region_impacto"
        if _grant_has_column(cursor, "region_impacto")
        else "NULL::text AS region_impacto"
    )
    finalidad_expr = (
        "g.finalidad AS finalidad"
        if _grant_has_column(cursor, "finalidad")
        else "NULL::text AS finalidad"
    )

    cursor.execute(
        f"""
        SELECT
            p.id,
            p.grant_id,
            p.status,
            p.notes,
            p.created_at,
            p.updated_at,
            g.titulo_corto,
            g.presupuesto,
            {importe_beneficiario_expr},
            {fecha_inicio_expr},
            g.fecha_finalizacion,
            g.fecha_de_publicacion,
            {region_expr},
            {finalidad_expr}
        FROM public.entity_grant_pipeline p
        JOIN public.grants g ON g.id = p.grant_id
        WHERE p.entity_id = %s
        ORDER BY
            CASE WHEN g.fecha_finalizacion IS NULL THEN 1 ELSE 0 END,
            g.fecha_finalizacion ASC NULLS LAST,
            p.updated_at DESC
        """,
        (entity_id,),
    )

    records = []
    for row in cursor.fetchall():
        deadline_display = _format_deadline_for_ui(row[10], row[11])
        records.append({
            "id": row[0],
            "grant_id": str(row[1]),
            "status": row[2],
            "notes": row[3],
            "created_at": _serialize_datetimeish(row[4]),
            "updated_at": _serialize_datetimeish(row[5]),
            "titulo_corto": row[6] or "Sin titulo",
            "fondos_totales": _format_amount_like_ui(row[7]),
            "importe_beneficiario": _format_amount_like_ui(row[8] if row[8] not in (None, "") else row[7]),
            "fecha_inicio": _serialize_datetimeish(row[9]),
            "fecha_inicio_display": (
                _parse_sql_date_like(row[9]).strftime("%d/%m/%Y")
                if _parse_sql_date_like(row[9]) else "No disponible"
            ),
            "fecha_limite": _serialize_datetimeish(row[10]),
            "fecha_limite_display": deadline_display,
            "region": row[12] or "No especificado",
            "finalidad": row[13] or "No especificado",
        })

    return records


def _coerce_progress(value, fallback=0.0):
    try:
        return max(0.0, min(float(value), 100.0))
    except (TypeError, ValueError):
        return fallback


def _set_entity_processing_status(
    entity_id,
    user_id,
    status,
    stage,
    message=None,
    progress=0.0,
    processed_items=0,
    total_items=0,
    matches_found=0,
    best_match_score=None,
    first_high_match_found=False,
    first_high_match_grant_id=None,
    first_high_match_score=None,
    started_at=None,
    completed_at=None,
    error=None,
):
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        _ensure_entity_processing_status_table(cursor)

        now = datetime.utcnow()
        safe_progress = _coerce_progress(progress, fallback=0.0)

        cursor.execute(
            """
            INSERT INTO public.entity_processing_status (
                entity_id,
                user_id,
                status,
                stage,
                message,
                progress,
                processed_items,
                total_items,
                matches_found,
                best_match_score,
                first_high_match_found,
                first_high_match_grant_id,
                first_high_match_score,
                started_at,
                completed_at,
                error,
                created_at,
                updated_at
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (entity_id) DO UPDATE SET
                user_id = EXCLUDED.user_id,
                status = EXCLUDED.status,
                stage = EXCLUDED.stage,
                message = EXCLUDED.message,
                progress = EXCLUDED.progress,
                processed_items = EXCLUDED.processed_items,
                total_items = EXCLUDED.total_items,
                matches_found = EXCLUDED.matches_found,
                best_match_score = EXCLUDED.best_match_score,
                first_high_match_found = EXCLUDED.first_high_match_found,
                first_high_match_grant_id = EXCLUDED.first_high_match_grant_id,
                first_high_match_score = EXCLUDED.first_high_match_score,
                started_at = COALESCE(public.entity_processing_status.started_at, EXCLUDED.started_at),
                completed_at = EXCLUDED.completed_at,
                error = EXCLUDED.error,
                updated_at = EXCLUDED.updated_at
            """,
            (
                str(entity_id),
                str(user_id),
                status,
                stage,
                message,
                safe_progress,
                int(processed_items or 0),
                int(total_items or 0),
                int(matches_found or 0),
                best_match_score,
                bool(first_high_match_found),
                first_high_match_grant_id,
                first_high_match_score,
                started_at or now,
                completed_at,
                error,
                now,
                now,
            ),
        )
        connection.commit()
    except Exception as status_error:
        logger.warning(f"No se pudo actualizar estado de procesamiento para entity_id={entity_id}: {status_error}")
        if connection:
            try:
                connection.rollback()
            except Exception:
                pass
    finally:
        if cursor:
            cursor.close()
        if connection:
            connection.close()


def _parse_bdns_concessions_response(payload):
    if not isinstance(payload, dict):
        return []

    candidate_lists = [
        payload.get("content"),
        payload.get("concesiones"),
        (payload.get("resultadoConcesionDTO") or {}).get("concesiones") if isinstance(payload.get("resultadoConcesionDTO"), dict) else None,
        (payload.get("data") or {}).get("concesiones") if isinstance(payload.get("data"), dict) else None,
        payload.get("items"),
    ]

    for candidate in candidate_lists:
        if isinstance(candidate, list):
            return candidate

    return []


def _fetch_bdns_tercero_id_by_cif(cif: str, ambito: str = "C"):
    normalized_cif = _normalize_cif_nif(cif)
    if not normalized_cif:
        return {"persona_id": None, "terceros": []}

    headers = {
        "Accept": "application/json",
        "User-Agent": "Grantify/1.0 (+https://grantify.local)",
    }
    params = {
        "ambito": ambito or "C",
        "busqueda": normalized_cif,
        "vpd": "GE",
    }

    response = requests.get(
        "https://www.pap.hacienda.gob.es/bdnstrans/api/terceros",
        headers=headers,
        params=params,
        timeout=20,
    )
    response.raise_for_status()

    payload = response.json()
    terceros = payload.get("terceros") if isinstance(payload, dict) else None
    terceros = terceros if isinstance(terceros, list) else []

    persona_id = None
    for tercero in terceros:
        if not isinstance(tercero, dict):
            continue
        if tercero.get("id") not in (None, ""):
            persona_id = tercero.get("id")
            break

    return {"persona_id": persona_id, "terceros": terceros}


def _fetch_bdns_concesiones_by_cif(cif: str, limit: int = 10):
    normalized_cif = _normalize_cif_nif(cif)
    if not normalized_cif:
        return {"items": [], "total": 0, "warning": "La entidad no tiene CIF/NIF informado."}

    safe_limit = max(1, min(limit, 50))
    headers = {
        "Accept": "application/json",
        "User-Agent": "Grantify/1.0 (+https://grantify.local)",
    }
    payload = None
    persona_id = None

    # Vía estable documentada: búsqueda directa por nifCif en BDNS (pap.hacienda).
    try:
        direct_params = {
            "nifCif": normalized_cif,
            "page": 0,
            "pageSize": safe_limit,
            "order": "fechaConcesion",
            "direccion": "desc",
        }
        response = requests.get(
            "https://www.pap.hacienda.gob.es/bdnstrans/api/concesiones/busqueda",
            headers=headers,
            params=direct_params,
            timeout=25,
        )
        response.raise_for_status()
        payload = response.json()
    except Exception as direct_error:
        logger.warning(
            f"Fallo consulta directa BDNS concesiones por nifCif={normalized_cif}. "
            f"Se intenta fallback legacy: {direct_error}"
        )

    # Fallback legacy (tercero -> beneficiario) para compatibilidad.
    if payload is None:
        tercero_lookup = _fetch_bdns_tercero_id_by_cif(normalized_cif)
        persona_id = tercero_lookup.get("persona_id")
        if not persona_id:
            return {
                "items": [],
                "total": 0,
                "persona_id": None,
                "warning": f"No se ha encontrado un tercero en BDNS para el CIF/NIF {normalized_cif}.",
            }

        concesiones_params = {
            "vpd": "GE",
            "beneficiario": persona_id,
            "page": 0,
            "pageSize": safe_limit,
            "order": "fechaConcesion",
            "direccion": "desc",
        }

        response = requests.get(
            "https://www.pap.hacienda.gob.es/bdnstrans/api/concesiones/busqueda",
            headers=headers,
            params=concesiones_params,
            timeout=25,
        )
        response.raise_for_status()
        payload = response.json()

    raw_items = _parse_bdns_concessions_response(payload)

    items = []
    for item in raw_items:
        if not isinstance(item, dict):
            continue

        importe = _safe_float(item.get("importe"))
        if importe is None:
            importe = _safe_float(item.get("ayudaEquivalente"))

        items.append({
            "id": item.get("id"),
            "referencia": item.get("codConcesion"),
            "cod_concesion": item.get("codConcesion"),
            "titulo": item.get("convocatoria") or item.get("descripcionCooficial") or "Concesión sin título",
            "convocatoria": item.get("convocatoria"),
            "descripcion_cooficial": item.get("descripcionCooficial"),
            "fecha_concesion": item.get("fechaConcesion"),
            "beneficiario": item.get("beneficiario"),
            "instrumento": item.get("instrumento"),
            "importe": importe,
            "ayuda_equivalente": _safe_float(item.get("ayudaEquivalente")),
            "nivel1": item.get("nivel1"),
            "nivel2": item.get("nivel2"),
            "nivel3": item.get("nivel3"),
            "region": item.get("nivel2"),
            "organo_convocante": item.get("nivel3"),
            "administracion": item.get("nivel1"),
            "url": item.get("urlBR"),
            "url_br": item.get("urlBR"),
            "id_convocatoria": item.get("idConvocatoria"),
            "numero_convocatoria": item.get("numeroConvocatoria"),
            "id_persona": item.get("idPersona"),
            "codigo_invente": item.get("codigoInvente"),
        })

    total = None
    if isinstance(payload, dict):
        total = payload.get("totalElements")
    if total is None:
        total = len(items)

    if persona_id is None:
        for item in items:
            if item.get("id_persona") not in (None, ""):
                persona_id = item.get("id_persona")
                break

    return {
        "items": items[:safe_limit],
        "total": total,
        "persona_id": persona_id,
        "warning": None,
    }


def _normalize_bdns_convocatoria_number(value):
    if value in (None, ""):
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value).strip()
    if isinstance(value, int):
        return str(value)

    text = str(value).strip()
    digits = re.sub(r"\D", "", text)
    return digits or text


def _split_bdns_beneficiario(raw_value):
    text = (str(raw_value).strip() if raw_value not in (None, "") else "")
    if not text:
        return {"beneficiario_cif": None, "beneficiario_nombre": None, "beneficiario": ""}

    parts = text.split(None, 1)
    if len(parts) == 1:
        return {"beneficiario_cif": None, "beneficiario_nombre": text, "beneficiario": text}

    cif_candidate, rest = parts[0].strip(), parts[1].strip()
    if re.fullmatch(r"[A-Za-z0-9]+", cif_candidate):
        return {
            "beneficiario_cif": cif_candidate,
            "beneficiario_nombre": rest or None,
            "beneficiario": text,
        }

    return {"beneficiario_cif": None, "beneficiario_nombre": text, "beneficiario": text}


def _fetch_bdns_concesiones_by_convocatoria(numero_convocatoria: str, page_size: int = 50):
    normalized_numero = _normalize_bdns_convocatoria_number(numero_convocatoria)
    if not normalized_numero:
        return {
            "items": [],
            "total": 0,
            "warning": "La subvención no tiene código BDNS de convocatoria disponible.",
        }

    safe_page_size = max(1, min(page_size, 200))
    headers = {
        "Accept": "application/json",
        "User-Agent": "Grantify/1.0 (+https://grantify.local)",
    }
    payload = None

    # Vía documentada oficial (simple): numeroConvocatoria + pageSize
    try:
        documented_params = {
            "numeroConvocatoria": normalized_numero,
            "pageSize": safe_page_size,
            "page": 0,
        }
        response = requests.get(
            "https://www.pap.hacienda.gob.es/bdnstrans/api/concesiones/busqueda",
            headers=headers,
            params=documented_params,
            timeout=25,
        )
        response.raise_for_status()
        payload = response.json()
    except Exception as documented_error:
        logger.warning(
            f"Fallo consulta BDNS concesiones por convocatoria (documentada) "
            f"numeroConvocatoria={normalized_numero}. Se intenta fallback extendido: {documented_error}"
        )

    # Fallback con parámetros extendidos (compatibilidad / orden explícito)
    if payload is None:
        fallback_params = {
            "vpd": "GE",
            "numeroConvocatoria": normalized_numero,
            "page": 0,
            "pageSize": safe_page_size,
            "order": "fechaConcesion",
            "direccion": "desc",
        }

        response = requests.get(
            "https://www.pap.hacienda.gob.es/bdnstrans/api/concesiones/busqueda",
            headers=headers,
            params=fallback_params,
            timeout=25,
        )
        response.raise_for_status()
        payload = response.json()
    raw_items = _parse_bdns_concessions_response(payload)
    items = []
    for item in raw_items:
        if not isinstance(item, dict):
            continue

        importe = _safe_float(item.get("importe"))
        if importe is None:
            importe = _safe_float(item.get("ayudaEquivalente"))

        beneficiario_split = _split_bdns_beneficiario(item.get("beneficiario"))
        items.append({
            "id": item.get("id"),
            "referencia": item.get("codConcesion"),
            "fecha_concesion": item.get("fechaConcesion"),
            "beneficiario": beneficiario_split.get("beneficiario"),
            "beneficiario_cif": beneficiario_split.get("beneficiario_cif"),
            "beneficiario_nombre": beneficiario_split.get("beneficiario_nombre"),
            "instrumento": item.get("instrumento"),
            "importe": importe,
            "ayuda_equivalente": _safe_float(item.get("ayudaEquivalente")),
            "url": item.get("urlBR"),
            "url_br": item.get("urlBR"),
            "numero_convocatoria": item.get("numeroConvocatoria"),
            "id_convocatoria": item.get("idConvocatoria"),
            "convocatoria": item.get("convocatoria") or item.get("descripcionCooficial"),
            "nivel1": item.get("nivel1"),
            "nivel2": item.get("nivel2"),
            "nivel3": item.get("nivel3"),
            "region": item.get("nivel2"),
            "organo_convocante": item.get("nivel3"),
        })

    total = payload.get("totalElements") if isinstance(payload, dict) else None
    if total is None:
        total = len(items)

    return {
        "items": items,
        "total": total,
        "warning": None,
        "numero_convocatoria": normalized_numero,
    }


def _safe_parse_iso_date(value):
    if not value:
        return None
    try:
        return datetime.strptime(str(value), "%Y-%m-%d").date()
    except (TypeError, ValueError):
        return None


def _date_years_ago(base_date: date, years: int) -> date:
    try:
        return base_date.replace(year=base_date.year - years)
    except ValueError:
        # Ajuste para 29 de febrero
        return base_date.replace(month=2, day=28, year=base_date.year - years)


def _fetch_bdns_minimis_by_cif(cif: str, page_size: int = 100, max_pages: int = 20):
    normalized_cif = _normalize_cif_nif(cif)
    if not normalized_cif:
        return {
            "items": [],
            "total": 0,
            "persona_id": None,
            "warning": "La entidad no tiene CIF/NIF informado.",
            "total_last_3_years_amount": 0.0,
            "total_amount": 0.0,
            "count_last_3_years": 0,
        }

    safe_page_size = max(1, min(page_size, 200))
    safe_max_pages = max(1, min(max_pages, 50))
    headers = {
        "Accept": "application/json",
        "User-Agent": "Grantify/1.0 (+https://grantify.local)",
    }

    def _build_minimis_result(all_items, total_elements, persona_id, warning, bdnstrans_notice):
        today = date.today()
        cutoff_date = _date_years_ago(today, 3)
        total_amount = 0.0
        total_last_3_years_amount = 0.0
        count_last_3_years = 0

        for item in all_items:
            importe = _safe_float(item.get("importe"))
            if importe is not None:
                total_amount += importe

            fecha_concesion = _safe_parse_iso_date(item.get("fecha_concesion"))
            if fecha_concesion and fecha_concesion >= cutoff_date:
                if importe is not None:
                    total_last_3_years_amount += importe
                count_last_3_years += 1

        return {
            "items": all_items,
            "total": total_elements if total_elements is not None else len(all_items),
            "persona_id": persona_id,
            "warning": warning,
            "total_last_3_years_amount": total_last_3_years_amount,
            "total_amount": total_amount,
            "count_last_3_years": count_last_3_years,
            "cutoff_date": cutoff_date.isoformat(),
            "bdnstrans_notice": bdnstrans_notice,
        }

    # Vía estable documentada: búsqueda directa por nifCif en BDNS (pap.hacienda).
    try:
        all_items = []
        total_elements = None
        warning = None
        bdnstrans_notice = None
        persona_id = None

        for page in range(safe_max_pages):
            minimis_params = {
                "nifCif": normalized_cif,
                "page": page,
                "pageSize": safe_page_size,
                "order": "fechaConcesion",
                "direccion": "desc",
            }

            response = requests.get(
                "https://www.pap.hacienda.gob.es/bdnstrans/api/minimis/busqueda",
                headers=headers,
                params=minimis_params,
                timeout=25,
            )
            response.raise_for_status()
            payload = response.json()

            if isinstance(payload, dict):
                if bdnstrans_notice is None:
                    bdnstrans_notice = payload.get("advertencia")
                if payload.get("totalElements") is not None:
                    total_elements = payload.get("totalElements")

            raw_items = _parse_bdns_concessions_response(payload)
            if not raw_items:
                break

            for item in raw_items:
                if not isinstance(item, dict):
                    continue

                if persona_id in (None, "") and item.get("idPersona") not in (None, ""):
                    persona_id = item.get("idPersona")

                importe = _safe_float(item.get("ayudaEquivalente"))
                if importe is None:
                    importe = _safe_float(item.get("importe"))

                codigo_concesion = item.get("codigoConcesion") or item.get("codConcesion")
                convocante = item.get("convocante")
                reglamento = item.get("reglamento")
                titulo = item.get("convocatoria")
                if not titulo:
                    titulo = f"Concesión minimis {codigo_concesion}" if codigo_concesion else "Concesión minimis"

                all_items.append({
                    "id": item.get("idConcesion") or item.get("id"),
                    "referencia": codigo_concesion,
                    "codigo_concesion": codigo_concesion,
                    "titulo": titulo,
                    "fecha_concesion": item.get("fechaConcesion"),
                    "fecha_registro": item.get("fechaRegistro"),
                    "beneficiario": item.get("beneficiario"),
                    "convocante": convocante,
                    "organo_convocante": convocante,
                    "reglamento": reglamento,
                    "instrumento": item.get("instrumento"),
                    "sector_actividad": item.get("sectorActividad"),
                    "sector_producto": item.get("sectorProducto"),
                    "importe": importe,
                    "ayuda_equivalente": importe,
                    "numero_convocatoria": item.get("numeroConvocatoria"),
                    "id_convocatoria": item.get("idConvocatoria"),
                    "id_persona": item.get("idPersona"),
                    "url": None,
                    "url_br": None,
                })

            is_last = bool(payload.get("last")) if isinstance(payload, dict) else False
            if is_last:
                break
            if total_elements is not None and len(all_items) >= int(total_elements):
                break

        return _build_minimis_result(
            all_items=all_items,
            total_elements=total_elements,
            persona_id=persona_id,
            warning=warning,
            bdnstrans_notice=bdnstrans_notice,
        )
    except Exception as direct_error:
        logger.warning(
            f"Fallo consulta directa BDNS minimis por nifCif={normalized_cif}. "
            f"Se intenta fallback legacy: {direct_error}"
        )

    tercero_lookup = _fetch_bdns_tercero_id_by_cif(normalized_cif, ambito="M")
    persona_id = tercero_lookup.get("persona_id")
    if not persona_id:
        return {
            "items": [],
            "total": 0,
            "persona_id": None,
            "warning": f"No se ha encontrado un tercero de minimis en BDNS para el CIF/NIF {normalized_cif}.",
            "total_last_3_years_amount": 0.0,
            "total_amount": 0.0,
            "count_last_3_years": 0,
        }

    all_items = []
    total_elements = None
    warning = None
    bdnstrans_notice = None

    for page in range(safe_max_pages):
        minimis_params = {
            "vpd": "GE",
            "beneficiario": persona_id,
            "page": page,
            "pageSize": safe_page_size,
            "order": "fechaConcesion",
            "direccion": "desc",
        }

        response = requests.get(
            "https://www.pap.hacienda.gob.es/bdnstrans/api/minimis/busqueda",
            headers=headers,
            params=minimis_params,
            timeout=25,
        )
        response.raise_for_status()
        payload = response.json()

        if isinstance(payload, dict):
            if bdnstrans_notice is None:
                bdnstrans_notice = payload.get("advertencia")
            if payload.get("totalElements") is not None:
                total_elements = payload.get("totalElements")

        raw_items = _parse_bdns_concessions_response(payload)
        if not raw_items:
            break

        for item in raw_items:
            if not isinstance(item, dict):
                continue

            importe = _safe_float(item.get("ayudaEquivalente"))
            if importe is None:
                importe = _safe_float(item.get("importe"))

            codigo_concesion = item.get("codigoConcesion") or item.get("codConcesion")
            convocante = item.get("convocante")
            reglamento = item.get("reglamento")
            titulo = item.get("convocatoria")
            if not titulo:
                titulo = f"Concesión minimis {codigo_concesion}" if codigo_concesion else "Concesión minimis"

            all_items.append({
                "id": item.get("idConcesion") or item.get("id"),
                "referencia": codigo_concesion,
                "codigo_concesion": codigo_concesion,
                "titulo": titulo,
                "fecha_concesion": item.get("fechaConcesion"),
                "fecha_registro": item.get("fechaRegistro"),
                "beneficiario": item.get("beneficiario"),
                "convocante": convocante,
                "organo_convocante": convocante,
                "reglamento": reglamento,
                "instrumento": item.get("instrumento"),
                "sector_actividad": item.get("sectorActividad"),
                "sector_producto": item.get("sectorProducto"),
                "importe": importe,
                "ayuda_equivalente": importe,
                "numero_convocatoria": item.get("numeroConvocatoria"),
                "id_convocatoria": item.get("idConvocatoria"),
                "id_persona": item.get("idPersona"),
                "url": None,
                "url_br": None,
            })

        is_last = bool(payload.get("last")) if isinstance(payload, dict) else False
        if is_last:
            break
        if total_elements is not None and len(all_items) >= int(total_elements):
            break

    return _build_minimis_result(
        all_items=all_items,
        total_elements=total_elements,
        persona_id=persona_id,
        warning=warning,
        bdnstrans_notice=bdnstrans_notice,
    )


def _extract_first_field(item, *keys):
    if not isinstance(item, dict):
        return None
    for key in keys:
        value = item.get(key)
        if value not in (None, ""):
            return value
    return None


def _normalize_bdns_concession_item(item):
    if not isinstance(item, dict):
        return None

    titulo = _extract_first_field(
        item,
        "titulo",
        "denominacion",
        "descripcion",
        "objeto",
        "nombreSubvencion",
        "descripcionSubvencion",
    )
    organismo = _extract_first_field(
        item,
        "organoConcedente",
        "organismo",
        "organo",
        "administracionConcedente",
    )
    fecha_concesion = _extract_first_field(
        item,
        "fechaConcesion",
        "fecha_concesion",
        "fechaResolucion",
        "fecha",
    )
    importe_raw = _extract_first_field(
        item,
        "importeConcedido",
        "importe",
        "importe_concedido",
        "cuantia",
        "monto",
    )
    finalidad = _extract_first_field(
        item,
        "finalidad",
        "descripcionFinalidad",
        "proyecto",
        "actuacion",
    )
    region = _extract_first_field(
        item,
        "comunidadAutonoma",
        "region",
        "ambitoTerritorial",
        "ccaa",
    )
    referencia = _extract_first_field(
        item,
        "id",
        "identificador",
        "codigo",
        "codigoConcesion",
        "numeroExpediente",
    )
    detalle_url = _extract_first_field(item, "url", "detalleUrl", "enlace")

    if not titulo:
        for candidate in item.values():
            if isinstance(candidate, str) and len(candidate.strip()) > 12:
                titulo = candidate.strip()
                break

    if not titulo:
        return None

    importe_num = _safe_float(importe_raw)
    if importe_num is not None:
        importe_value = importe_num
    elif importe_raw not in (None, ""):
        importe_value = str(importe_raw)
    else:
        importe_value = None

    return {
        "id": str(referencia) if referencia not in (None, "") else None,
        "referencia": str(referencia) if referencia not in (None, "") else None,
        "titulo": str(titulo).strip(),
        "organismo": str(organismo).strip() if organismo else None,
        "fecha_concesion": str(fecha_concesion).strip() if fecha_concesion else None,
        "importe": importe_value,
        "finalidad": str(finalidad).strip() if finalidad else None,
        "region": str(region).strip() if region else None,
        "url": str(detalle_url).strip() if detalle_url else None,
    }


def _fetch_bdns_concessions_by_nif(nif, limit=20):
    clean_nif = _normalize_cif_nif(nif)
    if not clean_nif:
        return {"awards": [], "source": None, "warning": "La entidad no tiene CIF/NIF informado."}

    configured_url = (os.getenv("BDNS_CONCESIONES_API_URL") or "").strip()
    candidate_urls = [configured_url] if configured_url else []
    candidate_urls.extend([
        "https://www.pap.hacienda.gob.es/bdnstrans/api/concesiones",
        "https://www.pap.hacienda.gob.es/bdnstrans/api/concesion",
    ])

    param_variants = [
        {"beneficiario": clean_nif, "page": 0, "size": limit},
        {"beneficiario": clean_nif, "limit": limit},
        {"cif": clean_nif, "page": 0, "size": limit},
        {"nif": clean_nif, "page": 0, "size": limit},
        {"nifBeneficiario": clean_nif, "page": 0, "size": limit},
    ]

    last_error = None

    for base_url in candidate_urls:
        for params in param_variants:
            try:
                req = Request(
                    f"{base_url}?{urlencode(params)}",
                    headers={
                        "Accept": "application/json, text/plain, */*",
                        "User-Agent": "Grantify/1.0",
                    },
                )
                with urlopen(req, timeout=12) as response:
                    raw = response.read()
                    content_type = response.headers.get("Content-Type", "")

                if "json" not in content_type.lower() and not raw.strip().startswith((b"{", b"[")):
                    last_error = f"Respuesta no JSON en {base_url}"
                    continue

                payload = json.loads(raw.decode("utf-8"))
                rows = _parse_bdns_concessions_response(payload)
                awards = []
                for row in rows:
                    normalized = _normalize_bdns_concession_item(row)
                    if normalized:
                        awards.append(normalized)
                    if len(awards) >= limit:
                        break

                return {"awards": awards, "source": base_url, "warning": None}
            except HTTPError as e:
                last_error = f"HTTP {e.code} ({base_url})"
            except URLError as e:
                last_error = f"URLError {e.reason} ({base_url})"
            except json.JSONDecodeError:
                last_error = f"JSON inválido ({base_url})"
            except Exception as e:
                last_error = str(e)

    logger.warning(f"No se pudo consultar BDNS para NIF {clean_nif}: {last_error}")
    return {
        "awards": [],
        "source": None,
        "warning": "No se pudo consultar BDNS en este momento. Si conoces el endpoint exacto, configúralo en BDNS_CONCESIONES_API_URL.",
    }


def get_grant_extra_metadata(cursor, grant_id):
    """
    Lee metadatos opcionales de grants sin romper si aún no existen columnas (p.ej. Documentacion).
    """
    metadata = {
        "codigobdns": None,
        "documentacion": None,
    }

    try:
        cursor.execute("SELECT codigobdns FROM grants WHERE id = %s LIMIT 1", (grant_id,))
        row = cursor.fetchone()
        metadata["codigobdns"] = row[0] if row else None
    except Exception as e:
        logger.warning(f"No se pudo leer codigobdns para grant {grant_id}: {e}")

    try:
        if _grant_has_column(cursor, "Documentacion"):
            cursor.execute('SELECT "Documentacion" FROM grants WHERE id = %s LIMIT 1', (grant_id,))
            row = cursor.fetchone()
            metadata["documentacion"] = row[0] if row else None
    except Exception as e:
        logger.warning(f"No se pudo leer Documentacion para grant {grant_id}: {e}")

    return metadata


def html_to_text_for_export(raw_html: str) -> str:
    if not raw_html:
        return ""
    text = pyhtml.unescape(raw_html)
    text = re.sub(r'<\s*br\s*/?>', '\n', text, flags=re.I)
    text = re.sub(r'</\s*(p|div|h1|h2|h3|h4|h5|h6|li)\s*>', '\n', text, flags=re.I)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def parse_documentacion_items(raw_documentacion):
    if raw_documentacion is None:
        return []

    if isinstance(raw_documentacion, list):
        return [str(item).strip() for item in raw_documentacion if str(item).strip()]

    if isinstance(raw_documentacion, dict):
        items = []
        for key, value in raw_documentacion.items():
            if isinstance(value, list):
                for v in value:
                    text = str(v).strip()
                    if text:
                        items.append(text)
            else:
                text = str(value).strip()
                if text:
                    items.append(f"{key}: {text}")
        return items

    raw_str = str(raw_documentacion).strip()
    if not raw_str:
        return []

    # Intentar JSON serializado
    try:
        parsed = json.loads(raw_str)
        return parse_documentacion_items(parsed)
    except Exception:
        pass

    # Intentar HTML -> líneas
    if "<" in raw_str and ">" in raw_str:
        text = html_to_text_for_export(raw_str)
        return [line.strip("-• \t") for line in text.splitlines() if line.strip()]

    # Texto plano (separado por líneas / ; )
    if ";" in raw_str:
        return [part.strip() for part in raw_str.split(";") if part.strip()]
    return [line.strip() for line in raw_str.splitlines() if line.strip()]


def get_grants_docs_bucket_name():
    try:
        ssm = boto3.client("ssm", region_name="eu-central-1")
        resp = ssm.get_parameter(Name="/grantify/s3/docssubvenciones", WithDecryption=False)
        return resp["Parameter"]["Value"]
    except ClientError as e:
        logger.warning(f"No se pudo leer bucket de subvenciones desde SSM: {e}")
        return os.getenv("S3_BUCKET") or os.getenv("GRANTIFY_GRANTS_S3_BUCKET")


def find_boe_s3_object_for_grant(codigobdns: str):
    if not codigobdns:
        return None

    bucket = get_grants_docs_bucket_name()
    if not bucket:
        return None

    prefixes_to_try = [f"{codigobdns}/", str(codigobdns)]
    candidates = []

    for prefix in prefixes_to_try:
        try:
            resp = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix, MaxKeys=200)
            for obj in resp.get("Contents", []):
                key = obj.get("Key", "")
                key_lower = key.lower()
                score = 0
                if "boe" in key_lower:
                    score += 10
                if "normativa" in key_lower or "bases" in key_lower or "convocatoria" in key_lower:
                    score += 5
                if key_lower.endswith(".pdf"):
                    score += 3
                if score > 0:
                    candidates.append((score, key))
        except Exception as e:
            logger.warning(f"Error buscando normativa en S3 para prefix {prefix}: {e}")

    if not candidates:
        return None

    candidates.sort(key=lambda item: item[0], reverse=True)
    return {"bucket": bucket, "key": candidates[0][1]}


_grant_chat_openai_client = None
_grant_chat_openai_lock = threading.Lock()


def _get_openai_client_for_grant_chat():
    global _grant_chat_openai_client
    if _grant_chat_openai_client is not None:
        return _grant_chat_openai_client

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY no configurada")

    with _grant_chat_openai_lock:
        if _grant_chat_openai_client is None:
            _grant_chat_openai_client = OpenAI(
                api_key=api_key,
                organization=os.getenv("OPENAI_ORG_ID"),
            )
    return _grant_chat_openai_client


def _format_amount_for_grant_chat(value) -> str:
    if value in (None, ""):
        return "No disponible"
    if isinstance(value, str) and "€" in value:
        return value
    numeric = _safe_float(value)
    if numeric is None:
        return str(value)
    formatted = f"{numeric:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    if formatted.endswith(",00"):
        formatted = formatted[:-3]
    return f"{formatted}€"


def _format_date_for_grant_chat(value) -> str:
    d = _parse_sql_date_like(value)
    if not d:
        return "No disponible"
    return d.strftime("%d/%m/%Y")


def _grant_chat_tokenize(text: str) -> list[str]:
    if not text:
        return []
    tokens = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9]{3,}", str(text).lower())
    return [t for t in tokens if t not in {"para", "sobre", "esta", "subvencion", "subvención", "con", "del"}]


def _split_text_for_grant_chat(text: str, max_chars: int = 1200, overlap: int = 120) -> list[str]:
    raw = (text or "").strip()
    if not raw:
        return []

    parts = [p.strip() for p in re.split(r"\n{2,}", raw) if p.strip()]
    if not parts:
        parts = [raw]

    chunks = []
    current = ""
    for part in parts:
        candidate = f"{current}\n\n{part}".strip() if current else part
        if len(candidate) <= max_chars:
            current = candidate
            continue

        if current:
            chunks.append(current.strip())
            tail = current[-overlap:] if overlap > 0 else ""
            current = (tail + "\n" + part).strip() if tail else part
        else:
            start = 0
            while start < len(part):
                end = min(start + max_chars, len(part))
                chunk = part[start:end].strip()
                if chunk:
                    chunks.append(chunk)
                if end >= len(part):
                    break
                start = max(0, end - overlap)
            current = ""

    if current:
        chunks.append(current.strip())

    dedup = []
    seen = set()
    for chunk in chunks:
        key = chunk.casefold()
        if key in seen:
            continue
        seen.add(key)
        dedup.append(chunk)
    return dedup


def _extract_beneficiarios_short_text(value) -> str:
    if not value:
        return "No especificado"
    try:
        data = value if isinstance(value, dict) else json.loads(value) if isinstance(value, str) else None
        if isinstance(data, dict):
            categorias = data.get("categorias")
            if isinstance(categorias, list):
                clean = [str(v).strip() for v in categorias if str(v).strip()]
                if clean:
                    return ", ".join(clean)
            if isinstance(categorias, str) and categorias.strip():
                return categorias.strip()
    except Exception:
        pass
    return "No especificado"


def _get_selected_entity_id_for_user(cursor, user_id):
    try:
        cursor.execute("""
            SELECT ue.entity_id
            FROM user_entities ue
            WHERE ue.user_id = %s
              AND ue.is_selected IS TRUE
            LIMIT 1
        """, (user_id,))
        row = cursor.fetchone()
        if row:
            return row[0]
    except Exception as e:
        logger.warning(f"No se pudo obtener entidad seleccionada para chat user_id={user_id}: {e}")
    return None


def _fetch_grant_chat_record(cursor, grant_id, user_id):
    has_titulo = _grant_has_column(cursor, "titulo")
    has_admin = _grant_has_column(cursor, "administracion_convocante")
    has_tipo = _grant_has_column(cursor, "tipo_ayuda")

    titulo_expr = "g.titulo" if has_titulo else "NULL::text AS titulo"
    admin_expr = "g.administracion_convocante" if has_admin else "NULL::text AS administracion_convocante"
    tipo_expr = "g.tipo_ayuda" if has_tipo else "NULL::text AS tipo_ayuda"
    beneficiarios_expr = 'g."Beneficiarios_Short"' if _grant_has_column(cursor, "Beneficiarios_Short") else "NULL::jsonb"

    cursor.execute(f"""
        SELECT
            g.id,
            g.titulo_corto,
            {titulo_expr},
            g.presupuesto,
            g.fecha_finalizacion,
            g.fecha_de_publicacion,
            g.resumen_completo,
            g.region_impacto,
            g.finalidad,
            {admin_expr},
            {tipo_expr},
            {beneficiarios_expr}
        FROM grants g
        WHERE g.id = %s
          AND g.resumen_completo IS NOT NULL
        LIMIT 1
    """, (grant_id,))
    row = cursor.fetchone()
    if not row:
        return None

    entity_id = _get_selected_entity_id_for_user(cursor, user_id)

    match_row = None
    if entity_id is not None:
        try:
            cursor.execute("""
                SELECT numero_match, justificacion, recomendacion
                FROM matches
                WHERE grant_id = %s
                  AND entity_id = %s
                ORDER BY numero_match DESC
                LIMIT 1
            """, (grant_id, entity_id))
            match_row = cursor.fetchone()
        except Exception as e:
            logger.warning(f"No se pudo leer match para grant {grant_id}, entity {entity_id}: {e}")

    extra_metadata = get_grant_extra_metadata(cursor, grant_id)
    documentacion_items = parse_documentacion_items(extra_metadata.get("documentacion"))
    deadline_ui = _format_deadline_for_ui(row[4], row[5]) or "No disponible"

    return {
        "grant_id": row[0],
        "titulo_corto": row[1] or "",
        "titulo_oficial": row[2] or "",
        "presupuesto": row[3],
        "fecha_limite": deadline_ui,
        "fecha_publicacion": _format_date_for_grant_chat(row[5]),
        "resumen_html": normalize_summary_html(row[6]) if row[6] else "",
        "resumen_texto": html_to_text_for_export(normalize_summary_html(row[6])) if row[6] else "",
        "region_impacto": row[7] or "No especificado",
        "finalidad": row[8] or "No especificado",
        "administracion_convocante": row[9] or "No especificado",
        "tipo_ayuda": row[10] or "No especificado",
        "beneficiarios": _extract_beneficiarios_short_text(row[11]),
        "numero_match": (int(round(float(match_row[0]) * 100)) if match_row and match_row[0] is not None and float(match_row[0]) <= 1
                         else int(round(float(match_row[0]))) if match_row and match_row[0] is not None
                         else None),
        "justificacion_texto": html_to_text_for_export(normalize_summary_html(match_row[1])) if match_row and match_row[1] else "",
        "recomendacion_texto": html_to_text_for_export(normalize_summary_html(match_row[2])) if match_row and match_row[2] else "",
        "documentacion_items": documentacion_items,
        "codigobdns": extra_metadata.get("codigobdns"),
    }


def _grant_chat_chunks_table_available(cursor) -> bool:
    required = ("grant_id", "content")
    return all(_table_has_column(cursor, "grant_chat_chunks", c) for c in required)


def _fetch_grant_chat_chunks_from_table(cursor, grant_id: int) -> list[dict]:
    if not _grant_chat_chunks_table_available(cursor):
        return []

    has_type = _table_has_column(cursor, "grant_chat_chunks", "chunk_type")
    has_order = _table_has_column(cursor, "grant_chat_chunks", "chunk_order")
    has_updated = _table_has_column(cursor, "grant_chat_chunks", "updated_at")

    type_expr = "chunk_type" if has_type else "'context'::text AS chunk_type"
    order_expr = "chunk_order" if has_order else "0 AS chunk_order"
    updated_expr = "updated_at" if has_updated else "NULL::timestamp AS updated_at"
    order_by = []
    if has_order:
        order_by.append("chunk_order ASC")
    if has_updated:
        order_by.append("updated_at DESC NULLS LAST")
    order_by.append("id ASC")

    try:
        cursor.execute(f"""
            SELECT id, {type_expr}, content, {order_expr}, {updated_expr}
            FROM grant_chat_chunks
            WHERE grant_id = %s
              AND content IS NOT NULL
              AND TRIM(content) <> ''
            ORDER BY {", ".join(order_by)}
            LIMIT 500
        """, (grant_id,))
        rows = cursor.fetchall()
    except Exception as e:
        logger.warning(f"No se pudieron leer chunks de grant_chat_chunks para grant {grant_id}: {e}")
        return []

    return [
        {
            "id": row[0],
            "chunk_type": row[1] or "context",
            "content": str(row[2]).strip(),
            "chunk_order": row[3] if row[3] is not None else 0,
        }
        for row in rows
        if row and row[2] and str(row[2]).strip()
    ]


def _build_grant_chat_fallback_chunks(record: dict) -> list[dict]:
    chunks: list[dict] = []

    metadata_lines = [
        f"Título corto: {record.get('titulo_corto') or 'No disponible'}",
        f"Título oficial: {record.get('titulo_oficial') or 'No disponible'}",
        f"Código BDNS: {record.get('codigobdns') or 'No disponible'}",
        f"Beneficiarios: {record.get('beneficiarios') or 'No disponible'}",
        f"Fondos totales: {_format_amount_for_grant_chat(record.get('presupuesto'))}",
        f"Plazo: {record.get('fecha_limite') or 'No disponible'}",
        f"Fecha de publicación: {record.get('fecha_publicacion') or 'No disponible'}",
        f"Región de impacto: {record.get('region_impacto') or 'No disponible'}",
        f"Finalidad: {record.get('finalidad') or 'No disponible'}",
        f"Administración convocante: {record.get('administracion_convocante') or 'No disponible'}",
        f"Tipo de ayuda: {record.get('tipo_ayuda') or 'No disponible'}",
    ]
    if record.get("numero_match") is not None:
        metadata_lines.append(f"Compatibilidad con la entidad: {record['numero_match']}%")

    chunks.append({
        "chunk_type": "metadata",
        "content": "\n".join(metadata_lines),
        "chunk_order": 0,
    })

    for idx, piece in enumerate(_split_text_for_grant_chat(record.get("resumen_texto", "")), start=1):
        chunks.append({"chunk_type": "resumen", "content": piece, "chunk_order": idx})

    for idx, piece in enumerate(_split_text_for_grant_chat(record.get("justificacion_texto", "")), start=1):
        chunks.append({"chunk_type": "justificacion", "content": piece, "chunk_order": idx})

    for idx, piece in enumerate(_split_text_for_grant_chat(record.get("recomendacion_texto", "")), start=1):
        chunks.append({"chunk_type": "recomendacion", "content": piece, "chunk_order": idx})

    doc_items = record.get("documentacion_items") or []
    if doc_items:
        doc_text = "\n".join([f"- {item}" for item in doc_items])
        for idx, piece in enumerate(_split_text_for_grant_chat(doc_text), start=1):
            chunks.append({"chunk_type": "documentacion", "content": piece, "chunk_order": idx})

    return chunks


def _rank_grant_chat_chunks(question: str, chunks: list[dict], top_k: int = 6) -> list[dict]:
    if not chunks:
        return []
    q = (question or "").strip()
    q_lower = q.lower()
    q_tokens = _grant_chat_tokenize(q)

    scored = []
    for idx, chunk in enumerate(chunks):
        content = str(chunk.get("content") or "").strip()
        if not content:
            continue
        text_lower = content.lower()
        score = 0.0

        if chunk.get("chunk_type") == "metadata":
            score += 1.2
        elif chunk.get("chunk_type") == "documentacion":
            score += 0.6

        if q_lower and q_lower in text_lower:
            score += 8.0

        if q_tokens:
            token_hits = 0
            for token in q_tokens:
                if token in text_lower:
                    token_hits += 1
            score += token_hits * 1.5
            if token_hits:
                score += min(2.5, token_hits / max(1, len(q_tokens)) * 3)

        scored.append((score, idx, chunk))

    scored.sort(key=lambda item: (-item[0], item[1]))

    result = []
    seen = set()
    # Siempre intenta incluir metadata primero para respuestas completas.
    metadata_chunk = next((c for c in chunks if c.get("chunk_type") == "metadata"), None)
    if metadata_chunk:
        result.append(metadata_chunk)
        seen.add(id(metadata_chunk))

    for score, _, chunk in scored:
        if len(result) >= top_k:
            break
        key = id(chunk)
        if key in seen:
            continue
        # Si no hay matches semánticos, igualmente coger contexto base.
        if score <= 0 and len(result) >= 2:
            continue
        result.append(chunk)
        seen.add(key)

    if not result:
        return chunks[:top_k]
    return result[:top_k]


def _sanitize_grant_chat_history(message_history) -> list[dict]:
    if not isinstance(message_history, list):
        return []
    sanitized = []
    for item in message_history[-8:]:
        if not isinstance(item, dict):
            continue
        role = item.get("role")
        if role not in ("user", "assistant"):
            continue
        raw = str(item.get("content") or "").strip()
        if not raw:
            continue
        content = html_to_text_for_export(raw) if ("<" in raw and ">" in raw) else raw
        content = re.sub(r"\s+", " ", content).strip()
        if not content:
            continue
        sanitized.append({"role": role, "content": content[:1600]})
    return sanitized


def _build_grant_chat_response(record: dict, question: str, chunks: list[dict], message_history=None) -> str:
    if not question or not str(question).strip():
        return "<p>Escribe una pregunta para que pueda ayudarte con esta subvención.</p>"

    context_blocks = []
    for idx, chunk in enumerate(chunks, start=1):
        chunk_type = str(chunk.get("chunk_type") or "contexto").strip()
        content = str(chunk.get("content") or "").strip()
        if not content:
            continue
        context_blocks.append(f"[{idx}] ({chunk_type})\n{content}")

    if not context_blocks:
        context_blocks = ["[1] (metadata)\nNo hay contexto documental disponible para esta subvención."]

    title = record.get("titulo_corto") or record.get("titulo_oficial") or f"Subvención {record.get('grant_id')}"
    system_prompt = (
        "Eres Granti, un asistente experto en subvenciones públicas. "
        "Responde SOLO con la información del contexto proporcionado para esta subvención concreta. "
        "Si una información no aparece en el contexto, dilo explícitamente. "
        "No inventes requisitos, plazos ni importes. "
        "Responde en español claro y usa HTML básico (<p>, <strong>, <ul>, <li>, <h3>) cuando ayude a la lectura."
    )

    user_prompt = (
        f"Subvención: {title}\n"
        f"Pregunta del usuario: {question.strip()}\n\n"
        "Contexto recuperado:\n"
        + "\n\n".join(context_blocks)
        + "\n\nInstrucciones:\n"
          "- Prioriza una respuesta breve y útil.\n"
          "- Si te preguntan por requisitos/documentación, enumera en lista.\n"
          "- Si hay duda o falta un dato, indícalo claramente.\n"
    )

    messages = [{"role": "system", "content": system_prompt}]
    messages.extend(_sanitize_grant_chat_history(message_history))
    messages.append({"role": "user", "content": user_prompt})

    try:
        client = _get_openai_client_for_grant_chat()
        completion = client.chat.completions.create(
            model=os.getenv("GRANT_CHAT_MODEL", "gpt-4o-mini"),
            messages=messages,
            temperature=0.2,
            timeout=30,
        )
        content = completion.choices[0].message.content if completion and completion.choices else None
        if content and str(content).strip():
            return str(content).strip()
    except Exception as e:
        logger.error(f"Error generando respuesta de chat de subvención para grant {record.get('grant_id')}: {e}", exc_info=True)

    # Fallback determinista si OpenAI falla.
    fallback_lines = [
        f"<p><strong>{pyhtml.escape(title)}</strong></p>",
        "<p>No he podido generar una respuesta completa en este momento. Te dejo la información disponible de la subvención:</p>",
        "<ul>",
        f"<li><strong>Beneficiarios:</strong> {pyhtml.escape(str(record.get('beneficiarios') or 'No disponible'))}</li>",
        f"<li><strong>Fondos totales:</strong> {pyhtml.escape(_format_amount_for_grant_chat(record.get('presupuesto')))}</li>",
        f"<li><strong>Plazo:</strong> {pyhtml.escape(str(record.get('fecha_limite') or 'No disponible'))}</li>",
        f"<li><strong>Región:</strong> {pyhtml.escape(str(record.get('region_impacto') or 'No disponible'))}</li>",
        f"<li><strong>Finalidad:</strong> {pyhtml.escape(str(record.get('finalidad') or 'No disponible'))}</li>",
        "</ul>",
    ]
    if record.get("documentacion_items"):
        fallback_lines.append("<p><strong>Documentación a aportar:</strong></p><ul>")
        for item in record["documentacion_items"][:12]:
            fallback_lines.append(f"<li>{pyhtml.escape(str(item))}</li>")
        fallback_lines.append("</ul>")
    return "".join(fallback_lines)


@app.route('/api/grants/<int:grant_id>/chat', methods=['POST'])
@require_auth
def chat_about_grant(user_id, grant_id):
    """Chat específico y rápido por subvención (SQL-first, sin Milvus)."""
    connection = None
    cursor = None
    try:
        data = request.get_json() or {}
        user_message = str(data.get("message") or "").strip()
        message_history = data.get("messageHistory") or []
        user_id_for_limits = str(user_id or data.get("user_id") or "anonymous")

        if not user_message:
            return jsonify({"error": "missing_message", "message": "Falta la pregunta del usuario"}), 400

        limit_check = check_user_limit(user_id_for_limits)
        if not limit_check['allowed']:
            reset_time_str = limit_check["reset_time"]
            return jsonify({
                'error': 'limit_exceeded',
                'message': f'Has alcanzado el límite diario de {8} preguntas. Se reseteará mañana a las {reset_time_str}.',
                'remaining': 0,
                'reset_time': limit_check['reset_time']
            }), 429

        connection = get_connection()
        cursor = connection.cursor()

        record = _fetch_grant_chat_record(cursor, grant_id, user_id)
        if not record:
            return jsonify({"error": "grant_not_found", "message": "Subvención no encontrada"}), 404

        stored_chunks = _fetch_grant_chat_chunks_from_table(cursor, grant_id)
        fallback_chunks = _build_grant_chat_fallback_chunks(record)
        chunks_pool = stored_chunks if stored_chunks else fallback_chunks
        selected_chunks = _rank_grant_chat_chunks(user_message, chunks_pool, top_k=6)

        response_html = _build_grant_chat_response(
            record=record,
            question=user_message,
            chunks=selected_chunks,
            message_history=message_history,
        )

        increment_usage(user_id_for_limits)
        updated_limit = check_user_limit(user_id_for_limits)

        return jsonify({
            "response": response_html,
            "remaining": updated_limit.get("remaining"),
            "reset_time": updated_limit.get("reset_time"),
            "grant_id": grant_id,
            "retrieval_source": "grant_chat_chunks" if stored_chunks else "grants_sql_fallback",
            "chunks_used": len(selected_chunks),
        })

    except Exception as e:
        logger.error(f"Error en chat de subvención grant_id={grant_id}: {e}", exc_info=True)
        return jsonify({"error": "chat_error", "message": str(e)}), 500
    finally:
        try:
            if cursor:
                cursor.close()
            if connection:
                connection.close()
        except Exception:
            pass

# Endpoint para la API de asistente
@app.route('/api/app_assistente', methods=['POST'])
def app_assistente():
    # Get the user's message from the JSON payload
    data = request.get_json()
    user_message = data.get('message')
    user_id = data.get('user_id', 'anonymous')  # Obtener user_id del request
    selected_grant = data.get('selectedGrant')  # Obtener subvención seleccionada
    # Obtener historial de mensajes si está disponible
    message_history = data.get('messageHistory', [])

    print(f"Message from user {user_id}: {user_message}")
    
    # Verificar límite de mensajes del usuario
    limit_check = check_user_limit(user_id)
    
    if not limit_check['allowed']:
        return jsonify({
            'error': 'limit_exceeded',
            'message': f'Has alcanzado el límite diario de {8} preguntas. Se reseteará mañana a las {limit_check["reset_time"]}.',
            'remaining': 0,
            'reset_time': limit_check['reset_time']
        }), 429

    # Use the LLM to determine if a function call is needed
    llm = LLM()
    tool_response, function_name, tool_id, args = llm.process_functions(user_message, message_history, selected_grant)
    
    # Initialize function_response as empty string to avoid UnboundLocalError
    function_response = ""
    
    if function_name is not None:
        try:
            # Si no faltan argumentos, proceder con la ejecución de la función
            if function_name == "get_preguntas_subvenciones":
                print("=" * 80)
                print(f"🔍 CHAT ASSISTANT - Procesando pregunta sobre subvención")
                print(f"   Título recibido: '{args['titulo']}'")
                print(f"   Mensaje del usuario: '{user_message}'")
                logger.info("=" * 80)
                logger.info(f"🔍 CHAT ASSISTANT - Procesando pregunta sobre subvención")
                logger.info(f"   Título recibido: '{args['titulo']}'")
                logger.info(f"   Mensaje del usuario: '{user_message}'")

                handler = get_preguntas(args["titulo"], user_message)
                function_response = handler.main()

                print(f"   Respuesta obtenida: {function_response is not None}")
                logger.info(f"   Tipo de respuesta: {type(function_response)}")
                if function_response:
                    logger.info(f"   Longitud de respuesta: {len(function_response) if isinstance(function_response, (list, str)) else 'N/A'}")
                logger.info("=" * 80)

                logger.info(f"   Respuesta obtenida: {function_response is not None}")
                if function_response is None:
                    logger.error(f"❌ CHAT ASSISTANT - No se obtuvo respuesta de Milvus para '{args['titulo']}'")
            elif function_name == "get_subvenciones_hoy":
                function_response = subvenciones_hoy.listar_subvenciones()
            elif function_name == "get_subvenciones_finalidad":
                function_response = sub_finalidad.finalidad(args["finalidad"])
            elif function_name == "get_subvenciones_region":
                function_response = sub_region.region(args["region"])
            elif function_name == "get_subvenciones_beneficiario":
                function_response = sub_beneficiario.beneficiario(args["tipo_juridico"])
            elif function_name == "get_resumen_subvenciones":
                function_response = GetResumen.resumen(args["titulo"])
            elif function_name == "get_match_entidad_subvencion":
                function_response = get_match.get_best_matches()
            else:
                function_response = "Función no implementada"
            
        except Exception as e:
            error_message = f"Error al procesar la función: {str(e)}"
            logger.error(f"❌ ERROR en función {function_name}: {error_message}", exc_info=True)
            print(error_message)
            function_response = error_message
    
    # Process the response regardless of whether there was a function call or not
    final_response = llm.process_response(tool_response, tool_id, function_response)
    
    # Incrementar el contador de uso después de procesar exitosamente
    increment_usage(user_id)
    
    # Obtener el nuevo estado del límite
    updated_limit = check_user_limit(user_id)
    
    return jsonify({
        'response': final_response,
        'remaining': updated_limit['remaining'],
        'reset_time': updated_limit['reset_time']
    })

# Endpoint para consultar los límites de chat sin enviar mensaje
@app.route('/api/get_chat_limits', methods=['GET'])
def get_chat_limits():
    """Endpoint para obtener los límites de chat actuales del usuario sin enviar mensaje"""
    try:
        user_id = request.args.get('user_id')

        if not user_id:
            return jsonify({
                'error': 'Se requiere user_id'
            }), 400

        # Verificar límite de mensajes del usuario
        limit_check = check_user_limit(user_id)

        return jsonify({
            'remaining': limit_check['remaining'],
            'reset_time': limit_check['reset_time'],
            'message_count': limit_check['message_count'],
            'allowed': limit_check['allowed']
        }), 200

    except Exception as e:
        logger.error(f"Error al obtener límites de chat: {e}")
        # En caso de error, devolver valores por defecto
        return jsonify({
            'remaining': 8,
            'reset_time': '00:00',
            'message_count': 0,
            'allowed': True
        }), 200

# Endpoint para obtener todas las subvenciones para el selector en el asistente
@app.route('/api/get_all_grants', methods=['GET'])
def get_all_grants():
    """Endpoint para obtener todas las subvenciones de la base de datos"""
    try:
        connection = get_connection()
        cursor = connection.cursor()
        print("Connected to the database successfully.")

        query = """
            SELECT id, titulo_corto
            FROM grants
            WHERE resumen_completo IS NOT NULL
            ORDER BY id DESC
        """
        cursor.execute(query)
        rows = cursor.fetchall()
        
        grants = []
        for row in rows:
            grants.append({
                "id": row[0],
                "title": row[1]
            })
        
        return jsonify(grants=grants)
    
    except Exception as e:
        print("An error occurred:", e)
        return jsonify(error=str(e)), 500
    finally:
        if connection:
            cursor.close()
            connection.close()
            print("Database connection closed.")

# Endpoints para el sistema de swipe de subvenciones
@app.route('/api/get_grants_for_swipe', methods=['GET'])
@require_auth
def get_grants_for_swipe(user_id):
    """Endpoint para obtener subvenciones para el sistema de swipe, excluyendo las ya vistas por el usuario."""
    logger.info(f"Getting grants for swipe for user_id: {user_id}")
    if not user_id:
        return jsonify({"error": "Falta parámetro user_id"}), 400

    try:
        conn = get_connection()
        cursor = conn.cursor()
        selected_entity_id = _get_selected_entity_id_for_user(cursor, user_id)

        # Selecciona las subvenciones que el usuario NO haya marcado todavía
        if selected_entity_id:
            cursor.execute("""
                SELECT
                    g.id,
                    g.titulo_corto,
                    g.presupuesto,
                    g.fecha_finalizacion,
                    g.fecha_de_publicacion,
                    g.resumen_completo,
                    g."Beneficiarios_Short",
                    g.region_impacto,
                    g.finalidad
                FROM public.grants g
                WHERE
                    g.resumen_completo IS NOT NULL
                    AND g.fecha_recepcion IS NOT NULL
                    AND (g.fecha_finalizacion IS NULL OR g.fecha_finalizacion::date >= CURRENT_DATE)
                    AND NOT EXISTS (
                        SELECT 1
                        FROM public.user_grant_preferences ugp
                        WHERE ugp.user_id = %s
                        AND ugp.entity_id = %s
                        AND CAST(ugp.grant_id AS INTEGER) = g.id)
                ORDER BY RANDOM()
                LIMIT 50;
            """, (user_id, selected_entity_id))
        else:
            cursor.execute("""
                SELECT
                    g.id,
                    g.titulo_corto,
                    g.presupuesto,
                    g.fecha_finalizacion,
                    g.fecha_de_publicacion,
                    g.resumen_completo,
                    g."Beneficiarios_Short",
                    g.region_impacto,
                    g.finalidad
                FROM public.grants g
                WHERE
                    g.resumen_completo IS NOT NULL
                    AND g.fecha_recepcion IS NOT NULL
                    AND (g.fecha_finalizacion IS NULL OR g.fecha_finalizacion::date >= CURRENT_DATE)
                    AND NOT EXISTS (
                        SELECT 1
                        FROM public.user_grant_preferences ugp
                        WHERE ugp.user_id = %s
                        AND CAST(ugp.grant_id AS INTEGER) = g.id)
                ORDER BY RANDOM()
                LIMIT 50;
            """, (user_id,))

        grants = cursor.fetchall()
        cursor.close()
        conn.close()

        # Formatear datos
        grants_list = []
        for grant in grants:
            # Extraer beneficiarios del JSONB
            beneficiarios_jsonb = grant[6]
            beneficiarios_value = ""

            try:
                if beneficiarios_jsonb:
                    if isinstance(beneficiarios_jsonb, dict):
                        categorias = beneficiarios_jsonb.get('categorias', [])
                        if categorias and isinstance(categorias, list):
                            beneficiarios_value = ', '.join(str(c) for c in categorias)
                        elif categorias and isinstance(categorias, str):
                            beneficiarios_value = categorias
                    elif isinstance(beneficiarios_jsonb, str):
                        try:
                            data = json.loads(beneficiarios_jsonb)
                            categorias = data.get('categorias', [])
                            if categorias and isinstance(categorias, list):
                                beneficiarios_value = ', '.join(str(c) for c in categorias)
                            elif categorias and isinstance(categorias, str):
                                beneficiarios_value = categorias
                        except:
                            pass
            except Exception as e:
                logger.error(f"Error processing beneficiarios for grant {grant[0]}: {e}")

            grants_list.append({
                "id":            grant[0],
                "title":         grant[1] or "Sin título",
                "amount":        grant[2] or "No especificado",
                "deadline":      _format_deadline_for_ui(grant[3], grant[4]),
                "description":   grant[5] or "Sin descripción",
                "beneficiario":  beneficiarios_value,
                "lugar":         grant[7] or "No especificado",
                "finalidad":     grant[8] or "No especificado"
            })
        logger.info(f"Returning {len(grants_list)} grants for swipe for user_id: {user_id}")
        return jsonify({"grants": grants_list})
    
    except Exception as e:
        print(f"Error en get_grants_for_swipe: {e}")
        return jsonify({"error": str(e)}), 500

# Endpoint para obtener los mejores matches para la entidad del usuario actual
@app.route('/api/get_best_matches', methods=['GET'])
@require_auth
def get_best_matches(user_id):
    """Endpoint para obtener los mejores matches para la entidad del usuario actual"""
    if not user_id:
        logger.info("Falta parámetro user_id en get_best_matches")
        return jsonify({"error": "Falta parámetro user_id"}), 400

    try:
        connection = get_connection()
        cursor = connection.cursor()
        logger.info("Connected to the database successfully.")

        # Obtener la entidad seleccionada del usuario actual
        query = """
            SELECT e.id
            FROM entities e
            JOIN user_entities ue ON e.id = ue.entity_id
            WHERE ue.user_id = %s
                AND ue.is_selected = TRUE
            ORDER BY ue.updated_at DESC
            LIMIT 1;
        """
        cursor.execute(query, (user_id,))
        row = cursor.fetchone()

        if not row:
            logger.warning("No entity found for the current user.")
            return jsonify(matches=[])

        entity_id = row[0]

        # 3) Matches de esa entidad, verificando que pertenece al usuario
        q_matches = """
            SELECT 
                g.id AS grant_id,
                g.titulo_corto AS title,
                g.presupuesto AS amount,
                g.importe_beneficiario,
                g.fecha_finalizacion AS deadline,
                g.fecha_de_publicacion,
                g.resumen_completo AS resumen,
                string_agg(DISTINCT b.value, ', ') AS beneficiario,
                g.region_impacto AS lugar,
                g.finalidad AS finalidad,
                g.administracion_convocante,
                g.tipo_ayuda,
                g.fecha_inicio_solicitud,
                g.fecha_finalizacion,
                m.justificacion,
                CASE
                    WHEN m.numero_match IS NULL THEN NULL
                    WHEN m.numero_match > 1 THEN m.numero_match / 100.0
                    ELSE m.numero_match
                END AS numero_match_ratio
            FROM matches m
            JOIN grants g ON m.grant_id = g.id
            LEFT JOIN LATERAL (
                SELECT value
                FROM jsonb_array_elements_text(
                    CASE 
                        WHEN jsonb_typeof(g."Beneficiarios_Short"->'categorias') = 'array' 
                        THEN g."Beneficiarios_Short"->'categorias'
                        ELSE '[]'::jsonb
                    END
                )
            ) AS b(value) ON TRUE
            JOIN user_entities ue ON ue.entity_id = m.entity_id
            WHERE 
                ue.user_id = %s
                AND m.entity_id = %s
                AND (
                    CASE
                        WHEN m.numero_match IS NULL THEN 0
                        WHEN m.numero_match > 1 THEN m.numero_match / 100.0
                        ELSE m.numero_match
                    END
                ) >= 0.8
                AND g.resumen_completo IS NOT NULL
            GROUP BY 
                g.id, g.titulo_corto, g.presupuesto, g.importe_beneficiario, g.fecha_finalizacion, g.fecha_de_publicacion,
                g.resumen_completo, g.region_impacto, g.finalidad, g.administracion_convocante,
                g.tipo_ayuda, g.fecha_inicio_solicitud, g.fecha_finalizacion, m.justificacion, m.numero_match
            ORDER BY
                CASE
                    WHEN m.numero_match IS NULL THEN 0
                    WHEN m.numero_match > 1 THEN m.numero_match / 100.0
                    ELSE m.numero_match
                END DESC,
                g.fecha_de_publicacion DESC;
        """

        cursor.execute(q_matches, (user_id, entity_id))
        matches_rows = cursor.fetchall()

        matches = []
        for row in matches_rows: 
            matches.append({
                "grant_id": row[0], 
                "title": row[1], 
                "amount": row[2] if row[2] else "No especificado", 
                "importe_beneficiario": row[3] if row[3] else "No especificado",
                "deadline": _format_deadline_for_ui(row[4], row[5]),
                "resumen": row[6] if row[6] else "Sin descripción disponible", 
                "beneficiario": row[7] if row[7] else "No especificado", 
                "lugar": row[8] if row[8] else "No especificado", 
                "finalidad": row[9] if row[9] else "No especificado",
                "administracion_convocante": row[10] if row[10] else "No especificado",
                "tipo_ayuda": row[11] if row[11] else "No especificado",
                "fecha_inicio_solicitud": row[12].isoformat() if isinstance(row[12], (datetime, date)) else (str(row[12]) if row[12] else None),
                "fecha_de_cierre": row[13].isoformat() if isinstance(row[13], (datetime, date)) else (str(row[13]) if row[13] else None),
                "justificacion": row[14] if row[14] else "Sin justificacióån disponible", 
                "numero_match": _normalize_match_ratio(row[15]) }) 
        
        logger.info(f"Returning {len(matches)} matches for user {user_id} and entity {entity_id}, matches details: {matches}")
            
        return jsonify(matches=matches)

    except Exception as e:
        # Usa logging.error en producción
        print("An error occurred:", e)
        return jsonify({"error": "No tienes ningún Match por el momento."}), 500
    finally:
        try:
            cursor.close()
            connection.close()
        except:
            pass

# Endpoint para obtener los detalles de un match específico por ID de subvención
@app.route('/api/get_match_detail/<string:id>', methods=['GET'])
@require_auth
def get_match_detail(user_id, id):
    """Endpoint para obtener los detalles de un match específico por ID de subvención"""
    try:
        connection = get_connection()
        cursor = connection.cursor()
        logger.info("Connected to the database successfully.")

        entity_id = _get_selected_entity_id_for_user(cursor, user_id)
        if not entity_id:
            logger.info("No selected entity for user_id=%s, trying latest linked entity.", user_id)
            cursor.execute("""
                SELECT ue.entity_id
                FROM user_entities ue
                WHERE ue.user_id = %s
                ORDER BY ue.is_selected DESC, ue.updated_at DESC NULLS LAST, ue.created_at DESC
                LIMIT 1
            """, (user_id,))
            entity_result = cursor.fetchone()
            entity_id = entity_result[0] if entity_result else None
            if not entity_id:
                logger.warning("No entity found for user_id=%s.", user_id)
                return jsonify(error="No se encontró una entidad asociada al usuario actual"), 404

        # --- Consulta principal ---
        query = """
            SELECT 
                g.id,
                g.titulo_corto,
                g.presupuesto,
                g.fecha_finalizacion,
                g.fecha_de_publicacion,
                g.resumen_completo,
                string_agg(DISTINCT b.value, ', ') AS beneficiarios,
                g.region_impacto,
                m.numero_match, 
                m.justificacion, 
                m.recomendacion
            FROM grants g
            LEFT JOIN LATERAL jsonb_array_elements_text(
                    g."Beneficiarios_Short"->'categorias'
                ) AS b(value) ON TRUE
            LEFT JOIN LATERAL (
                SELECT 
                    m.numero_match, 
                    m.justificacion, 
                    m.recomendacion
                FROM matches m
                WHERE m.grant_id = g.id
                AND m.entity_id = %s
                ORDER BY
                    CASE
                        WHEN m.numero_match IS NULL THEN 0
                        WHEN m.numero_match > 1 THEN m.numero_match / 100.0
                        ELSE m.numero_match
                    END DESC
                LIMIT 1
            ) m ON TRUE
            WHERE 
                g.id = %s
                AND g.resumen_completo IS NOT NULL
            GROUP BY 
                g.id, g.titulo_corto, g.presupuesto, g.fecha_finalizacion, g.fecha_de_publicacion,
                g.resumen_completo, g.region_impacto,
                m.numero_match, m.justificacion, m.recomendacion
            LIMIT 1;
        """

        cursor.execute(query, (entity_id, id))
        match_row = cursor.fetchone()
        
        if not match_row:
            logger.warning(f"Match with ID {id} not found or incomplete.")
            return jsonify(error="Match no encontrado o sin datos suficientes"), 404

        # --- Procesar campos ---
        deadline = _format_deadline_for_ui(match_row[3], match_row[4])

        amount = match_row[2]
        if isinstance(amount, (int, float)):
            amount = f"{amount}€"
        elif isinstance(amount, str) and not amount.endswith("€") and amount != "No especificado":
            amount = f"{amount}€"

        extra_metadata = get_grant_extra_metadata(cursor, match_row[0])

        match_detail = {
            "id": match_row[0],
            "grant_id": match_row[0],  # Add grant_id for consistency with frontend expectations
            "titulo_corto": match_row[1],
            "presupuesto": amount or "No especificado",
            "fecha_limite": deadline,
            "resumen_completo": normalize_summary_html(match_row[5]) if match_row[5] else "Sin descripción disponible",
            "beneficiarios": match_row[6] or "No especificado",
            "region_impacto": match_row[7] or "No especificado",
            "numero_match": _normalize_match_ratio(match_row[8]),
            "justificacion": normalize_summary_html(match_row[9]) if match_row[9] else "Sin justificación disponible",
            "recomendacion": normalize_summary_html(match_row[10]) if match_row[10] else "Sin recomendación disponible",
            "documentacion": extra_metadata.get("documentacion"),
            "codigobdns": extra_metadata.get("codigobdns")
        }

        logger.info(f"Returning match detail: {match_detail}")
        return jsonify(match=match_detail)

    except Exception as e:
        logger.error(f"Error al obtener detalles del match: {e}")
        return jsonify(error=str(e)), 500

    finally:
        if cursor:
            cursor.close()
        if connection:
            connection.close()
            logger.info("Database connection closed.")

# Agregar este endpoint en app.py
@app.route('/api/get_grant_detail/<string:id>', methods=['GET'])
@require_auth
def get_grant_detail(user_id, id):
    """Endpoint para obtener los detalles de una subvención por ID"""
    try:
        connection = get_connection()
        cursor = connection.cursor()
        logger.info("Connected to the database successfully.")

        entity_id = _get_selected_entity_id_for_user(cursor, user_id)
        if not entity_id:
            cursor.execute("""
                SELECT ue.entity_id
                FROM user_entities ue
                WHERE ue.user_id = %s
                ORDER BY ue.is_selected DESC, ue.updated_at DESC NULLS LAST, ue.created_at DESC
                LIMIT 1
            """, (user_id,))
            entity_result = cursor.fetchone()
            entity_id = entity_result[0] if entity_result else None
            if not entity_id:
                logger.warning("No entity found for user_id=%s.", user_id)
                return jsonify(error="No se encontró una entidad asociada al usuario actual"), 404

        # --- Consulta principal ---
        query = """
            SELECT 
                g.id,
                g.titulo_corto,
                g.presupuesto,
                g.fecha_finalizacion,
                g.fecha_de_publicacion,
                g.resumen_completo,
                string_agg(DISTINCT b.value, ', ') AS beneficiarios,
                g.region_impacto,
                m.numero_match, 
                m.justificacion, 
                m.recomendacion
            FROM grants g
            LEFT JOIN LATERAL jsonb_array_elements_text(
                    g."Beneficiarios_Short"->'categorias'
                ) AS b(value) ON TRUE
            LEFT JOIN LATERAL (
                SELECT 
                    m.numero_match, 
                    m.justificacion, 
                    m.recomendacion
                FROM matches m
                WHERE m.grant_id = g.id
                AND m.entity_id = %s
                ORDER BY
                    CASE
                        WHEN m.numero_match IS NULL THEN 0
                        WHEN m.numero_match > 1 THEN m.numero_match / 100.0
                        ELSE m.numero_match
                    END DESC
                LIMIT 1
            ) m ON TRUE
            WHERE 
                g.id = %s
                AND g.resumen_completo IS NOT NULL
            GROUP BY 
                g.id, g.titulo_corto, g.presupuesto, g.fecha_finalizacion, g.fecha_de_publicacion,
                g.resumen_completo, g.region_impacto,
                m.numero_match, m.justificacion, m.recomendacion
            LIMIT 1;
        """

        cursor.execute(query, (entity_id, id))
        grant_row = cursor.fetchone()
        logger.info(grant_row)
        
        if not grant_row:
            logger.warning(f"Grant with ID {id} not found or incomplete.")
            return jsonify(error="Subvención no encontrada o sin datos suficientes"), 404

        # --- Formatear campos ---
        deadline = _format_deadline_for_ui(grant_row[3], grant_row[4])

        amount = grant_row[2]
        if isinstance(amount, (int, float)):
            amount = f"{amount}€"
        elif isinstance(amount, str) and not amount.endswith("€") and amount != "No especificado":
            amount = f"{amount}€"

        extra_metadata = get_grant_extra_metadata(cursor, grant_row[0])

        grant_detail = {
            "id": grant_row[0],
            "titulo_corto": grant_row[1],
            "presupuesto": amount or "No especificado",
            "fecha_limite": deadline,
            "resumen_completo": normalize_summary_html(grant_row[5]) if grant_row[5] else "Sin descripción disponible",
            "beneficiarios": grant_row[6] or "No especificado",
            "region_impacto": grant_row[7] or "No especificado",
            "numero_match": _normalize_match_ratio(grant_row[8]) or 0,
            "justificacion": normalize_summary_html(grant_row[9]) if grant_row[9] else "Sin justificación disponible",
            "recomendacion": normalize_summary_html(grant_row[10]) if grant_row[10] else "Sin recomendación disponible",
            "documentacion": extra_metadata.get("documentacion"),
            "codigobdns": extra_metadata.get("codigobdns")
        }


        logger.info(f"Returning grant detail: {grant_detail}")
        return jsonify(grant=grant_detail)

    except Exception as e:
        logger.error(f"Error al obtener detalles de la subvención: {e}")
        return jsonify(error=str(e)), 500

    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()
            logger.info("Database connection closed.")

@app.route('/api/store_user_preference', methods=['POST'])
@require_auth
def store_preference(user_id):
    """
    Store user preference for a grant.
    user_id comes from @require_auth decorator (authenticated user from JWT token)
    """
    try:
        data = request.json
        grant_id = data.get('grant_id')
        action = data.get('action')

        # Validate grant_id (must be present and not empty)
        if not grant_id or grant_id == '':
            logger.error(f"Missing or empty grant_id in request: {data}")
            return jsonify({"error": "Se requiere grant_id válido"}), 400

        # Validate action
        if not action or action not in ['interesa', 'no interesa']:
            logger.error(f"Invalid action in request: {action}")
            return jsonify({"error": "Action debe ser 'interesa' o 'no interesa'"}), 400

        # user_id already comes from @require_auth, don't overwrite it
        # This ensures we use the authenticated user's ID, not what the client sends
        if not user_id:
            logger.error("No user_id from authentication")
            return jsonify({"error": "Usuario no autenticado"}), 401

        logger.info(f"Storing preference: user_id={user_id}, grant_id={grant_id}, action={action}")
        success = store_user_preference(grant_id, action, user_id)

        if success:
            return jsonify({"status": "success", "message": "Preferencia guardada correctamente"}), 200
        else:
            return jsonify({"error": "Error al almacenar preferencia en la base de datos"}), 500
    except Exception as e:
        logger.error(f"Error in store_preference endpoint: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route('/api/reco/events', methods=['POST'])
@require_auth
def log_reco_event_endpoint(user_id):
    """Register recommendation interaction events (append-only)."""
    try:
        ensure_recommendation_tables()
        data = request.get_json(silent=True) or {}

        event_type = (data.get('event_type') or '').strip()
        grant_id = data.get('grant_id')
        surface = data.get('surface')
        position = data.get('position')
        session_id = data.get('session_id')
        metadata = data.get('metadata') if isinstance(data.get('metadata'), dict) else {}
        event_value = data.get('event_value')

        if not event_type:
            return jsonify({"error": "Se requiere event_type"}), 400

        if grant_id in (None, ""):
            return jsonify({"error": "Se requiere grant_id"}), 400

        try:
            grant_id = int(str(grant_id).strip())
        except Exception:
            return jsonify({"error": "grant_id debe ser entero"}), 400

        if position is not None:
            try:
                position = int(position)
            except Exception:
                position = None

        ok = record_reco_event(
            user_id=user_id,
            grant_id=grant_id,
            event_type=event_type,
            surface=surface,
            position=position,
            session_id=session_id,
            metadata=metadata,
            event_value=event_value,
        )
        if not ok:
            return jsonify({"error": "No se pudo registrar el evento"}), 500

        return jsonify({"status": "ok"}), 200
    except Exception as e:
        logger.error(f"Error in log_reco_event endpoint: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route('/api/frontend_log', methods=['POST'])
@require_auth
def frontend_log_endpoint(user_id):
    """Persist frontend diagnostic logs in backend logger output."""
    try:
        payload = request.get_json(silent=True) or {}

        level = str(payload.get('level') or 'info').strip().lower()
        if level not in {'info', 'warning', 'error'}:
            level = 'info'

        context = str(payload.get('context') or 'frontend').strip() or 'frontend'
        message = str(payload.get('message') or '').strip()
        if not message:
            return jsonify({"error": "Se requiere message"}), 400

        details = payload.get('details')
        if details is None:
            details = {}
        if not isinstance(details, dict):
            details = {"raw_details": str(details)}

        log_payload = {
            "user_id": user_id,
            "context": context,
            "message": message,
            "details": details,
        }
        serialized = json.dumps(log_payload, ensure_ascii=False, default=str)

        if level == 'error':
            logger.error("FRONTEND_LOG %s", serialized)
        elif level == 'warning':
            logger.warning("FRONTEND_LOG %s", serialized)
        else:
            logger.info("FRONTEND_LOG %s", serialized)

        return jsonify({"status": "ok"}), 200
    except Exception as e:
        logger.error(f"Error in frontend_log endpoint: {e}", exc_info=True)
        return jsonify({"error": "No se pudo registrar log frontend"}), 500


@app.route('/api/get_user_preferences', methods=['POST'])
@require_auth
def get_preferences(user_id):
    try:
        from utils.recommendation_model import get_selected_entity_id_for_user, get_user_preferences
        # get_user_preferences returns a tuple: (preferences_dict, already_seen_list)
        preferences_dict, already_seen_list = get_user_preferences(user_id)

        # Get the grant IDs for likes and dislikes separately
        conn = get_connection()
        cursor = conn.cursor()
        selected_entity_id = get_selected_entity_id_for_user(user_id, cursor=cursor)

        # Get grants marked as 'interesa' (likes)
        if selected_entity_id:
            cursor.execute("""
                SELECT grant_id FROM user_grant_preferences
                WHERE user_id = %s AND entity_id = %s AND action = 'interesa'
                ORDER BY timestamp DESC
            """, (user_id, selected_entity_id))
        else:
            cursor.execute("""
                SELECT grant_id FROM user_grant_preferences
                WHERE user_id = %s AND action = 'interesa'
                ORDER BY timestamp DESC
            """, (user_id,))
        likes = [row[0] for row in cursor.fetchall()]

        # Get grants marked as 'no interesa' (dislikes)
        if selected_entity_id:
            cursor.execute("""
                SELECT grant_id FROM user_grant_preferences
                WHERE user_id = %s AND entity_id = %s AND action = 'no interesa'
                ORDER BY timestamp DESC
            """, (user_id, selected_entity_id))
        else:
            cursor.execute("""
                SELECT grant_id FROM user_grant_preferences
                WHERE user_id = %s AND action = 'no interesa'
                ORDER BY timestamp DESC
            """, (user_id,))
        dislikes = [row[0] for row in cursor.fetchall()]

        cursor.close()
        conn.close()

        # Return in the format expected by the frontend
        return jsonify({
            "preferences": {
                "likes": likes,
                "dislikes": dislikes
            },
            "insights": {
                "beneficiarios": preferences_dict.get("beneficiarios", []),
                "regiones": preferences_dict.get("region_impacto", []),
                "finalidades": preferences_dict.get("finalidad", []),
                "totalLikes": len(likes),
            }
        }), 200
    except Exception as e:
        logger.error(f"Error in get_preferences: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/reset_user_interests', methods=['POST'])
@require_auth
def reset_user_interests(user_id):
    """Reset all persisted recommendation preferences and swipe history signals for the current user."""
    connection = None
    cursor = None
    try:
        if not user_id:
            return jsonify({"error": "No se encontró el usuario"}), 400

        ensure_recommendation_tables()

        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute(
            """
            SELECT entity_id
            FROM user_entities
            WHERE user_id = %s
            """,
            (user_id,),
        )
        entity_ids = [str(row[0]) for row in cursor.fetchall() if row and row[0] is not None]

        cursor.execute("DELETE FROM user_grant_preferences WHERE user_id = %s", (user_id,))
        deleted_preferences = cursor.rowcount or 0

        cursor.execute("DELETE FROM reco_events WHERE user_id = %s", (user_id,))
        deleted_events = cursor.rowcount or 0

        connection.commit()

        if entity_ids:
            refresh_reco_entity_state()

        logger.info(
            "Reset de intereses completado para user_id=%s. deleted_preferences=%s deleted_events=%s",
            user_id,
            deleted_preferences,
            deleted_events,
        )

        return jsonify(
            {
                "status": "success",
                "deleted_preferences": deleted_preferences,
                "deleted_events": deleted_events,
            }
        ), 200
    except Exception as e:
        if connection:
            connection.rollback()
        logger.error(f"Error reseteando intereses para user_id={user_id}: {e}", exc_info=True)
        return jsonify({"error": "No se pudieron resetear los intereses"}), 500
    finally:
        try:
            if cursor:
                cursor.close()
            if connection:
                connection.close()
        except Exception:
            pass

@app.route('/api/get_grant_preference/<string:grant_id>', methods=['GET'])
@require_auth
def get_grant_preference(user_id, grant_id):
    """
    Get user's preference for a specific grant.
    Returns: { "preference": "interesa" | "no interesa" | null }
    """
    try:
        if not grant_id:
            return jsonify({"error": "Se requiere grant_id"}), 400

        conn = get_connection()
        cursor = conn.cursor()
        from utils.recommendation_model import get_selected_entity_id_for_user
        selected_entity_id = get_selected_entity_id_for_user(user_id, cursor=cursor)

        # Get the user's preference for this specific grant
        if selected_entity_id:
            cursor.execute("""
                SELECT action FROM user_grant_preferences
                WHERE user_id = %s AND entity_id = %s AND grant_id = %s
                ORDER BY timestamp DESC
                LIMIT 1
            """, (user_id, selected_entity_id, grant_id))
        else:
            cursor.execute("""
                SELECT action FROM user_grant_preferences
                WHERE user_id = %s AND grant_id = %s
                ORDER BY timestamp DESC
                LIMIT 1
            """, (user_id, grant_id))

        result = cursor.fetchone()
        cursor.close()
        conn.close()

        preference = result[0] if result else None

        return jsonify({"preference": preference}), 200
    except Exception as e:
        logger.error(f"Error in get_grant_preference: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

# Endpoint para actualizar preferencias de usuario (perfil)
@app.route('/api/update_user_preferences', methods=['POST'])
def api_update_user_preferences():
    try:
        data = request.json
        user_id = data.get('user_id')
        
        if not user_id:
            return jsonify({"error": "Se requiere user_id"}), 400
        
        # Aquí puedes guardar las preferencias de notificaciones, email, etc.
        # en una tabla de configuración de usuario si existe
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/get_recommended_grants', methods=['GET'])
@require_auth
def get_recommendations(user_id):
    """Endpoint para obtener subvenciones recomendadas"""
    logger.info(f"Getting recommended grants for user_id: {user_id}")
    try:
        # Get limit from query params, default to 50
        limit = request.args.get('limit', 50, type=int)
        logger.info(f"Fetching {limit} recommended grants")

        recommended_grants = get_recommended_grants(user_id, limit)
        logger.info(f"Recommended grants found: {len(recommended_grants)}")

        if recommended_grants:
            logger.info(f"Sample grant: {recommended_grants[0] if len(recommended_grants) > 0 else 'None'}")

        return jsonify({"grants": recommended_grants})

    except Exception as e:
        logger.error(f"Error en get_recommendations: {e}", exc_info=True)
        return jsonify({"grants": [], "error": str(e)}), 500

"""USER & ENTITIES ENDPOINTS"""
""" ENTITIES """
# Endpoint para obtener las entidades asociadas al usuario (un usuario puede tener varias entidades)
@app.route('/api/get_user_entities', methods=['GET'])
@require_auth
def get_user_entities(user_id):
    """Endpoint para obtener las entidades asociadas al usuario"""
    try:
        print(f"=== GET USER ENTITIES for user {user_id} ===")

        connection = get_connection()
        cursor = connection.cursor()
        print("Connected to the database successfully.")

        # Si no se proporciona user_id, obtenemos el usuario actual
        if not user_id:
            cursor.execute("SELECT id FROM app_user LIMIT 1")
            print("No user_id provided, fetching current user.")
            return jsonify(error="No se han cargado tus entidades, por favor intenta de nuevo."), 400

        # Obtener las entidades del usuario
        query = """
            SELECT e.id, e.razon_social, e.nif, ue.is_selected
            FROM entities e
            JOIN user_entities ue ON e.id = ue.entity_id
            WHERE ue.user_id = %s
            ORDER BY ue.is_selected DESC, e.razon_social
        """
        cursor.execute(query, (user_id,))
        rows = cursor.fetchall()
        
        entities = []
        for row in rows:
            entities.append({
                "id": str(row[0]),
                "razon_social": row[1],
                "nif": row[2] if row[2] else "",
                "is_selected": row[3]
            })
        
        return jsonify(entities=entities)
    
    except Exception as e:
        print("An error occurred:", e)
        return jsonify(error=str(e)), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()
            print("Database connection closed.")


@app.route('/api/select_user_entity', methods=['POST'])
@require_auth
def select_entity_for_user(user_id):
    """Select active entity for a user (single is_selected=true)."""
    try:
        data = request.get_json(silent=True) or {}
        entity_id = str(data.get('entity_id') or '').strip()

        if not entity_id:
            return jsonify({"error": "Se requiere entity_id"}), 400

        ensure_recommendation_tables()
        ok = select_user_entity(user_id=user_id, entity_id=entity_id)
        if not ok:
            return jsonify({"error": "Entidad no encontrada para este usuario"}), 404

        return jsonify({"status": "success", "entity_id": entity_id}), 200
    except Exception as e:
        logger.error(f"Error selecting entity for user={user_id}: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

# Endpoint para obtener subvenciones ya concedidas a una entidad (BDNS)
@app.route('/api/get_entity_awarded_grants', methods=['GET'])
@require_auth
def get_entity_awarded_grants(user_id):
    connection = None
    cursor = None
    try:
        entity_id = request.args.get('entity_id')
        limit = max(1, min(request.args.get('limit', 20, type=int) or 20, 50))

        connection = get_connection()
        cursor = connection.cursor()

        if entity_id:
            cursor.execute("""
                SELECT e.id, e.nif, e.razon_social
                FROM entities e
                JOIN user_entities ue ON e.id = ue.entity_id
                WHERE e.id = %s AND ue.user_id = %s
                LIMIT 1
            """, (entity_id, user_id))
        else:
            cursor.execute("""
                SELECT e.id, e.nif, e.razon_social
                FROM entities e
                JOIN user_entities ue ON e.id = ue.entity_id
                WHERE ue.user_id = %s
                ORDER BY ue.is_selected DESC, e.razon_social
                LIMIT 1
            """, (user_id,))

        row = cursor.fetchone()
        if not row:
            return jsonify(success=False, message="Entidad no encontrada para este usuario", awards=[]), 404

        entity_id_value, raw_nif, razon_social = row
        nif = _normalize_cif_nif(raw_nif)
        if not nif:
            return jsonify(
                success=True,
                awards=[],
                source=None,
                warning="La entidad no tiene CIF/NIF informado. Completa ese campo para consultar las concesiones en BDNS.",
                entity={"id": str(entity_id_value), "razon_social": razon_social or "", "nif": ""},
            )

        result = _fetch_bdns_concesiones_by_cif(nif, limit=limit)

        return jsonify(
            success=True,
            grants=result.get("items", []),
            awards=result.get("items", []),  # Compatibilidad con frontend previo
            total=result.get("total", 0),
            source="BDNS",
            warning=result.get("warning"),
            source_url="https://www.pap.hacienda.gob.es/bdnstrans/GE/es/concesiones/consulta",
            entity={"id": str(entity_id_value), "razon_social": razon_social or "", "nif": nif},
            bdnstrans_beneficiario_id=result.get("persona_id"),
        )
    except requests.RequestException as e:
        logger.error(f"Error consultando BDNS para entidad {request.args.get('entity_id')}: {e}")
        return jsonify(
            success=False,
            message="No se pudo consultar la BDNS en este momento",
            grants=[],
            awards=[],
        ), 502
    except Exception as e:
        logger.error(f"Error en get_entity_awarded_grants: {e}", exc_info=True)
        return jsonify(success=False, message=str(e), grants=[], awards=[]), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()

# Endpoint para obtener concesiones de minimis de una entidad (BDNS)
@app.route('/api/get_entity_minimis_grants', methods=['GET'])
@require_auth
def get_entity_minimis_grants(user_id):
    connection = None
    cursor = None
    try:
        entity_id = request.args.get('entity_id')
        page_size = max(10, min(request.args.get('page_size', 100, type=int) or 100, 200))
        max_pages = max(1, min(request.args.get('max_pages', 20, type=int) or 20, 50))

        connection = get_connection()
        cursor = connection.cursor()

        if entity_id:
            cursor.execute("""
                SELECT e.id, e.nif, e.razon_social
                FROM entities e
                JOIN user_entities ue ON e.id = ue.entity_id
                WHERE e.id = %s AND ue.user_id = %s
                LIMIT 1
            """, (entity_id, user_id))
        else:
            cursor.execute("""
                SELECT e.id, e.nif, e.razon_social
                FROM entities e
                JOIN user_entities ue ON e.id = ue.entity_id
                WHERE ue.user_id = %s
                ORDER BY ue.is_selected DESC, e.razon_social
                LIMIT 1
            """, (user_id,))

        row = cursor.fetchone()
        if not row:
            return jsonify(success=False, message="Entidad no encontrada para este usuario", minimis_awards=[]), 404

        entity_id_value, raw_nif, razon_social = row
        nif = _normalize_cif_nif(raw_nif)
        if not nif:
            return jsonify(
                success=True,
                minimis_grants=[],
                minimis_awards=[],
                total=0,
                total_last_3_years_amount=0,
                total_amount=0,
                count_last_3_years=0,
                source=None,
                warning="La entidad no tiene CIF/NIF informado. Completa ese campo para consultar minimis en BDNS.",
                entity={"id": str(entity_id_value), "razon_social": razon_social or "", "nif": ""},
            )

        result = _fetch_bdns_minimis_by_cif(nif, page_size=page_size, max_pages=max_pages)

        # Persistimos el total histórico de minimis en la entidad para tenerlo cacheado/en perfil.
        # Si existe la nueva columna `minimis`, se actualiza ahí. Mantiene compatibilidad con
        # `concesion_minimis` si aún está presente en algunos entornos.
        total_minimis_amount = _safe_float(result.get("total_amount"))
        if total_minimis_amount is not None:
            entity_columns_to_update = []
            try:
                if _table_has_column(cursor, "entities", "minimis"):
                    entity_columns_to_update.append("minimis")
                if _table_has_column(cursor, "entities", "concesion_minimis"):
                    entity_columns_to_update.append("concesion_minimis")
            except Exception as schema_check_error:
                logger.warning(f"No se pudo inspeccionar columnas minimis en entities: {schema_check_error}")

            if entity_columns_to_update:
                try:
                    set_clause = ", ".join([f"{col} = %s" for col in entity_columns_to_update])
                    params = [total_minimis_amount] * len(entity_columns_to_update) + [entity_id_value]
                    cursor.execute(
                        f"UPDATE entities SET {set_clause} WHERE id = %s",
                        tuple(params)
                    )
                    connection.commit()
                except Exception as update_error:
                    try:
                        connection.rollback()
                    except Exception:
                        pass
                    logger.warning(
                        f"No se pudo actualizar total minimis en entities (entity_id={entity_id_value}): {update_error}"
                    )

        return jsonify(
            success=True,
            minimis_grants=result.get("items", []),
            minimis_awards=result.get("items", []),  # Compatibilidad futura
            total=result.get("total", 0),
            total_last_3_years_amount=result.get("total_last_3_years_amount", 0),
            total_amount=result.get("total_amount", 0),
            count_last_3_years=result.get("count_last_3_years", 0),
            cutoff_date=result.get("cutoff_date"),
            source="BDNS",
            warning=result.get("warning"),
            source_url="https://www.pap.hacienda.gob.es/bdnstrans/GE/es/minimis/consulta",
            entity={"id": str(entity_id_value), "razon_social": razon_social or "", "nif": nif},
            bdnstrans_beneficiario_id=result.get("persona_id"),
        )
    except requests.RequestException as e:
        logger.error(f"Error consultando BDNS minimis para entidad {request.args.get('entity_id')}: {e}")
        return jsonify(
            success=False,
            message="No se pudo consultar la BDNS de minimis en este momento",
            minimis_grants=[],
            minimis_awards=[],
        ), 502
    except Exception as e:
        logger.error(f"Error en get_entity_minimis_grants: {e}", exc_info=True)
        return jsonify(success=False, message=str(e), minimis_grants=[], minimis_awards=[]), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()

# Endpoint para obtener el perfil de una entidad
@app.route('/api/get_entity_profile', methods=['GET'])
def get_entity_profile():
    """Endpoint para obtener los datos del perfil de una entidad específica"""
    print("=== GET ENTITY PROFILE START ===")
    try:
        entity_id = request.args.get('entity_id')
        print(f"Received entity_id: {entity_id}")
        
        if not entity_id:
            # Si no se proporciona entity_id, obtenemos la entidad seleccionada del usuario actual
            connection = get_connection()
            cursor = connection.cursor()
            print("Connected to the database successfully.")
            
            cursor.execute("""
                SELECT e.id FROM entities e
                JOIN user_entities ue ON e.id = ue.entity_id
                JOIN app_user u ON ue.user_id = u.id
                LIMIT 1
            """)
            
            entity_row = cursor.fetchone()
            if entity_row:
                entity_id = entity_row[0]
            else:
                # Si no hay una entidad seleccionada, buscar cualquier entidad del usuario
                cursor.execute("""
                    SELECT e.id FROM entities e
                    JOIN user_entities ue ON e.id = ue.entity_id
                    JOIN app_user u ON ue.user_id = u.id
                    ORDER BY ue.created_at DESC
                    LIMIT 1
                """)
                entity_row = cursor.fetchone()
        
        connection = get_connection()
        cursor = connection.cursor()
        print("Connected to the database successfully.")

        has_minimis = _table_has_column(cursor, 'entities', 'minimis')
        has_concesion_minimis = _table_has_column(cursor, 'entities', 'concesion_minimis')

        select_fields = [
            "razon_social", "nombre_representante", "pagina_web", "comunidad_autonoma",
            "comunidad_autonoma_centro_trabajo", "telefono", "correo",
            "nif", "descripcion", "tipo_empresa", "fecha_constitucion", "personal_en_linea", "liderado_por_mujeres",
            "porcentaje_liderado_por_mujeres", "sector", "facturacion_anual",
            "direccion_social", "cnae", "objeto_social", "administrador_cargo", "administrador_año"
        ]
        if has_minimis:
            select_fields.append("minimis")
        if has_concesion_minimis:
            select_fields.append("concesion_minimis")

        query = f"""
            SELECT {", ".join(select_fields)}
            FROM entities
            WHERE id = %s
        """
 
        cursor.execute(query, (entity_id,))
        row = cursor.fetchone()
        
        if not row:
            return jsonify(error="Entidad no encontrada"), 404
        
        row_data = dict(zip(select_fields, row))

        # Formatear la fecha de constitución si existe
        fecha_constitucion = row_data.get("fecha_constitucion")
        if fecha_constitucion and isinstance(fecha_constitucion, datetime):
            fecha_constitucion = fecha_constitucion.strftime("%Y-%m-%d")

        minimis_value = row_data.get("minimis")
        if minimis_value is None:
            minimis_value = row_data.get("concesion_minimis")
        minimis_value = float(minimis_value) if minimis_value is not None else None
        
        profile = {
            "razon_social": row_data.get("razon_social") or "",
            "nombre_representante": row_data.get("nombre_representante") or "",
            "pagina_web": row_data.get("pagina_web") or "",
            "comunidad_autonoma": row_data.get("comunidad_autonoma") or "",
            "comunidad_autonoma_centro_trabajo": row_data.get("comunidad_autonoma_centro_trabajo") or "",
            "telefono": row_data.get("telefono") or "",
            "correo": row_data.get("correo") or "",
            "nif": row_data.get("nif") or "",
            "descripcion": row_data.get("descripcion") or "",
            "tipo_empresa": row_data.get("tipo_empresa") or "",
            "fecha_constitucion": fecha_constitucion if fecha_constitucion else "",
            "personal_en_linea": row_data.get("personal_en_linea") if row_data.get("personal_en_linea") is not None else 0,
            "liderado_por_mujeres": row_data.get("liderado_por_mujeres") if row_data.get("liderado_por_mujeres") is not None else False,
            "porcentaje_liderado_por_mujeres": row_data.get("porcentaje_liderado_por_mujeres") if row_data.get("porcentaje_liderado_por_mujeres") is not None else 0,
            "sector": row_data.get("sector") or "",
            "facturacion_anual": f"{row_data.get('facturacion_anual')}€" if row_data.get("facturacion_anual") else "No especificado",
            "direccion_social": row_data.get("direccion_social") or "",
            "cnae": row_data.get("cnae") or "",
            "objeto_social": row_data.get("objeto_social") or "",
            "administrador_cargo": row_data.get("administrador_cargo") or "",
            "administrador_año": row_data.get("administrador_año") or "",
            "minimis": minimis_value,
            "concesion_minimis": minimis_value,
        }
 
        return jsonify(profile=profile)
    
    except Exception as e:
        print("An error occurred:", e)
        return jsonify(error=str(e)), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()
            print("Database connection closed.")

# Endpoint para actualizar el perfil de una entidad
@app.route('/api/update_entity_profile', methods=['POST'])
def update_entity_profile():
    """Endpoint para actualizar los datos del perfil de una entidad"""
    try:
        data = request.json
        entity_id = data['entity_id']
        profile = data['profile']

        print(f"Profile data to update: {profile}")

        # Extract and validate facturacion_anual
        facturacion_anual = profile.get('facturacion_anual')

        connection = get_connection()
        with connection:
            with connection.cursor() as cursor:
                # Verificar si la entidad existe
                cursor.execute("SELECT id FROM entities WHERE id = %s", (entity_id,))
                if not cursor.fetchone():
                    return jsonify(success=False, message="Entidad no encontrada"), 404
                
                # Ensure numeric fields are properly parsed
                personal_en_linea = profile.get('personal_en_linea')
                if personal_en_linea == '':
                    personal_en_linea = None
                elif personal_en_linea is not None:
                    try:
                        personal_en_linea = int(personal_en_linea)
                    except (ValueError, TypeError):
                        personal_en_linea = None
                        
                porcentaje_liderado_por_mujeres = profile.get('porcentaje_liderado_por_mujeres')
                if porcentaje_liderado_por_mujeres == '':
                    porcentaje_liderado_por_mujeres = None
                elif porcentaje_liderado_por_mujeres is not None:
                    try:
                        porcentaje_liderado_por_mujeres = int(porcentaje_liderado_por_mujeres)
                    except (ValueError, TypeError):
                        porcentaje_liderado_por_mujeres = None

                minimis_value = profile.get('minimis', profile.get('concesion_minimis'))
                if minimis_value == '':
                    minimis_value = None
                elif minimis_value is not None:
                    try:
                        minimis_value = float(minimis_value)
                    except (ValueError, TypeError):
                        minimis_value = None
                
                # Format date in ISO format if provided
                fecha_constitucion = profile.get('fecha_constitucion')
                print(f"Received fecha_constitucion: {fecha_constitucion}")
                if fecha_constitucion:
                    try:
                        # Ensure date is in ISO format for PostgreSQL
                        from datetime import datetime
                        # Parse any potential date format into a datetime object
                        if '/' in fecha_constitucion:
                            # Handle DD/MM/YYYY format
                            parts = fecha_constitucion.split('/')
                            if len(parts) == 3:
                                fecha_constitucion = f"{parts[2]}-{parts[1]}-{parts[0]}"
                        # Ensure it's a valid date
                        datetime.fromisoformat(fecha_constitucion.replace('Z', '+00:00'))
                        print(f"Parsed fecha_constitucion: {fecha_constitucion}")
                    except (ValueError, TypeError) as e:
                        print(f"Date parsing error: {e}")
                        # Keep the original value if parsing fails
                else:
                    fecha_constitucion = None
                    print(f"Received fecha_constitucion: {fecha_constitucion}")

                has_minimis = _table_has_column(cursor, 'entities', 'minimis')
                has_concesion_minimis = _table_has_column(cursor, 'entities', 'concesion_minimis')

                update_fields = [
                    "razon_social = %s",
                    "nombre_representante = %s",
                    "pagina_web = %s",
                    "comunidad_autonoma = %s",
                    "comunidad_autonoma_centro_trabajo = %s",
                    "telefono = %s",
                    "correo = %s",
                    "nif = %s",
                    "descripcion = %s",
                    "tipo_empresa = %s",
                    "fecha_constitucion = %s",
                    "personal_en_linea = %s",
                    "liderado_por_mujeres = %s",
                    "porcentaje_liderado_por_mujeres = %s",
                    "sector = %s",
                    "facturacion_anual = %s",
                    "direccion_social = %s",
                    "cnae = %s",
                    "objeto_social = %s",
                    "administrador_cargo = %s",
                    "administrador_año = %s",
                ]
                update_values = [
                    profile.get('razon_social'),
                    profile.get('nombre_representante'),
                    profile.get('pagina_web'),
                    profile.get('comunidad_autonoma'),
                    profile.get('comunidad_autonoma_centro_trabajo'),
                    profile.get('telefono'),
                    profile.get('correo'),
                    profile.get('nif'),
                    profile.get('descripcion'),
                    profile.get('tipo_empresa'),
                    fecha_constitucion,
                    personal_en_linea,
                    profile.get('liderado_por_mujeres'),
                    porcentaje_liderado_por_mujeres,
                    profile.get('sector'),
                    facturacion_anual,
                    profile.get('direccion_social'),
                    profile.get('cnae'),
                    profile.get('objeto_social'),
                    profile.get('administrador_cargo'),
                    profile.get('administrador_año'),
                ]

                if has_minimis:
                    update_fields.append("minimis = %s")
                    update_values.append(minimis_value)

                if has_concesion_minimis:
                    update_fields.append("concesion_minimis = %s")
                    update_values.append(minimis_value)

                update_query = f"""
                    UPDATE entities
                    SET {", ".join(update_fields)}
                    WHERE id = %s
                """

                cursor.execute(update_query, (*update_values, entity_id))
                
                connection.commit()

        return jsonify(success=True, message="Perfil actualizado correctamente")
    except Exception as e:
        print("Error al actualizar el perfil de la entidad:", e)
        return jsonify(success=False, message=str(e)), 500

# Endpoint para crear una nueva entidad del usuario
@app.route('/api/create_entity', methods=['POST'])
@require_auth
def create_entity(user_id):
    """Endpoint para crear una nueva entidad con datos completos"""
    try:
        logger.info(f"=== CREATE ENTITY START ===")
        logger.info(f"Received user_id from auth: {user_id}")
        
        # Campos
        razon_social    = request.form.get('razon_social')
        nif             = request.form.get('nif')
  

        # Validar campos obligatorios
        if not razon_social or not nif:
            logger.info("Faltan datos obligatorios: razon_social, nif")
            return jsonify(success=False, message="Faltan datos obligatorios (razon_social, nif)"), 400
        
        if request.files.getlist('files'):
            files = request.files.getlist('files')
            logger.info(f"Archivos recibidos: {len(files)}")
        else:
            files = ""
        
        if request.form.get('pagina_web'):
            pagina_web      = request.form.get('pagina_web')
            logger.info(f"Página web recibida: {pagina_web}")
        else:
            pagina_web = ""

        if request.form.get('descripcion_usuario'):
            descripcion_usuario = request.form.get('descripcion_usuario', '')
            logger.info("Descripción de usuario recibida.")
        else:
            descripcion_usuario = ""

        logger.info(f"Razon Social: {razon_social}, NIF: {nif}")
        connection = get_connection()
        cursor = connection.cursor()
        logger.info("Conexión a la base de datos establecida.")

        # Si no se proporciona user_id, se busca el usuario actual
        if not user_id:
            logger.error("No se pudo validar el usuario autenticado en create_entity.")
            return jsonify(success=False, message="Usuario no autenticado"), 401

        # 0. Verificar si ya existe una entidad con el mismo NIF para el usuario
        check_query = """
            SELECT e.id 
            FROM entities e
            WHERE e.nif = %s
        """
        cursor.execute(check_query, (nif,))
        existing = cursor.fetchone()
        if existing:
            logger.info(f"Ya existe una entidad con nif: {nif} para el usuario {user_id}")
            return jsonify(success=False, message="Ya existe una entidad con ese NIF"), 409

        logger.info(f"No existe una entidad con nif: {nif} para el usuario {user_id}")

        # 1. Crear la entidad primero
        insert_entity_query = """
            INSERT INTO entities (razon_social, nif, pagina_web, created_at, updated_at)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id
        """
        
        cursor.execute(insert_entity_query, (
            razon_social,
            nif,
            pagina_web,
            datetime.now(),
            datetime.now()
        ))
        
        entity_id = cursor.fetchone()[0]
        connection.commit()
        
        logger.info(f"Entidad creada con ID: {entity_id}")
        
        # 2. Procesar archivos si existen
        if files and len(files) > 0:
            uploaded_files = []
            logger.info(f"Procesando {len(files)} archivos para la entidad {entity_id}")
            
            for file in files:
                if file.filename == '':
                    logger.info("Archivo sin nombre, omitiendo...")
                    continue
                    
                # Generar nombre único para el archivo
                file_extension = os.path.splitext(file.filename)[1]
                unique_filename = f"{entity_id}/{uuid.uuid4()}{file_extension}"
                safe_filename = secure_filename(file.filename)

                # 1) Lee todo el contenido binario de golpe
                raw_bytes = file.read()

                # 2) Prepara un BytesIO para la extracción de texto
                bio_for_text = io.BytesIO(raw_bytes)
                bio_for_text.seek(0)
                try:
                    file_text = extract_text(bio_for_text, file_extension)
                    if not file_text:
                        logger.info(f"El archivo {safe_filename} no contiene texto.")
                        continue
                    logger.info(f"Contenido extraído: {file_text[:200]}…")
                except Exception as e:
                    logger.error("Error al extraer texto:", e)
                    continue

                logger.info(f"Extrayendo contenido del archivo: {safe_filename} como {unique_filename}")

                # 3) Prepara otro BytesIO para subir a S3
                bio_for_s3 = io.BytesIO(raw_bytes)
                bio_for_s3.seek(0)
                try:
                    logger.info("Subiendo archivo a S3...")
                    # Subir archivo a S3
                    s3_client.upload_fileobj(
                        Fileobj=bio_for_s3,             # o file.stream
                        Bucket=S3_BUCKET_NAME,
                        Key=unique_filename,
                        ExtraArgs={
                            'Metadata': {
                                'entity_id': str(entity_id),
                                'original_filename': safe_filename,
                                'upload_timestamp': datetime.now().isoformat()
                            }
                        }
                    )
                    

                    # Obtener el tamaño del archivo
                    file.seek(0, 2)  # Ir al final del archivo
                    file_size = file.tell()
                    file.seek(0)  # Volver al inicio

                except Exception as s3_error:
                    logger.error(f"Error al subir archivo {safe_filename} a S3: {str(s3_error)}")
                    continue
                
                logger.info(f"Subido `{unique_filename}` a s3://{S3_BUCKET_NAME}/{unique_filename}")

                try:
                    logger.info("Guardando en sql...")
                    # Guardar referencia en la base de datos
                    insert_file_query = """
                        INSERT INTO entity_documents (entity_id, original_filename, s3_key, s3_bucket, file_size, analysis_result, upload_date, status)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING id
                    """
                    
                    cursor.execute(insert_file_query, (
                        entity_id,
                        safe_filename,
                        unique_filename,
                        S3_BUCKET_NAME,
                        file_size,
                        file_text,
                        datetime.now(),
                        'uploaded'
                    ))
                    
                    file_id = cursor.fetchone()[0]
                    
                    uploaded_files.append({
                        'id': file_id,
                        'filename': safe_filename,
                        's3_key': unique_filename,
                        'status': 'uploaded'
                    })
                    
                    logger.info(f"Archivo subido exitosamente: {safe_filename} -> {unique_filename}")
                    
                except Exception as upload_error:
                    logger.error(f"Error al subir archivo {safe_filename}: {str(upload_error)}")
                    continue
            connection.commit()
        else:
            uploaded_files = ""
            logger.info("No se proporcionaron archivos para subir.")
        
        # Registrar la nueva entidad y marcarla como seleccionada (única seleccionada por usuario)
        cursor.execute(
            """
            UPDATE user_entities
            SET is_selected = FALSE
            WHERE user_id = %s
            """,
            (user_id,),
        )
        link_query = """
            INSERT INTO user_entities (user_id, entity_id, is_selected)
            VALUES (%s, %s, TRUE)
        """
        cursor.execute(link_query, (user_id, entity_id))
        connection.commit()

        logger.info(f"Entidad creada con éxito: ID {entity_id}")

        _set_entity_processing_status(
            entity_id=entity_id,
            user_id=user_id,
            status="running",
            stage="queued",
            message="Entidad creada. Iniciando análisis automático...",
            progress=5,
            processed_items=0,
            total_items=0,
            matches_found=0,
            started_at=datetime.utcnow(),
        )

        def run_entity_pipeline_background(
            entity_id_value,
            user_id_value,
            pagina_web_value,
            razon_social_value,
            nif_value,
            descripcion_usuario_value,
            uploaded_files_value,
        ):
            status_snapshot = {
                "processed_items": 0,
                "total_items": 0,
                "matches_found": 0,
                "best_match_score": None,
                "first_high_match_found": False,
                "first_high_match_grant_id": None,
                "first_high_match_score": None,
            }
            try:
                _set_entity_processing_status(
                    entity_id=entity_id_value,
                    user_id=user_id_value,
                    status="running",
                    stage="scraping",
                    message="Analizando la información de tu entidad...",
                    progress=12,
                    processed_items=0,
                    total_items=0,
                    matches_found=0,
                    started_at=datetime.utcnow(),
                )

                final_response = scrape_main(
                    pagina_web_value,
                    razon_social_value,
                    nif_value,
                    entity_id_value,
                    descripcion_usuario_value,
                    uploaded_files_value,
                    OPENAI_PROMPT_SCRAPE_ENTITY,
                )
                logger.info(f"[ENTITY PIPELINE] Scraping output entity_id={entity_id_value}: {type(final_response)}")

                if final_response:
                    bg_connection = None
                    bg_cursor = None
                    try:
                        bg_connection = get_connection()
                        bg_cursor = bg_connection.cursor()

                        fecha_constitucion = _parse_sql_date_like(final_response.get("fecha_constitucion"))

                        update_entity_query = """
                            UPDATE entities
                            SET
                            nombre_representante                   = %s,
                            comunidad_autonoma                     = %s,
                            comunidad_autonoma_centro_trabajo      = %s,
                            telefono                               = %s,
                            correo                                 = %s,
                            descripcion                            = %s,
                            tipo_empresa                           = %s,
                            fecha_constitucion                     = %s,
                            personal_en_linea                      = %s,
                            liderado_por_mujeres                   = %s,
                            porcentaje_liderado_por_mujeres        = %s,
                            sector                                 = %s,
                            facturacion_anual                      = %s,
                            direccion_social                       = %s,
                            cnae                                   = %s,
                            objeto_social                          = %s,
                            administrador_cargo                    = %s,
                            administrador_año                      = %s,
                            updated_at                             = %s
                            WHERE id = %s
                        """
                        bg_cursor.execute(
                            update_entity_query,
                            (
                                final_response.get("nombre_representante"),
                                final_response.get("comunidad_autonoma"),
                                final_response.get("comunidad_autonoma_centro_trabajo"),
                                final_response.get("telefono"),
                                final_response.get("correo"),
                                final_response.get("descripcion"),
                                final_response.get("tipo_entidad"),
                                fecha_constitucion,
                                final_response.get("personal_en_linea"),
                                final_response.get("liderado_por_mujeres"),
                                final_response.get("porcentaje_liderado_por_mujeres"),
                                final_response.get("sector"),
                                final_response.get("facturacion_anual"),
                                final_response.get("direccion_social"),
                                final_response.get("cnae"),
                                final_response.get("objeto_social"),
                                final_response.get("administrador_cargo"),
                                final_response.get("administrador_año"),
                                datetime.now(),
                                entity_id_value,
                            ),
                        )
                        bg_connection.commit()
                        logger.info("[ENTITY PIPELINE] Entidad actualizada con datos de scraping: %s", entity_id_value)
                        _set_entity_processing_status(
                            entity_id=entity_id_value,
                            user_id=user_id_value,
                            status="running",
                            stage="scraping",
                            message="Perfil enriquecido correctamente. Iniciando matching...",
                            progress=32,
                            processed_items=0,
                            total_items=0,
                            matches_found=0,
                        )
                    except Exception as update_error:
                        logger.warning(
                            "[ENTITY PIPELINE] No se pudo actualizar la entidad con scraping (entity_id=%s): %s",
                            entity_id_value,
                            update_error,
                        )
                        if bg_connection:
                            try:
                                bg_connection.rollback()
                            except Exception:
                                pass
                    finally:
                        if bg_cursor:
                            bg_cursor.close()
                        if bg_connection:
                            bg_connection.close()
                else:
                    _set_entity_processing_status(
                        entity_id=entity_id_value,
                        user_id=user_id_value,
                        status="running",
                        stage="scraping",
                        message="No se pudo completar todo el scraping. Continuamos con el matching.",
                        progress=28,
                        processed_items=0,
                        total_items=0,
                        matches_found=0,
                    )

                _set_entity_processing_status(
                    entity_id=entity_id_value,
                    user_id=user_id_value,
                    status="running",
                    stage="matching",
                    message="Buscando subvenciones compatibles...",
                    progress=35,
                    processed_items=0,
                    total_items=0,
                    matches_found=0,
                )

                def matching_progress_callback(event):
                    if not isinstance(event, dict):
                        return

                    total_items = int(event.get("total", status_snapshot["total_items"]) or 0)
                    processed_items = int(event.get("processed", status_snapshot["processed_items"]) or 0)
                    matches_found = int(event.get("matches_found", status_snapshot["matches_found"]) or 0)
                    best_match_score = event.get("best_match_score", status_snapshot["best_match_score"])

                    if total_items > 0:
                        status_snapshot["total_items"] = total_items
                    status_snapshot["processed_items"] = processed_items
                    status_snapshot["matches_found"] = matches_found
                    status_snapshot["best_match_score"] = best_match_score

                    event_type = event.get("type")
                    phase_label = event.get("phase_label")
                    message = phase_label or "Buscando subvenciones compatibles..."

                    if event_type == "first_high_match":
                        status_snapshot["first_high_match_found"] = True
                        status_snapshot["first_high_match_grant_id"] = event.get("grant_id")
                        status_snapshot["first_high_match_score"] = event.get("score")
                        message = "¡Encontramos una subvención con más del 85% de compatibilidad!"

                    total_for_progress = max(status_snapshot["total_items"], 1)
                    processed_for_progress = min(status_snapshot["processed_items"], total_for_progress)
                    matching_progress = 35 + (processed_for_progress / total_for_progress) * 60

                    _set_entity_processing_status(
                        entity_id=entity_id_value,
                        user_id=user_id_value,
                        status="running",
                        stage="matching",
                        message=message,
                        progress=matching_progress,
                        processed_items=status_snapshot["processed_items"],
                        total_items=status_snapshot["total_items"],
                        matches_found=status_snapshot["matches_found"],
                        best_match_score=status_snapshot["best_match_score"],
                        first_high_match_found=status_snapshot["first_high_match_found"],
                        first_high_match_grant_id=status_snapshot["first_high_match_grant_id"],
                        first_high_match_score=status_snapshot["first_high_match_score"],
                    )

                matching_result = main_match_new_client(
                    entity_id_value,
                    progress_callback=matching_progress_callback,
                )

                if isinstance(matching_result, dict) and matching_result.get("status") == "completed":
                    first_high = matching_result.get("first_high_match") or {}
                    _set_entity_processing_status(
                        entity_id=entity_id_value,
                        user_id=user_id_value,
                        status="completed",
                        stage="completed",
                        message="Análisis finalizado. Te enviaremos el correo con todos tus matches.",
                        progress=100,
                        processed_items=matching_result.get("processed", status_snapshot["processed_items"]),
                        total_items=matching_result.get("total", status_snapshot["total_items"]),
                        matches_found=matching_result.get("matches_found", status_snapshot["matches_found"]),
                        best_match_score=matching_result.get("best_match_score", status_snapshot["best_match_score"]),
                        first_high_match_found=bool(first_high),
                        first_high_match_grant_id=first_high.get("grant_id"),
                        first_high_match_score=first_high.get("score"),
                        completed_at=datetime.utcnow(),
                    )
                else:
                    error_message = (
                        matching_result.get("error")
                        if isinstance(matching_result, dict)
                        else "No se pudo completar el matching."
                    )
                    _set_entity_processing_status(
                        entity_id=entity_id_value,
                        user_id=user_id_value,
                        status="error",
                        stage="matching",
                        message="Hubo un error al generar los matches.",
                        progress=100,
                        processed_items=status_snapshot["processed_items"],
                        total_items=status_snapshot["total_items"],
                        matches_found=status_snapshot["matches_found"],
                        best_match_score=status_snapshot["best_match_score"],
                        first_high_match_found=status_snapshot["first_high_match_found"],
                        first_high_match_grant_id=status_snapshot["first_high_match_grant_id"],
                        first_high_match_score=status_snapshot["first_high_match_score"],
                        error=error_message,
                        completed_at=datetime.utcnow(),
                    )
            except Exception as pipeline_error:
                logger.error(
                    "[ENTITY PIPELINE] Error en background para entity_id=%s: %s",
                    entity_id_value,
                    pipeline_error,
                    exc_info=True,
                )
                _set_entity_processing_status(
                    entity_id=entity_id_value,
                    user_id=user_id_value,
                    status="error",
                    stage="failed",
                    message="No se pudo completar el procesamiento de la entidad.",
                    progress=100,
                    processed_items=status_snapshot["processed_items"],
                    total_items=status_snapshot["total_items"],
                    matches_found=status_snapshot["matches_found"],
                    best_match_score=status_snapshot["best_match_score"],
                    first_high_match_found=status_snapshot["first_high_match_found"],
                    first_high_match_grant_id=status_snapshot["first_high_match_grant_id"],
                    first_high_match_score=status_snapshot["first_high_match_score"],
                    error=str(pipeline_error),
                    completed_at=datetime.utcnow(),
                )

        pipeline_thread = threading.Thread(
            target=run_entity_pipeline_background,
            args=(
                entity_id,
                user_id,
                pagina_web,
                razon_social,
                nif,
                descripcion_usuario,
                uploaded_files,
            ),
            daemon=True,
        )
        pipeline_thread.start()
        logger.info("[ENTITY PIPELINE] Hilo iniciado para entity_id=%s", entity_id)

        return jsonify(
            success=True,
            message="La entidad se ha creado correctamente. El análisis está en curso.",
            entity_id=str(entity_id),
            processing_started=True,
        )

    except Exception as e:
        logger.error("Ocurrió un error:", e)
        return jsonify(success=False, error=str(e)), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()
            logger.info("Conexión a la base de datos cerrada.")

# Endpoint para consultar estado de procesamiento de una entidad (scraping + matching)
@app.route('/api/entity_processing_status', methods=['GET'])
@require_auth
def get_entity_processing_status(user_id):
    connection = None
    cursor = None
    try:
        entity_id = request.args.get('entity_id')
        if not entity_id:
            return jsonify(success=False, message="entity_id es obligatorio"), 400

        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute(
            """
            SELECT 1
            FROM user_entities
            WHERE user_id = %s AND entity_id = %s
            LIMIT 1
            """,
            (user_id, entity_id),
        )
        if not cursor.fetchone():
            return jsonify(success=False, message="Entidad no encontrada para este usuario"), 404

        _ensure_entity_processing_status_table(cursor)
        cursor.execute(
            """
            SELECT
                status,
                stage,
                message,
                progress,
                processed_items,
                total_items,
                matches_found,
                best_match_score,
                first_high_match_found,
                first_high_match_grant_id,
                first_high_match_score,
                started_at,
                completed_at,
                error,
                updated_at
            FROM public.entity_processing_status
            WHERE entity_id = %s AND user_id = %s
            LIMIT 1
            """,
            (str(entity_id), str(user_id)),
        )
        row = cursor.fetchone()

        if not row:
            return jsonify(
                success=True,
                status={
                    "entity_id": str(entity_id),
                    "status": "idle",
                    "stage": "idle",
                    "message": "Sin procesos en ejecución",
                    "progress": 0,
                    "processed_items": 0,
                    "total_items": 0,
                    "matches_found": 0,
                    "best_match_score": None,
                    "first_high_match": None,
                    "started_at": None,
                    "completed_at": None,
                    "error": None,
                    "updated_at": None,
                },
            )

        (
            status_value,
            stage_value,
            message_value,
            progress_value,
            processed_items,
            total_items,
            matches_found,
            best_match_score,
            first_high_match_found,
            first_high_match_grant_id,
            first_high_match_score,
            started_at,
            completed_at,
            error_value,
            updated_at,
        ) = row

        first_high_match_payload = None
        if first_high_match_found:
            first_high_match_payload = {
                "grant_id": str(first_high_match_grant_id) if first_high_match_grant_id is not None else None,
                "score": float(first_high_match_score) if first_high_match_score is not None else None,
            }

        return jsonify(
            success=True,
            status={
                "entity_id": str(entity_id),
                "status": status_value,
                "stage": stage_value,
                "message": message_value,
                "progress": float(progress_value) if progress_value is not None else 0.0,
                "processed_items": int(processed_items or 0),
                "total_items": int(total_items or 0),
                "matches_found": int(matches_found or 0),
                "best_match_score": float(best_match_score) if best_match_score is not None else None,
                "first_high_match": first_high_match_payload,
                "started_at": started_at.isoformat() if started_at else None,
                "completed_at": completed_at.isoformat() if completed_at else None,
                "error": error_value,
                "updated_at": updated_at.isoformat() if updated_at else None,
            },
        )
    except Exception as e:
        logger.error(f"Error en get_entity_processing_status: {e}", exc_info=True)
        return jsonify(success=False, message="No se pudo obtener el estado de procesamiento", error=str(e)), 500
    finally:
        if cursor:
            cursor.close()
        if connection:
            connection.close()

# Endpoint para eliminar una entidad
@app.route('/api/delete_entity', methods=['DELETE'])
def delete_entity():
    """Endpoint para eliminar una entidad específica"""
    try:
        data = request.json
        entity_id = data.get('entity_id')

        logger.info(f"=== DELETE ENTITY START ===")
        logger.info(f"Received entity_id: {entity_id}")
        
        if not entity_id:
            logger.error("No se ha proporcionado entity_id")
            return jsonify(success=False, message="No se ha podido eliminar la entidad"), 400

        connection = get_connection()
        cursor = connection.cursor()
        logger.info("Connected to the database successfully.")
        
        # Verificar que la entidad existe y obtener información sobre las entidades del usuario
        cursor.execute("""
            SELECT COUNT(*) as entity_count
            FROM entities e
            JOIN user_entities ue ON e.id = ue.entity_id
            JOIN app_user u ON ue.user_id = u.id
        """)
        total_entities = cursor.fetchone()[0]

        logger.info(f"Total entities associated with the user: {total_entities}")
        
        try:
            # Eliminar relaciones en user_entities
            cursor.execute("DELETE FROM user_entities WHERE entity_id = %s", (entity_id,))
            logger.info(f"Deleted FROM user_entities for entity_id: {entity_id}")
            
            # Eliminar matches asociados a la entidad
            cursor.execute("DELETE FROM matches WHERE entity_id = %s", (entity_id,))
            logger.info(f"Deleted FROM matches for entity_id: {entity_id}")

            # Eliminar estado de procesamiento asociado a la entidad
            try:
                _ensure_entity_processing_status_table(cursor)
                cursor.execute("DELETE FROM entity_processing_status WHERE entity_id = %s", (entity_id,))
                logger.info(f"Deleted FROM entity_processing_status for entity_id: {entity_id}")
            except Exception as status_delete_error:
                logger.warning(f"No se pudo limpiar entity_processing_status para entity_id={entity_id}: {status_delete_error}")
            
            # Eliminar la entidad
            cursor.execute("DELETE FROM entities WHERE id = %s", (entity_id,))
            logger.info(f"Deleted FROM entities with id: {entity_id}")
            
            # Confirmar la transacción
            connection.commit()
            
            return jsonify(success=True, message="Entidad eliminada correctamente")
            
        except Exception as e:
            raise e
            
    except Exception as e:
        logger.error("Error al eliminar la entidad:", e)
        return jsonify(success=False, message=str(e)), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()
            logger.info("Database connection closed.")
    
# Endpoint para obtener los documentos de una entidad
@app.route('/api/get_entity_documents/<string:entity_id>', methods=['GET'])
@require_auth
def get_entity_documents(user_id, entity_id):
    """Endpoint para obtener los documentos asociados a una entidad"""
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        _ensure_entity_documents_table(cursor)

        cursor.execute(
            """
            SELECT 1
            FROM user_entities
            WHERE user_id = %s AND entity_id = %s
            LIMIT 1
            """,
            (user_id, entity_id),
        )
        if cursor.fetchone() is None:
            return jsonify(error="Entidad no encontrada para este usuario"), 404

        _ensure_corporate_document_tables(cursor)
        _sync_legacy_corporate_documents(cursor)
        _backfill_inferred_entity_document_types(cursor, entity_id)
        connection.commit()
        
        query = """
            SELECT
                d.id,
                d.original_filename,
                d.s3_key,
                d.s3_bucket,
                d.file_size,
                d.upload_date,
                d.status,
                d.document_type_code,
                d.mime_type,
                c.label
            FROM public.entity_documents d
            LEFT JOIN public.corporate_document_catalog c
                ON c.code = d.document_type_code
            WHERE d.entity_id = %s
              AND COALESCE(d.is_current, TRUE) IS TRUE
            ORDER BY d.upload_date DESC
        """
        cursor.execute(query, (entity_id,))
        rows = cursor.fetchall()
        
        documents = []
        for row in rows:
            upload_date = row[5]
            if upload_date and isinstance(upload_date, datetime):
                upload_date = upload_date.isoformat()
            
            documents.append({
                "id": row[0],
                "filename": row[1],
                "s3_key": row[2],
                "s3_bucket": row[3],
                "file_size": row[4],
                "upload_date": upload_date,
                "status": row[6],
                "document_type_code": row[7],
                "mime_type": row[8],
                "document_type_label": row[9],
            })
        
        return jsonify(documents=documents)
    
    except Exception as e:
        logger.error(
            "Error fetching entity documents for user_id=%s entity_id=%s: %s",
            user_id,
            entity_id,
            e,
            exc_info=True,
        )
        return jsonify(error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass

# Endpoint para subir un documento y analizarlo con OpenAI
@app.route('/api/upload_entity_document', methods=['POST'])
@require_auth
def upload_entity_document(user_id):
    """Endpoint para subir un documento, guardarlo en S3, analizarlo con OpenAI y actualizar la entidad"""
    connection = None
    cursor = None
    try:
        entity_id = request.form.get('entity_id')
        document_type_code = (request.form.get('document_type_code') or '').strip() or None
        files = request.files.getlist('file')
        connection = get_connection()
        cursor = connection.cursor()
        logger.info("Conexión a la base de datos establecida.")
        _ensure_entity_documents_table(cursor)
        _ensure_corporate_document_tables(cursor)

        cursor.execute(
            """
            SELECT 1
            FROM user_entities
            WHERE user_id = %s AND entity_id = %s
            LIMIT 1
            """,
            (user_id, entity_id),
        )
        if cursor.fetchone() is None:
            return jsonify(success=False, message="Entidad no encontrada para este usuario"), 404

        if document_type_code:
            cursor.execute("""
                SELECT 1
                FROM public.corporate_document_catalog
                WHERE code = %s
                LIMIT 1
            """, (document_type_code,))
            if cursor.fetchone() is None:
                return jsonify(success=False, message="Tipo documental no válido"), 400

        if entity_id and len(files) > 0:
            uploaded_files = []
            logger.info(f"Procesando {len(files)} archivos para la entidad {entity_id}")
            
            for file in files:
                if file.filename == '':
                    logger.info("Archivo sin nombre, omitiendo...")
                    continue
                try:
                    stored_document = _store_entity_document(
                        cursor,
                        entity_id=entity_id,
                        uploaded_file=file,
                        uploaded_by=user_id,
                        document_type_code=document_type_code,
                    )
                    uploaded_files.append(stored_document)
                    logger.info(
                        "Archivo subido exitosamente: %s -> %s",
                        stored_document["filename"],
                        stored_document["s3_key"],
                    )
                except Exception as upload_error:
                    logger.error("Error al subir archivo %s: %s", file.filename, upload_error)
                    continue
            connection.commit()
        else:
            uploaded_files = []
            logger.info("No se proporcionaron archivos para subir.")

        if not uploaded_files:
            return jsonify(success=False, message="No se pudo subir ningún documento"), 400
        
        return jsonify(
            success=True, 
            message="Documento subido correctamente. Se está analizando en segundo plano.",
            document_id=uploaded_files[-1]["id"],
            documents=uploaded_files,
        )
    
    except Exception as e:
        logger.error(
            "Error uploading entity document for user_id=%s entity_id=%s: %s",
            user_id,
            request.form.get('entity_id'),
            e,
            exc_info=True,
        )
        return jsonify(success=False, error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass

# Preferencias manuales para refinar intereses en la página de matches
@app.route('/api/matches/refine-interests', methods=['GET'])
@require_auth
def get_match_refine_interests(user_id):
    try:
        connection = get_connection()
        cursor = connection.cursor()
        has_admin_col = _table_has_column(cursor, "user_refine_interests", "administraciones_convocantes")
        has_tipo_col = _table_has_column(cursor, "user_refine_interests", "tipos_ayuda")
        has_fecha_inicio_col = _table_has_column(cursor, "user_refine_interests", "fecha_inicio_solicitud")
        has_fecha_cierre_col = _table_has_column(cursor, "user_refine_interests", "fecha_de_cierre")

        admin_select = "COALESCE(administraciones_convocantes, '[]'::jsonb)," if has_admin_col else "'[]'::jsonb AS administraciones_convocantes,"
        tipo_select = "COALESCE(tipos_ayuda, '[]'::jsonb)" if has_tipo_col else "'[]'::jsonb AS tipos_ayuda"
        fecha_inicio_select = "fecha_inicio_solicitud," if has_fecha_inicio_col else "NULL::date AS fecha_inicio_solicitud,"
        fecha_cierre_select = "fecha_de_cierre," if has_fecha_cierre_col else "NULL::date AS fecha_de_cierre,"

        cursor.execute(f"""
            SELECT
                importe_min,
                importe_max,
                {fecha_inicio_select}
                {fecha_cierre_select}
                COALESCE(beneficiarios, '[]'::jsonb),
                COALESCE(regiones, '[]'::jsonb),
                COALESCE(finalidades, '[]'::jsonb),
                {admin_select}
                {tipo_select}
            FROM user_refine_interests
            WHERE user_id = %s
            LIMIT 1
        """, (user_id,))
        row = cursor.fetchone()

        if not row:
            return jsonify({
                "interests": {
                    "importeMin": "",
                    "importeMax": "",
                    "fechaDeInicio": "",
                    "fechaDeCierre": "",
                    "beneficiarios": [],
                    "regiones": [],
                    "finalidades": [],
                    "administraciones_convocantes": [],
                    "tipos_ayuda": [],
                }
            }), 200

        (
            importe_min,
            importe_max,
            fecha_inicio_solicitud,
            fecha_de_cierre,
            beneficiarios,
            regiones,
            finalidades,
            administraciones_convocantes,
            tipos_ayuda,
        ) = row

        def parse_json_array(value):
            if isinstance(value, list):
                return value
            if isinstance(value, str):
                try:
                    parsed = json.loads(value)
                    return parsed if isinstance(parsed, list) else []
                except Exception:
                    return []
            return []

        return jsonify({
            "interests": {
                "importeMin": "" if importe_min in (None, "") else str(importe_min),
                "importeMax": "" if importe_max in (None, "") else str(importe_max),
                "fechaDeInicio": fecha_inicio_solicitud.isoformat() if isinstance(fecha_inicio_solicitud, (datetime, date)) else (str(fecha_inicio_solicitud) if fecha_inicio_solicitud else ""),
                "fechaDeCierre": fecha_de_cierre.isoformat() if isinstance(fecha_de_cierre, (datetime, date)) else (str(fecha_de_cierre) if fecha_de_cierre else ""),
                "beneficiarios": parse_json_array(beneficiarios),
                "regiones": parse_json_array(regiones),
                "finalidades": parse_json_array(finalidades),
                "administraciones_convocantes": parse_json_array(administraciones_convocantes),
                "tipos_ayuda": parse_json_array(tipos_ayuda),
            }
        }), 200

    except Exception as e:
        logger.error(f"Error al obtener intereses refinados de matches: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            cursor.close()
            connection.close()
        except Exception:
            pass


@app.route('/api/matches/refine-interests', methods=['POST'])
@require_auth
def save_match_refine_interests(user_id):
    try:
        data = request.get_json() or {}

        def clean_list(value):
            if not isinstance(value, list):
                return []
            result = []
            for item in value:
                if item is None:
                    continue
                normalized = str(item).strip()
                if normalized and normalized not in result:
                    result.append(normalized)
            return result

        def clean_amount(value):
            if value in (None, ""):
                return None
            cleaned = str(value).strip()
            if not cleaned:
                return None
            return _safe_float(cleaned)

        interests_payload = {
            "importeMin": str(data.get("importeMin", "") or "").strip(),
            "importeMax": str(data.get("importeMax", "") or "").strip(),
            "fechaDeInicio": str(data.get("fechaDeInicio", "") or "").strip(),
            "fechaDeCierre": str(data.get("fechaDeCierre", "") or "").strip(),
            "beneficiarios": clean_list(data.get("beneficiarios")),
            "regiones": clean_list(data.get("regiones")),
            "finalidades": clean_list(data.get("finalidades")),
            "administraciones_convocantes": clean_list(data.get("administraciones_convocantes")),
            "tipos_ayuda": clean_list(data.get("tipos_ayuda")),
        }

        connection = get_connection()
        cursor = connection.cursor()
        has_admin_col = _table_has_column(cursor, "user_refine_interests", "administraciones_convocantes")
        has_tipo_col = _table_has_column(cursor, "user_refine_interests", "tipos_ayuda")
        has_fecha_inicio_col = _table_has_column(cursor, "user_refine_interests", "fecha_inicio_solicitud")
        has_fecha_cierre_col = _table_has_column(cursor, "user_refine_interests", "fecha_de_cierre")

        optional_missing = []
        if not has_admin_col:
            optional_missing.append("administraciones_convocantes")
        if not has_tipo_col:
            optional_missing.append("tipos_ayuda")
        if not has_fecha_inicio_col:
            optional_missing.append("fecha_inicio_solicitud")
        if not has_fecha_cierre_col:
            optional_missing.append("fecha_de_cierre")
        if optional_missing:
            logger.warning(
                "user_refine_interests sin columnas opcionales %s. Guardando refine interests sin esos campos.",
                ", ".join(optional_missing),
            )

        insert_columns = [
            "user_id",
            "importe_min",
            "importe_max",
            "beneficiarios",
            "regiones",
            "finalidades",
        ]
        insert_values_sql = ["%s", "%s", "%s", "%s::jsonb", "%s::jsonb", "%s::jsonb"]
        update_assignments = [
            "importe_min = EXCLUDED.importe_min",
            "importe_max = EXCLUDED.importe_max",
            "beneficiarios = EXCLUDED.beneficiarios",
            "regiones = EXCLUDED.regiones",
            "finalidades = EXCLUDED.finalidades",
        ]
        params = [
            user_id,
            clean_amount(interests_payload["importeMin"]),
            clean_amount(interests_payload["importeMax"]),
            json.dumps(interests_payload["beneficiarios"]),
            json.dumps(interests_payload["regiones"]),
            json.dumps(interests_payload["finalidades"]),
        ]

        if has_admin_col:
            insert_columns.append("administraciones_convocantes")
            insert_values_sql.append("%s::jsonb")
            update_assignments.append("administraciones_convocantes = EXCLUDED.administraciones_convocantes")
            params.append(json.dumps(interests_payload["administraciones_convocantes"]))

        if has_tipo_col:
            insert_columns.append("tipos_ayuda")
            insert_values_sql.append("%s::jsonb")
            update_assignments.append("tipos_ayuda = EXCLUDED.tipos_ayuda")
            params.append(json.dumps(interests_payload["tipos_ayuda"]))

        if has_fecha_inicio_col:
            insert_columns.append("fecha_inicio_solicitud")
            insert_values_sql.append("%s")
            update_assignments.append("fecha_inicio_solicitud = EXCLUDED.fecha_inicio_solicitud")
            params.append(_safe_parse_iso_date(interests_payload["fechaDeSolicitud"]))

        if has_fecha_cierre_col:
            insert_columns.append("fecha_de_cierre")
            insert_values_sql.append("%s")
            update_assignments.append("fecha_de_cierre = EXCLUDED.fecha_de_cierre")
            params.append(_safe_parse_iso_date(interests_payload["fechaDeCierre"]))

        cursor.execute(f"""
            INSERT INTO user_refine_interests (
                {", ".join(insert_columns)},
                is_active,
                created_at,
                updated_at
            )
            VALUES ({", ".join(insert_values_sql)}, TRUE, NOW(), NOW())
            ON CONFLICT (user_id) DO UPDATE SET
                {", ".join(update_assignments)},
                updated_at = NOW()
            RETURNING user_id
        """, tuple(params))

        updated_user = cursor.fetchone()
        connection.commit()

        if not updated_user:
            return jsonify({"error": "Usuario no encontrado"}), 404

        return jsonify({
            "status": "success",
            "message": "Preferencias refinadas guardadas correctamente",
            "interests": interests_payload
        }), 200

    except Exception as e:
        try:
            connection.rollback()
        except Exception:
            pass
        logger.error(f"Error al guardar intereses refinados de matches: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            cursor.close()
            connection.close()
        except Exception:
            pass


@app.route('/api/matches/export', methods=['POST'])
@require_auth
def export_matches_excel(user_id):
    """Exporta en Excel las subvenciones visibles en la pantalla de matches."""
    connection = None
    cursor = None
    try:
        data = request.get_json() or {}
        raw_grant_ids = data.get("grant_ids", [])

        if not isinstance(raw_grant_ids, list) or not raw_grant_ids:
            return jsonify({"error": "Se requiere una lista de grant_ids para exportar"}), 400

        # Normalizar IDs y preservar orden de pantalla
        grant_ids = []
        seen = set()
        for grant_id in raw_grant_ids:
            try:
                parsed = int(str(grant_id).strip())
            except Exception:
                continue
            if parsed not in seen:
                seen.add(parsed)
                grant_ids.append(parsed)

        if not grant_ids:
            return jsonify({"error": "No hay subvenciones válidas para exportar"}), 400

        connection = get_connection()
        cursor = connection.cursor()

        # Obtener la entidad seleccionada del usuario (para incluir datos de match)
        cursor.execute("""
            SELECT ue.entity_id
            FROM user_entities ue
            WHERE ue.user_id = %s AND ue.is_selected = TRUE
            ORDER BY ue.updated_at DESC NULLS LAST, ue.created_at DESC
            LIMIT 1
        """, (user_id,))
        entity_row = cursor.fetchone()

        if not entity_row:
            cursor.execute("""
                SELECT ue.entity_id
                FROM user_entities ue
                WHERE ue.user_id = %s
                ORDER BY ue.created_at DESC
                LIMIT 1
            """, (user_id,))
            entity_row = cursor.fetchone()

        entity_id = entity_row[0] if entity_row else -1

        def _find_grants_column_name_case_insensitive(candidate_names):
            try:
                lowered_candidates = [str(c).lower() for c in candidate_names if c]
                cursor.execute("""
                    SELECT column_name
                    FROM information_schema.columns
                    WHERE table_schema = 'public'
                      AND table_name = 'grants'
                """)
                for (column_name,) in cursor.fetchall():
                    if str(column_name).lower() in lowered_candidates:
                        return column_name
            except Exception as e:
                logger.warning(f"No se pudo inspeccionar columnas de grants para export: {e}")
            return None

        organo_column = _find_grants_column_name_case_insensitive([
            "organo_convocante",
            "organo",
            "organismo",
            "entidad_convocante",
            "organo_concedente",
        ])

        organo_select = f'g."{organo_column}"' if organo_column else "NULL::text"

        # Traer solo columnas legibles para exportación + datos de match
        cursor.execute(f"""
            SELECT
                g.id AS grant_id,
                g.titulo_corto AS titulo,
                g.presupuesto AS fondo_disponible,
                g.fecha_finalizacion AS plazo,
                g.finalidad AS finalidad,
                g.region_impacto AS region_impacto,
                {organo_select} AS organo_convocante,
                g.resumen_completo AS resumen_html,
                COALESCE((
                    SELECT string_agg(DISTINCT b.value, ', ')
                    FROM jsonb_array_elements_text(
                        CASE
                            WHEN jsonb_typeof(g."Beneficiarios_Short"->'categorias') = 'array'
                            THEN g."Beneficiarios_Short"->'categorias'
                            ELSE '[]'::jsonb
                        END
                    ) AS b(value)
                ), '') AS beneficiarios_texto,
                m.numero_match AS match_numero_match,
                m.justificacion AS match_justificacion_html,
                m.recomendacion AS match_recomendacion_html
            FROM grants g
            LEFT JOIN LATERAL (
                SELECT numero_match, justificacion, recomendacion
                FROM matches
                WHERE matches.grant_id = g.id
                  AND matches.entity_id = %s
                ORDER BY numero_match DESC
                LIMIT 1
            ) m ON TRUE
            WHERE g.id = ANY(%s)
              AND g.resumen_completo IS NOT NULL
        """, (entity_id, grant_ids))

        rows = cursor.fetchall()
        if not rows:
            return jsonify({"error": "No se encontraron subvenciones para exportar"}), 404

        column_names = [desc[0] for desc in cursor.description]
        order_map = {gid: index for index, gid in enumerate(grant_ids)}

        def serialize_value(value):
            if value is None:
                return ""
            if isinstance(value, (datetime, date)):
                return value.isoformat()
            if isinstance(value, (dict, list)):
                return json.dumps(value, ensure_ascii=False)
            if isinstance(value, bytes):
                try:
                    return value.decode("utf-8")
                except Exception:
                    return str(value)
            return str(value)

        records = []
        for row in rows:
            record = {}
            for idx, col_name in enumerate(column_names):
                record[col_name] = row[idx]
            records.append(record)

        records.sort(key=lambda item: order_map.get(int(item.get("grant_id", 0)), 10**9))

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            openpyxl_available = True
        except Exception as import_error:
            logger.warning(f"openpyxl no disponible; usando xlsxwriter para exportación XLSX: {import_error}")
            openpyxl_available = False
        try:
            import xlsxwriter  # type: ignore
            xlsxwriter_available = True
        except Exception as xlsxwriter_error:
            logger.warning(f"xlsxwriter no disponible: {xlsxwriter_error}")
            xlsxwriter_available = False

        filters_applied = data.get("filters_applied") or {}

        def _display_value(value):
            if value is None:
                return "No disponible"
            text = serialize_value(value).strip()
            return text if text else "No disponible"

        export_rows = []
        for record in records:
            resumen_html = record.get("resumen_html") or ""
            justificacion_html = record.get("match_justificacion_html") or ""
            recomendacion_html = record.get("match_recomendacion_html") or ""

            resumen_txt = html_to_text_for_export(normalize_summary_html(str(resumen_html))) if resumen_html else ""
            justificacion_txt = html_to_text_for_export(normalize_summary_html(str(justificacion_html))) if justificacion_html else ""
            recomendacion_txt = html_to_text_for_export(normalize_summary_html(str(recomendacion_html))) if recomendacion_html else ""

            export_rows.append({
                "ID subvención": _display_value(record.get("grant_id")),
                "Título": _display_value(record.get("titulo")),
                "Beneficiario": _display_value(record.get("beneficiarios_texto")),
                "Fondo disponible": _display_value(record.get("fondo_disponible")),
                "Plazo": _display_value(record.get("plazo")),
                "Finalidad": _display_value(record.get("finalidad")),
                "Órgano convocante": _display_value(record.get("organo_convocante")),
                "Región de impacto": _display_value(record.get("region_impacto")),
                "Compatibilidad (%)": _display_value(record.get("match_numero_match")),
                "Resumen": resumen_txt or "No disponible",
                "Justificación match": justificacion_txt or "No disponible",
                "Recomendación": recomendacion_txt or "No disponible",
            })

        headers = list(export_rows[0].keys()) if export_rows else [
            "ID subvención",
            "Título",
            "Beneficiario",
            "Fondo disponible",
            "Plazo",
            "Finalidad",
            "Órgano convocante",
            "Región de impacto",
            "Compatibilidad (%)",
            "Resumen",
            "Justificación match",
            "Recomendación",
        ]

        def _filters_amount_label(filters_dict):
            amount_range = filters_dict.get("amountRange") if isinstance(filters_dict, dict) else None
            if isinstance(amount_range, list) and len(amount_range) >= 2:
                try:
                    return f"{amount_range[0]} - {amount_range[1]}"
                except Exception:
                    return str(amount_range)
            return str(filters_dict.get("amountBand", "all")) if isinstance(filters_dict, dict) else "all"

        if openpyxl_available:
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Subvenciones compatibles"

            if isinstance(filters_applied, dict):
                meta_sheet = workbook.create_sheet("Filtros aplicados")
                meta_sheet.append(["Filtro", "Valor"])
                meta_sheet.append(["Importe", _filters_amount_label(filters_applied)])
                meta_sheet.append(["Beneficiarios", ", ".join(filters_applied.get("beneficiarios", [])) if isinstance(filters_applied.get("beneficiarios"), list) else ""])
                meta_sheet.append(["Regiones", ", ".join(filters_applied.get("regiones", [])) if isinstance(filters_applied.get("regiones"), list) else ""])
                meta_sheet.append(["Finalidades", ", ".join(filters_applied.get("finalidades", [])) if isinstance(filters_applied.get("finalidades"), list) else ""])

            sheet.append(headers)

            header_fill = PatternFill(fill_type="solid", fgColor="E8EEF9")
            header_font = Font(bold=True)

            for cell in sheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(vertical="top", wrap_text=True)

            for export_row in export_rows:
                row_values = [serialize_value(export_row.get(col)) for col in headers]
                sheet.append(row_values)

            # Ajuste básico de ancho de columnas (limitado para no romper el layout)
            for column_cells in sheet.columns:
                max_length = 0
                column_letter = column_cells[0].column_letter
                for cell in column_cells[:150]:
                    value = cell.value
                    if value is None:
                        continue
                    max_length = max(max_length, len(str(value)))
                sheet.column_dimensions[column_letter].width = min(max(max_length + 2, 14), 60)

            sheet.freeze_panes = "A2"
            sheet.auto_filter.ref = sheet.dimensions

            output = io.BytesIO()
            workbook.save(output)
            output.seek(0)

            filename = f"subvenciones_compatibles_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            return send_file(
                output,
                as_attachment=True,
                download_name=filename,
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        if xlsxwriter_available:
            output = io.BytesIO()
            workbook = xlsxwriter.Workbook(output, {'in_memory': True})
            sheet = workbook.add_worksheet("Subvenciones compatibles")

            header_fmt = workbook.add_format({
                'bold': True,
                'bg_color': '#E8EEF9',
                'text_wrap': True,
                'valign': 'top',
                'border': 1,
            })
            cell_fmt = workbook.add_format({
                'text_wrap': True,
                'valign': 'top',
                'border': 1,
            })

            if isinstance(filters_applied, dict):
                meta_sheet = workbook.add_worksheet("Filtros aplicados")
                meta_rows = [
                    ("Filtro", "Valor"),
                    ("Importe", _filters_amount_label(filters_applied)),
                    ("Beneficiarios", ", ".join(filters_applied.get("beneficiarios", [])) if isinstance(filters_applied.get("beneficiarios"), list) else ""),
                    ("Regiones", ", ".join(filters_applied.get("regiones", [])) if isinstance(filters_applied.get("regiones"), list) else ""),
                    ("Finalidades", ", ".join(filters_applied.get("finalidades", [])) if isinstance(filters_applied.get("finalidades"), list) else ""),
                ]
                for r_idx, (left, right) in enumerate(meta_rows):
                    fmt = header_fmt if r_idx == 0 else cell_fmt
                    meta_sheet.write(r_idx, 0, left, fmt)
                    meta_sheet.write(r_idx, 1, right, fmt)
                meta_sheet.set_column(0, 0, 22)
                meta_sheet.set_column(1, 1, 80)

            for col_idx, header in enumerate(headers):
                sheet.write(0, col_idx, header, header_fmt)

            for row_idx, export_row in enumerate(export_rows, start=1):
                for col_idx, header in enumerate(headers):
                    sheet.write(row_idx, col_idx, serialize_value(export_row.get(header)), cell_fmt)

            width_map = {
                "ID subvención": 14,
                "Título": 50,
                "Beneficiario": 34,
                "Fondo disponible": 18,
                "Plazo": 16,
                "Finalidad": 36,
                "Órgano convocante": 40,
                "Región de impacto": 26,
                "Compatibilidad (%)": 18,
                "Resumen": 70,
                "Justificación match": 70,
                "Recomendación": 70,
            }
            for col_idx, header in enumerate(headers):
                sheet.set_column(col_idx, col_idx, width_map.get(header, 24))

            sheet.freeze_panes(1, 0)
            workbook.close()
            output.seek(0)

            filename = f"subvenciones_compatibles_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            return send_file(
                output,
                as_attachment=True,
                download_name=filename,
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        return jsonify({"error": "No se pudo generar XLSX: faltan dependencias de Excel (openpyxl/xlsxwriter)"}), 500

    except Exception as e:
        logger.error(f"Error exportando matches a Excel: {e}", exc_info=True)
        return jsonify({"error": "Error al generar el Excel de subvenciones"}), 500
    finally:
        try:
            cursor.close()
            connection.close()
        except Exception:
            pass

""" USER """
# Endpoint para actualizar el perfil del usuario
@app.route('/api/update_user_profile', methods=['POST'])
def update_user_profile():
    """Endpoint para actualizar los datos del perfil del usuario"""
    try:
        data = request.json
        print(f"=== UPDATE USER PROFILE START ===")
        print(f"Received data: {data}")

        user_id = data.get('user_id')

        if not user_id:
            return jsonify(success=False, message="No se encontraron datos o user_id"), 400

        name = data.get('name')
        email = data.get('email')
        phone = data.get('phone')
        preferences = data.get('preferences')

        # Serializamos el dict a JSON
        prefs_json = json.dumps(preferences)
        print(prefs_json)

        connection = get_connection()
        cursor = connection.cursor()
        print("Connected to the database successfully.")

        # Verificar que el usuario existe
        cursor.execute("SELECT id FROM app_user WHERE id = %s", (user_id,))
        if not cursor.fetchone():
            return jsonify(success=False, message="Usuario no encontrado"), 404

        # Actualizar los datos del usuario
        update_query = """
            UPDATE app_user
            SET name = %s,
                email = %s,
                phone = %s,
                preferences = %s,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = %s
        """
        
        cursor.execute(update_query, (
            name,
            email,
            phone,
            prefs_json,
            user_id
        ))
        
        connection.commit()
        
        return jsonify(success=True, message="Perfil de usuario actualizado correctamente")
    
    except Exception as e:
        print("Error al actualizar el perfil del usuario:", e)
        return jsonify(success=False, message=str(e)), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()
            print("Database connection closed.")

@app.route('/api/update_user_preferences', methods=['POST'])
def update_user_preferences():
    try:
        data = request.get_json() or {}
        print("=== UPDATE USER PREFERENCES START ===", data)

        user_id = data.get('user_id')
        if not user_id:
            return jsonify(success=False, message="Faltan datos del usuario"), 400

        # Sólo nos quedamos con las claves válidas
        allowed = {'smsUpdates', 'emailUpdates', 'darkMode'}
        incoming = {k: data[k] for k in allowed if k in data}

        if not incoming:
            return jsonify(success=False, message="No se enviaron preferencias para actualizar"), 400

        # Serializamos sólo el trocito que cambió
        incoming_json = json.dumps(incoming)

        conn = get_connection()
        cur = conn.cursor()
        print("Connected to the database successfully.")

        # Verificamos que el usuario existe
        cur.execute("SELECT 1 FROM app_user WHERE id = %s", (user_id,))
        if not cur.fetchone():
            return jsonify(success=False, message="Usuario no encontrado"), 404

        # Aquí viene la magia: 
        #   preferences || '{"notifications":false}'::jsonb
        # concatena el JSON viejo con el nuevo, sobrescribiendo sólo esas claves.
        update_sql = """
            UPDATE app_user
            SET preferences = preferences || %s::jsonb,
                updated_at   = CURRENT_TIMESTAMP
            WHERE id = %s
        """
        cur.execute(update_sql, (incoming_json, user_id))
        conn.commit()

        return jsonify(success=True, message="Preferencias actualizadas correctamente")

    except Exception as e:
        print("Error al actualizar las preferencias:", e)
        return jsonify(success=False, message=str(e)), 500

    finally:
        try:
            cur.close()
            conn.close()
            print("Database connection closed.")
        except:
            pass

# Endpoint para eliminar completamente un usuario
@app.route('/api/delete_user_account', methods=['DELETE'])
def delete_user_account():
    """Endpoint para eliminar completamente la cuenta de un usuario"""
    try:
        data = request.json
        user_id = data.get('user_id')
        logger.info(f"=== DELETE USER ACCOUNT START ===")
        logger.info(f"Received user_id from auth: {user_id}")

        if not user_id:
            return jsonify(success=False, message="No se encontró el usuario"), 400
        
        connection = get_connection()
        cursor = connection.cursor()
        logger.info("Connected to the database successfully.")

        # Verificar que el usuario existe
        cursor.execute("SELECT id, email FROM app_user WHERE id = %s", (user_id,))
        user_data = cursor.fetchone()
        if not user_data:
            logger.warning(f"User not found: {user_id}")
            return jsonify(success=False, message="Usuario no encontrado"), 404

        try:
            # Obtener las entidades del usuario
            cursor.execute("SELECT entity_id FROM user_entities WHERE user_id = %s", (user_id,))
            entity_ids = [row[0] for row in cursor.fetchall()]
            
            # Eliminar matches de las entidades del usuario
            #for entity_id in entity_ids:
                #cursor.execute("DELETE FROM matches WHERE entity_id = %s", (entity_id,))
            
            # Eliminar preferencias del usuario
            cursor.execute("DELETE FROM user_grant_preferences WHERE user_id = %s", (user_id,))
            
            # Eliminar relaciones usuario-entidad
            cursor.execute("DELETE FROM user_entities WHERE user_id = %s", (user_id,))
            
            # Eliminar entidades que no tienen otros usuarios asociados
            #for entity_id in entity_ids:
                #cursor.execute("SELECT COUNT(*) FROM user_entities WHERE entity_id = %s", (entity_id,))
                #if cursor.fetchone()[0] == 0:
                    #cursor.execute("DELETE FROM entities WHERE id = %s", (entity_id,))
            
            # Eliminar tokens de autenticación del usuario
            cursor.execute("DELETE FROM refresh_token WHERE user_id = %s", (user_id,))
            
            # Finalmente, eliminar el usuario
            cursor.execute("DELETE FROM app_user WHERE id = %s", (user_id,))
            
            # Confirmar la transacción
            connection.commit()
            
            logger.info(f"Usuario {user_data[1]} (ID: {user_id}) eliminado completamente")
            
            return jsonify(success=True, message="Cuenta de usuario eliminada completamente")
            
        except Exception as e:
            # Revertir la transacción en caso de error
            connection.rollback()
            raise e
    
    except Exception as e:
        print("Error al eliminar la cuenta del usuario:", e)
        return jsonify(success=False, message=str(e)), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()
            print("Database connection closed.")

# Endpoint para obtener los datos completos del usuario
@app.route('/api/get_user_profile', methods=['GET'])
@require_auth
def get_user_profile(user_id):
    """Endpoint para obtener los datos completos del perfil del usuario"""
    try:
        print(f"=== GET USER PROFILE START ===")
        print(f"Received user_id from auth: {user_id}")
        if not user_id:
            return jsonify(success=False, message="No se encontró el usuario"), 400

        connection = get_connection()
        cursor = connection.cursor()
        print("Connected to the database successfully.")

        # Obtener datos del usuario
        query = """
            SELECT id, email, name, phone, created_at, preferences, email_verified
            FROM app_user
            WHERE id = %s
        """
        
        cursor.execute(query, (user_id,))
        user_row = cursor.fetchone()
        
        if not user_row:
            return jsonify(success=False, message="Usuario no encontrado"), 404
        
        # Formatear la fecha de creación
        created_at = user_row[4]
        if created_at and isinstance(created_at, datetime):
            created_at = created_at.strftime("%Y-%m-%d")
        
        # Parsear preferencias JSON
        print(f"User preferences JSON: {user_row[5]}")
        
        user_profile = {
            "id": user_row[0],
            "email": user_row[1],
            "name": user_row[2] if user_row[2] else "",
            "phone": user_row[3] if user_row[3] else "",
            "created_at": created_at,
            "preferences": user_row[5] if user_row[5] else {'smsUpdates': False, 'emailUpdates': False, 'darkMode': False},
            "email_verified": user_row[6] if user_row[6] is not None else False
        }
        
        print(f"User profile data: {user_profile}")
        return jsonify(success=True, profile=user_profile)
    
    except Exception as e:
        print("Error al obtener el perfil del usuario:", e)
        return jsonify(success=False, message=str(e)), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()
            print("Database connection closed.")

# Endpoint para cambiar la contraseña del usuario
@app.route('/api/change_password', methods=['POST'])
def change_password():
    # Creamos un singleton para reutilizar
    ph = PasswordHasher()

    data = request.get_json() or {}
    user_id         = data.get('user_id')
    current_pw      = data.get('current_password')
    new_pw          = data.get('new_password')

    if not all([user_id, current_pw, new_pw]):
        return jsonify(success=False, message="Rellena todos los campos"), 400

    conn = get_connection()
    cur  = conn.cursor()
    # 1) Recuperamos el hash actual de la BBDD
    cur.execute("SELECT password_hash FROM app_user WHERE id = %s", (user_id,))
    row = cur.fetchone()
    if not row:
        return jsonify(success=False, message="Usuario no encontrado"), 404

    stored_hash = row[0]

    # 2) Verificamos la contraseña actual
    try:
        ph.verify(stored_hash, current_pw)
    except argon2_exceptions.VerifyMismatchError:
        return jsonify(success=False, message="Contraseña actual incorrecta"), 401
    except argon2_exceptions.VerificationError as e:
        # Error interno de verificación
        return jsonify(success=False, message="Error al verificar contraseña"), 500

    # 3) Generamos el nuevo hash y actualizamos
    try:
        new_hash = ph.hash(new_pw)
        cur.execute("""
            UPDATE app_user
            SET password_hash = %s,
                updated_at    = CURRENT_TIMESTAMP
            WHERE id = %s
        """, (new_hash, user_id))
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify(success=False, message="Error al actualizar contraseña"), 500
    finally:
        cur.close()
        conn.close()

    return jsonify(success=True, message="Contraseña actualizada correctamente")

# Endpoint para solicitar un restablecimiento de contraseña
@app.route('/api/forgot_password', methods=['POST'])
def forgot_password():
    data = request.get_json() or {}
    email = data.get('email')
    if not email:
        return jsonify(success=False, message="El email es obligatorio"), 400

    conn = get_connection()
    cur = conn.cursor()
    # 1) ¿Existe el usuario?
    cur.execute("SELECT id FROM app_user WHERE email = %s", (email,))
    row = cur.fetchone()
    if not row:
        # Respondemos OK de todas formas para no revelar usuarios
        cur.close()
        conn.close()
        return jsonify(success=True, message="Si el email existe, recibirás instrucciones"), 200

    user_id = row[0]
    # 2) Generamos un token (UUID) y lo guardamos con TTL
    token = str(uuid.uuid4())
    cur.execute("""
        INSERT INTO password_reset_tokens (user_id, token, expires_at)
        VALUES (%s, %s, NOW() + INTERVAL '1 hour')
        ON CONFLICT (user_id) DO UPDATE
          SET token = EXCLUDED.token,
              expires_at = EXCLUDED.expires_at
    """, (user_id, token))
    conn.commit()
    cur.close()
    conn.close()

    # 3) Enviamos el correo con el enlace
    # reset_link = f"https://tu-dominio.com/reset-password?token={token}"
    # send_reset_email(email, reset_link)

    return jsonify(success=True, message="Si el email existe, recibirás instrucciones"), 200

# Endpoint para completar el onboarding del usuario
@app.route('/api/auth/complete-onboarding', methods=['POST'])
def complete_onboarding():
    """Endpoint para marcar el onboarding como completado"""
    try:
        # Obtener el token de autorización
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"error": "Token de autorización requerido"}), 401
        
        token = auth_header.split(' ')[1]
        
        # Importar función de validación
        from utils.auth import validate_token
        
        # Validar token y obtener user_id
        user_id, error = validate_token(token)
        if error:
            return jsonify({"error": error}), 401
        
        # Conectar a la base de datos
        connection = get_connection()
        cursor = connection.cursor()
        
        # Actualizar el estado de onboarding
        query = """
            UPDATE app_user 
            SET onboarding_completed = TRUE 
            WHERE id = %s
            RETURNING id, email, email_verified, onboarding_completed
        """
        
        cursor.execute(query, (user_id,))
        result = cursor.fetchone()
        connection.commit()
        
        if not result:
            return jsonify({"error": "Usuario no encontrado"}), 404
        
        return jsonify({
            "message": "Onboarding completado exitosamente",
            "user": {
                "id": str(result[0]),
                "email": result[1],
                "emailVerified": result[2],
                "onboardingCompleted": result[3]
            }
        }), 200
    
    except Exception as e:
        print("Error completing onboarding:", e)
        return jsonify({"error": f"Error al completar onboarding: {str(e)}"}), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()
            print("Database connection closed.")

""" GRANT SEARCH & FILTERS"""
# Endpoint para búsqueda en el marketplace con filtros y paginación
@app.route('/api/search_grants_marketplace', methods=['POST'])
@require_auth
def search_grants_marketplace(user_id):
    data = {}
    page = 1
    limit = 20
    try:
        data = request.get_json(silent=True) or {}
        
        # Extraer filtros del objeto filters
        filters = data.get('filters', {})
        beneficiarios = filters.get('beneficiarios', [])
        regiones = filters.get('regiones', [])
        finalidades = filters.get('finalidades', [])
        administraciones_convocantes = filters.get('administraciones_convocantes', [])
        tipos_ayuda = filters.get('tipos_ayuda', [])
        amount_range = data.get('amount_range', None)
        date_range = data.get('date_range', {}) if isinstance(data, dict) else {}
        amount_min = None
        amount_max = None
        fecha_inicio = None
        fecha_cierre = None
        if isinstance(amount_range, (list, tuple)) and len(amount_range) >= 2:
            try:
                amount_min = max(0.0, float(amount_range[0]))
                amount_max = max(0.0, float(amount_range[1]))
                if amount_min > amount_max:
                    amount_min, amount_max = amount_max, amount_min
            except (TypeError, ValueError):
                amount_min = None
                amount_max = None
        if isinstance(date_range, dict):
            raw_fecha_inicio = str(date_range.get('fecha_inicio_solicitud') or '').strip()
            raw_fecha_cierre = str(date_range.get('fecha_de_cierre') or '').strip()
            parsed_fecha_inicio = _safe_parse_iso_date(raw_fecha_inicio)
            parsed_fecha_cierre = _safe_parse_iso_date(raw_fecha_cierre)
            if parsed_fecha_inicio:
                fecha_inicio = parsed_fecha_inicio.isoformat()
            if parsed_fecha_cierre:
                fecha_cierre = parsed_fecha_cierre.isoformat()
            if fecha_inicio and fecha_cierre and fecha_inicio > fecha_cierre:
                fecha_inicio, fecha_cierre = fecha_cierre, fecha_inicio
        
        search_query = data.get('search_query')
        page = data.get('page', 1)
        limit = data.get('limit', 20)
        try:
            page = max(1, int(page))
        except (TypeError, ValueError):
            page = 1
        try:
            limit = max(1, min(int(limit), 100))
        except (TypeError, ValueError):
            limit = 20
        order_by = data.get('order_by', 'preferences')
        if order_by not in {'preferences', 'match', 'amount', 'deadline'}:
            order_by = 'preferences'
        sort_direction = data.get('sort_direction', 'desc')
        if sort_direction not in {'asc', 'desc'}:
            sort_direction = 'desc'

        logger.info(
            f"MARKETPLACE search - User: {user_id}, Filters: beneficiarios={beneficiarios}, "
            f"regiones={regiones}, finalidades={finalidades}, "
            f"administraciones_convocantes={administraciones_convocantes}, tipos_ayuda={tipos_ayuda}, "
            f"amount_range={[amount_min, amount_max] if amount_min is not None and amount_max is not None else None}, "
            f"date_range={[fecha_inicio, fecha_cierre] if fecha_inicio or fecha_cierre else None}, "
            f"query={search_query}, order_by={order_by}, sort_direction={sort_direction}"
        )

        result = GrantSearch.search_marketplace(
            user_id=user_id,
            beneficiarios=beneficiarios if beneficiarios else None,
            regiones=regiones if regiones else None,
            finalidades=finalidades if finalidades else None,
            administraciones_convocantes=administraciones_convocantes if administraciones_convocantes else None,
            tipos_ayuda=tipos_ayuda if tipos_ayuda else None,
            amount_min=amount_min,
            amount_max=amount_max,
            fecha_inicio=fecha_inicio,
            fecha_cierre=fecha_cierre,
            order_by=order_by,
            sort_direction=sort_direction,
            search_query=search_query,
            page=page,
            limit=limit
        )
        
        # Añadir total_pages al resultado
        total_pages = (result['total_count'] + limit - 1) // limit if result.get('total_count', 0) > 0 else 0
        result['total_pages'] = total_pages
        
        return jsonify(result)

    except Exception as e:
        logger.exception("Error in search_grants_marketplace")
        # Fallback payload to avoid hard-failing UI on transient backend issues.
        return jsonify({
            "grants": [],
            "has_more": False,
            "total_count": 0,
            "page": page,
            "total_pages": 0,
            "error": str(e),
        }), 200

# ─── Barometro de subvenciones (público) ────────────────────────────────
_barometro_cache = {"data": None, "timestamp": None}

@app.route('/api/barometro', methods=['GET'])
def get_barometro_data():
    """Endpoint público: estadísticas agregadas de subvenciones para el Barómetro."""
    import time as _time

    # Cache de 1 hora
    now = _time.time()
    if _barometro_cache["data"] and _barometro_cache["timestamp"] and (now - _barometro_cache["timestamp"] < 3600):
        return jsonify(_barometro_cache["data"])

    connection = None
    try:
        connection = get_connection()
        cursor = connection.cursor()

        base_where = """
            resumen_completo IS NOT NULL
            AND titulo_corto IS NOT NULL AND titulo_corto <> ''
            AND "Beneficiarios_Short" IS NOT NULL
            AND "Beneficiarios_Short"::text NOT IN ('{}', '{"categorias": []}', '')
            AND region_impacto IS NOT NULL AND region_impacto <> ''
        """

        # KPIs globales
        cursor.execute(f"""
            SELECT
              COUNT(*),
              COUNT(*) FILTER (WHERE fecha_de_publicacion >= CURRENT_DATE - INTERVAL '30 days'),
              COUNT(*) FILTER (WHERE fecha_de_publicacion >= CURRENT_DATE - INTERVAL '7 days')
            FROM grants WHERE {base_where}
        """)
        row = cursor.fetchone()
        kpis = {
            "total_grants": row[0] or 0,
            "grants_last_30_days": row[1] or 0,
            "grants_last_7_days": row[2] or 0,
        }

        # Por finalidad (top 10)
        cursor.execute(f"""
            SELECT finalidad, COUNT(*) AS cnt
            FROM grants
            WHERE finalidad IS NOT NULL AND finalidad <> ''
              AND {base_where}
            GROUP BY finalidad
            ORDER BY cnt DESC
            LIMIT 10
        """)
        by_finalidad = [{"name": r[0], "count": r[1]} for r in cursor.fetchall()]

        # Por región (top 17)
        cursor.execute(f"""
            SELECT region_impacto, COUNT(*) AS cnt
            FROM grants
            WHERE {base_where}
            GROUP BY region_impacto
            ORDER BY cnt DESC
            LIMIT 17
        """)
        by_region = [{"name": r[0], "count": r[1]} for r in cursor.fetchall()]

        # Por tipo de beneficiario
        cursor.execute(f"""
            SELECT b.value AS beneficiario, COUNT(DISTINCT g.id) AS cnt
            FROM grants g
            LEFT JOIN LATERAL jsonb_array_elements_text(
              CASE
                WHEN jsonb_typeof(g."Beneficiarios_Short"->'categorias') = 'array'
                  THEN g."Beneficiarios_Short"->'categorias'
                ELSE '[]'::jsonb
              END
            ) AS b(value) ON TRUE
            WHERE {base_where}
              AND b.value IS NOT NULL AND TRIM(b.value) <> ''
            GROUP BY b.value
            ORDER BY cnt DESC
        """)
        by_beneficiario = [{"name": r[0], "count": r[1]} for r in cursor.fetchall()]

        # Tendencia mensual (últimos 12 meses)
        cursor.execute(f"""
            SELECT
              DATE_TRUNC('month', fecha_de_publicacion) AS month,
              COUNT(*) AS cnt
            FROM grants
            WHERE fecha_de_publicacion IS NOT NULL
              AND fecha_de_publicacion >= CURRENT_DATE - INTERVAL '12 months'
              AND {base_where}
            GROUP BY month
            ORDER BY month ASC
        """)
        monthly_trend = [{"month": r[0].strftime("%Y-%m"), "count": r[1]} for r in cursor.fetchall()]

        # Por sector (top 10)
        cursor.execute(f"""
            SELECT sector, COUNT(*) AS cnt
            FROM grants
            WHERE sector IS NOT NULL AND sector <> ''
              AND {base_where}
            GROUP BY sector
            ORDER BY cnt DESC
            LIMIT 10
        """)
        by_sector = [{"name": r[0], "count": r[1]} for r in cursor.fetchall()]

        result = {
            "kpis": kpis,
            "by_finalidad": by_finalidad,
            "by_region": by_region,
            "by_beneficiario": by_beneficiario,
            "monthly_trend": monthly_trend,
            "by_sector": by_sector,
            "generated_at": datetime.utcnow().isoformat(),
        }

        _barometro_cache["data"] = result
        _barometro_cache["timestamp"] = now

        return jsonify(result)

    except Exception as e:
        logger.exception("Error in get_barometro_data")
        return jsonify({"error": str(e)}), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


# Endpoint para obtener las opciones de filtros disponibles
@app.route('/api/get_filter_options', methods=['GET'])
def get_filter_options():
    """Endpoint para obtener las opciones únicas de cada filtro desde la base de datos"""
    try:
        logger.info("Fetching filter options for marketplace")
        options = GrantSearch.get_filter_options()
        logger.info(
            "get_filter_options response counts: beneficiarios=%s, regiones=%s, finalidades=%s, administraciones_convocantes=%s, tipos_ayuda=%s",
            len(options.get("beneficiarios", []) or []),
            len(options.get("regiones", []) or []),
            len(options.get("finalidades", []) or []),
            len(options.get("administraciones_convocantes", []) or []),
            len(options.get("tipos_ayuda", []) or []),
        )
        return jsonify(options)

    except Exception as e:
        logger.error("Error in get_filter_options:", e)
        return jsonify({"error": str(e)}), 500

# Endpoint para obtener sugerencias de títulos en el buscador
@app.route('/api/search_grants_suggestions', methods=['POST'])
@require_auth
def search_grants_suggestions(user_id):
    """Endpoint para obtener sugerencias de títulos basadas en el texto escrito"""
    try:
        data = request.get_json()
        search_query = data.get('search_query', '').strip()
        try:
            limit = int(data.get('limit', 10))
        except (TypeError, ValueError):
            limit = 10
        limit = max(1, min(limit, 20))  # Limitar a 20 sugerencias

        if not search_query or len(search_query) < 2:
            return jsonify({"suggestions": []})

        logger.info(f"Fetching title suggestions for query: '{search_query}'")

        conn = get_connection()
        cursor = conn.cursor()

        normalized_query = re.sub(r"[\"'`´‘’“”«»]", " ", search_query.lower())
        normalized_query = re.sub(r"[^0-9a-záéíóúüñ]+", " ", normalized_query)
        normalized_query = re.sub(r"\s+", " ", normalized_query).strip()
        stopwords = {
            "de", "del", "la", "las", "el", "los", "y", "e", "o", "u",
            "para", "por", "con", "sin", "en", "un", "una", "unos", "unas",
        }
        tokens = [
            t
            for t in normalized_query.split(" ")
            if t
            and t not in stopwords
            and (len(t) >= 3 or (len(t) >= 2 and any(ch.isdigit() for ch in t)))
        ]

        normalized_title_expr = (
            "LOWER(regexp_replace(COALESCE(titulo_corto, '') || ' ' || COALESCE(titulo, ''), "
            "'[^[:alnum:]áéíóúüñ]+', ' ', 'g'))"
        )

        where_parts = [
            "(COALESCE(titulo_corto, '') ILIKE %s OR COALESCE(titulo, '') ILIKE %s)"
        ]
        query_params = [f"%{search_query}%", f"%{search_query}%"]

        if normalized_query:
            where_parts.append(f"{normalized_title_expr} LIKE %s")
            query_params.append(f"%{normalized_query}%")

        if tokens:
            token_conditions = " OR ".join([f"{normalized_title_expr} LIKE %s"] * len(tokens))
            where_parts.append(f"({token_conditions})")
            query_params.extend([f"%{token}%" for token in tokens])

        query = f"""
            SELECT
                id,
                COALESCE(NULLIF(titulo_corto, ''), titulo) AS titulo_sugerencia
            FROM grants
            WHERE (
                {" OR ".join(where_parts)}
            )
            AND COALESCE(NULLIF(titulo_corto, ''), titulo) IS NOT NULL
            AND COALESCE(NULLIF(titulo_corto, ''), titulo) <> ''
            ORDER BY
                CASE
                    WHEN COALESCE(titulo_corto, '') ILIKE %s THEN 0
                    WHEN COALESCE(titulo, '') ILIKE %s THEN 1
                    WHEN COALESCE(titulo_corto, '') ILIKE %s THEN 2
                    WHEN COALESCE(titulo, '') ILIKE %s THEN 3
                    WHEN {normalized_title_expr} LIKE %s THEN 4
                    ELSE 5
                END,
                COALESCE(NULLIF(titulo_corto, ''), titulo)
            LIMIT %s
        """

        query_params.extend(
            [
                search_query,
                search_query,
                f"{search_query}%",
                f"{search_query}%",
                f"%{normalized_query or search_query.lower()}%",
                limit,
            ]
        )

        cursor.execute(query, tuple(query_params))
        rows = cursor.fetchall()

        suggestions = []
        for row in rows:
            suggestions.append({
                "id": str(row[0]),
                "titulo_corto": row[1]
            })

        cursor.close()
        conn.close()

        logger.info(f"Found {len(suggestions)} suggestions")
        return jsonify({"suggestions": suggestions})

    except Exception as e:
        logger.error(f"Error in search_grants_suggestions: {e}")
        return jsonify({"suggestions": [], "error": str(e)}), 500

@app.route('/webhook/n8n', methods=['POST'])
def webhook_n8n():
    try:
        run_main()
        return jsonify({"status": "Proceso ejecutado correctamente"}), 200
    except Exception as e:
        print("❌ Error ejecutando run_main():", e)
        return jsonify({"error": str(e)}), 500


@app.route('/api/entities/download_technical_sheet_pdf', methods=['POST'])
@require_auth
def download_entity_technical_sheet_pdf(user_id):
    """Genera una ficha técnica PDF de la entidad activa y la descarga directa."""
    connection = None
    cursor = None
    try:
        payload = request.get_json(silent=True) or {}
        entity_id = payload.get("entity_id")
        if entity_id in (None, ""):
            return jsonify({"error": "Falta entity_id"}), 400

        connection = get_connection()
        cursor = connection.cursor()
        cursor.execute(
            """
            SELECT e.id, COALESCE(e.razon_social, ''), COALESCE(e.nif, '')
            FROM user_entities ue
            JOIN entities e ON e.id = ue.entity_id
            WHERE ue.user_id = %s
              AND ue.entity_id = %s
            LIMIT 1
            """,
            (user_id, entity_id),
        )
        entity_row = cursor.fetchone()
        if not entity_row:
            return jsonify({"error": "Entidad no encontrada para este usuario"}), 404

        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.lib.utils import ImageReader
            from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except Exception as e:
            python_executable = sys.executable
            logger.error(
                "reportlab no disponible para generar PDF de entidad. python=%s error=%s",
                python_executable,
                e,
            )
            return jsonify({
                "error": f"No se pudo generar el PDF (reportlab import failed en {python_executable}: {type(e).__name__}: {e})"
            }), 500

        def _safe_text(value, fallback="-"):
            if value is None:
                return fallback
            text = str(value).strip()
            return text if text else fallback

        def _normalize_rows(raw_rows, expected_keys):
            normalized = []
            if not isinstance(raw_rows, list):
                return normalized
            for item in raw_rows:
                if not isinstance(item, dict):
                    continue
                normalized.append({key: _safe_text(item.get(key)) for key in expected_keys})
            return normalized

        entity_name = _safe_text(payload.get("entity_name") or entity_row[1], "Entidad")
        entity_nif = _safe_text(payload.get("entity_nif") or entity_row[2], "No disponible")

        generated_at_raw = payload.get("generated_at")
        generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
        if isinstance(generated_at_raw, str) and generated_at_raw.strip():
            try:
                generated_at = datetime.fromisoformat(generated_at_raw.replace("Z", "+00:00")).strftime("%d/%m/%Y %H:%M")
            except Exception:
                generated_at = _safe_text(generated_at_raw, generated_at)

        chart_filters = payload.get("chart_filters") if isinstance(payload.get("chart_filters"), dict) else {}
        chart_summary = payload.get("chart_summary") if isinstance(payload.get("chart_summary"), dict) else {}

        profile_rows = _normalize_rows(payload.get("profile_rows"), ("label", "value"))
        awarded_rows = _normalize_rows(payload.get("awarded_rows"), ("titulo", "organo", "fecha", "importe", "enlace"))
        minimis_rows = _normalize_rows(payload.get("minimis_rows"), ("titulo", "convocante", "fecha", "importe", "enlace_bdns"))

        chart_image_bytes = None
        chart_data_url = payload.get("chart_image_data_url")
        if isinstance(chart_data_url, str) and chart_data_url.startswith("data:image"):
            try:
                _, encoded_data = chart_data_url.split(",", 1)
                chart_image_bytes = base64.b64decode(encoded_data, validate=True)
            except Exception as decode_error:
                logger.warning(f"No se pudo decodificar chart_image_data_url en PDF de entidad: {decode_error}")

        output = io.BytesIO()
        doc = SimpleDocTemplate(
            output,
            pagesize=A4,
            leftMargin=16 * mm,
            rightMargin=16 * mm,
            topMargin=14 * mm,
            bottomMargin=14 * mm,
            title=f"Ficha técnica - {entity_name}",
        )

        primary = colors.HexColor("#4F46E5")
        primary_soft = colors.HexColor("#EEF2FF")
        text_main = colors.HexColor("#111827")
        text_muted = colors.HexColor("#6B7280")
        border = colors.HexColor("#D1D5DB")
        white = colors.HexColor("#FFFFFF")

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "EntityPdfTitle",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=19,
            leading=22,
            textColor=white,
            spaceAfter=2,
        )
        subtitle_style = ParagraphStyle(
            "EntityPdfSubtitle",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=white,
        )
        meta_style = ParagraphStyle(
            "EntityPdfMeta",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#E0E7FF"),
        )
        section_style = ParagraphStyle(
            "EntityPdfSection",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=14,
            textColor=primary,
        )
        cell_style = ParagraphStyle(
            "EntityPdfCell",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            textColor=text_main,
        )
        header_cell_style = ParagraphStyle(
            "EntityPdfHeaderCell",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=white,
        )
        kpi_label_style = ParagraphStyle(
            "EntityPdfKpiLabel",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8,
            leading=10,
            textColor=colors.black,
        )
        kpi_value_style = ParagraphStyle(
            "EntityPdfKpiValue",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=13,
            textColor=text_main,
        )

        def _p(value, style):
            escaped = pyhtml.escape(_safe_text(value), quote=False).replace("\n", "<br/>")
            return Paragraph(escaped, style)

        def _section_title(text):
            section_table = Table([[_p(text, section_style)]], colWidths=[doc.width])
            section_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), primary_soft),
                        ("BOX", (0, 0), (-1, -1), 0.8, border),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            return section_table

        story = []

        header_table = Table(
            [
                [_p("Ficha técnica de entidad", title_style)],
                [_p(entity_name, subtitle_style)],
                [_p(f"CIF/NIF: {entity_nif} · Generado: {generated_at}", meta_style)],
            ],
            colWidths=[doc.width],
        )
        header_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), primary),
                    ("BOX", (0, 0), (-1, -1), 0.8, primary),
                    ("LEFTPADDING", (0, 0), (-1, -1), 12),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        story.append(header_table)
        story.append(Spacer(1, 10))

        story.append(_section_title("Información de la empresa"))
        story.append(Spacer(1, 6))
        if not profile_rows:
            profile_rows = [{"label": "Información", "value": "No disponible"}]
        profile_data = [[_p("Campo", header_cell_style), _p("Valor", header_cell_style)]]
        profile_data.extend([[_p(item["label"], cell_style), _p(item["value"], cell_style)] for item in profile_rows])
        profile_table = Table(profile_data, colWidths=[doc.width * 0.33, doc.width * 0.67], repeatRows=1)
        profile_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), primary),
                    ("BOX", (0, 0), (-1, -1), 0.7, border),
                    ("INNERGRID", (0, 0), (-1, -1), 0.35, border),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(profile_table)
        story.append(Spacer(1, 10))

        story.append(_section_title("Resumen económico"))
        story.append(Spacer(1, 6))
        kpi_table = Table(
            [
                [
                    _p("TOTAL NO MINIMIS", kpi_label_style),
                    _p("TOTAL MINIMIS", kpi_label_style),
                    _p("TOTAL COMBINADO", kpi_label_style),
                ],
                [
                    _p(chart_summary.get("total_non_minimis", "No disponible"), kpi_value_style),
                    _p(chart_summary.get("total_minimis", "No disponible"), kpi_value_style),
                    _p(chart_summary.get("total_combined", "No disponible"), kpi_value_style),
                ],
            ],
            colWidths=[doc.width / 3.0, doc.width / 3.0, doc.width / 3.0],
        )
        kpi_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#A78BFA")),
                    ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#3966DF")),
                    ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#F3F4F6")),
                    ("TEXTCOLOR", (0, 0), (1, 0), colors.white),
                    ("BOX", (0, 0), (-1, -1), 0.8, border),
                    ("INNERGRID", (0, 0), (-1, -1), 0.35, border),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        story.append(kpi_table)
        filters_text = f"Filtros aplicados: Año={_safe_text(chart_filters.get('year'), 'Todos')} · Subvención={_safe_text(chart_filters.get('grant'), 'Todas')}"
        story.append(Spacer(1, 4))
        story.append(_p(filters_text, ParagraphStyle("EntityPdfFilter", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=text_muted)))
        story.append(Spacer(1, 10))

        story.append(_section_title("Evolución anual de subvenciones"))
        story.append(Spacer(1, 4))
        story.append(_p("Diferenciación entre minimis y no minimis, con filtros por año y subvención.", ParagraphStyle("EntityPdfChartSubtitle", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=text_muted)))
        story.append(Spacer(1, 6))
        if chart_image_bytes:
            image_reader = ImageReader(io.BytesIO(chart_image_bytes))
            image_width, image_height = image_reader.getSize()
            max_width = doc.width
            max_height = 110 * mm
            scale = min(max_width / float(image_width), max_height / float(image_height))
            chart_image = Image(io.BytesIO(chart_image_bytes), width=image_width * scale, height=image_height * scale)
            chart_container = Table([[chart_image]], colWidths=[doc.width])
            chart_container.setStyle(
                TableStyle(
                    [
                        ("BOX", (0, 0), (-1, -1), 0.8, border),
                        ("BACKGROUND", (0, 0), (-1, -1), white),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ]
                )
            )
            story.append(chart_container)
        else:
            story.append(_p("No se pudo incrustar el gráfico en esta descarga.", ParagraphStyle("EntityPdfChartWarning", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=12, textColor=text_muted)))
        story.append(Spacer(1, 10))

        story.append(_section_title("Subvenciones conseguidas"))
        story.append(Spacer(1, 6))
        if not awarded_rows:
            awarded_rows = [{"titulo": "No hay subvenciones concedidas", "organo": "-", "fecha": "-", "importe": "-", "enlace": "-"}]
        awarded_data = [[
            _p("Título", header_cell_style),
            _p("Órgano", header_cell_style),
            _p("Fecha", header_cell_style),
            _p("Importe", header_cell_style),
            _p("Enlace", header_cell_style),
        ]]
        for item in awarded_rows:
            awarded_data.append([
                _p(item["titulo"], cell_style),
                _p(item["organo"], cell_style),
                _p(item["fecha"], cell_style),
                _p(item["importe"], cell_style),
                _p(item["enlace"], cell_style),
            ])
        awarded_table = Table(
            awarded_data,
            colWidths=[doc.width * 0.34, doc.width * 0.20, doc.width * 0.11, doc.width * 0.12, doc.width * 0.23],
            repeatRows=1,
        )
        awarded_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), primary),
                    ("BOX", (0, 0), (-1, -1), 0.7, border),
                    ("INNERGRID", (0, 0), (-1, -1), 0.35, border),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(awarded_table)
        story.append(Spacer(1, 10))

        story.append(_section_title("Subvenciones minimis"))
        story.append(Spacer(1, 6))
        if not minimis_rows:
            minimis_rows = [{"titulo": "No hay subvenciones minimis", "convocante": "-", "fecha": "-", "importe": "-", "enlace_bdns": "-"}]
        minimis_data = [[
            _p("Título", header_cell_style),
            _p("Convocante", header_cell_style),
            _p("Fecha", header_cell_style),
            _p("Importe", header_cell_style),
            _p("Ver en BDNS", header_cell_style),
        ]]
        for item in minimis_rows:
            minimis_data.append([
                _p(item["titulo"], cell_style),
                _p(item["convocante"], cell_style),
                _p(item["fecha"], cell_style),
                _p(item["importe"], cell_style),
                _p(item["enlace_bdns"], cell_style),
            ])
        minimis_table = Table(
            minimis_data,
            colWidths=[doc.width * 0.34, doc.width * 0.20, doc.width * 0.11, doc.width * 0.12, doc.width * 0.23],
            repeatRows=1,
        )
        minimis_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), primary),
                    ("BOX", (0, 0), (-1, -1), 0.7, border),
                    ("INNERGRID", (0, 0), (-1, -1), 0.35, border),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(minimis_table)

        doc.build(story)
        output.seek(0)

        safe_name = re.sub(r'[^A-Za-z0-9_-]+', '_', entity_name)[:70].strip("_")
        filename = f"{safe_name or 'ficha_entidad'}.pdf"
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype="application/pdf",
        )
    except Exception as e:
        logger.error(f"Error generando ficha técnica PDF para entidad {request.json.get('entity_id') if request.is_json else 'unknown'}: {e}", exc_info=True)
        return jsonify({"error": "No se pudo generar la ficha técnica en PDF"}), 500
    finally:
        try:
            cursor.close()
            connection.close()
        except Exception:
            pass


@app.route('/api/grants/<int:grant_id>/download-pack', methods=['POST'])
@require_auth
def download_grant_pack(user_id, grant_id):
    """
    Genera un documento Word con resumen, justificación y documentación de la subvención.
    """
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute("""
            SELECT
                g.id,
                g.titulo_corto,
                g.presupuesto,
                g.fecha_finalizacion,
                g.resumen_completo,
                g.region_impacto,
                string_agg(DISTINCT b.value, ', ') AS beneficiarios,
                m.numero_match,
                m.justificacion
            FROM grants g
            LEFT JOIN LATERAL (
                SELECT value
                FROM jsonb_array_elements_text(
                    CASE
                        WHEN jsonb_typeof(g."Beneficiarios_Short"->'categorias') = 'array'
                        THEN g."Beneficiarios_Short"->'categorias'
                        ELSE '[]'::jsonb
                    END
                )
            ) AS b(value) ON TRUE
            LEFT JOIN LATERAL (
                SELECT ue.entity_id
                FROM user_entities ue
                WHERE ue.user_id = %s
                ORDER BY ue.is_selected DESC, ue.updated_at DESC NULLS LAST, ue.created_at DESC
                LIMIT 1
            ) ue_sel ON TRUE
            LEFT JOIN LATERAL (
                SELECT numero_match, justificacion
                FROM matches
                WHERE grant_id = g.id AND entity_id = ue_sel.entity_id
                ORDER BY numero_match DESC
                LIMIT 1
            ) m ON TRUE
            WHERE g.id = %s
              AND g.resumen_completo IS NOT NULL
            GROUP BY g.id, g.titulo_corto, g.presupuesto, g.fecha_finalizacion, g.resumen_completo, g.region_impacto, m.numero_match, m.justificacion
            LIMIT 1
        """, (user_id, grant_id))
        row = cursor.fetchone()

        if not row:
            return jsonify({"error": "Subvención no encontrada"}), 404

        extra = get_grant_extra_metadata(cursor, grant_id)
        documentacion_items = parse_documentacion_items(extra.get("documentacion"))

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import RGBColor
        except Exception as e:
            logger.error(f"python-docx no disponible: {e}")
            return jsonify({"error": "No se pudo generar el documento Word (falta python-docx)"}), 500

        brand_lila = RGBColor(124, 58, 237)  # violeta principal

        def add_lila_heading(document, text, level=1):
            paragraph = document.add_heading(text, level=level)
            for run in paragraph.runs:
                run.font.color.rgb = brand_lila
            return paragraph

        def add_justified_paragraph(document, text):
            paragraph = document.add_paragraph(text)
            return paragraph

        class _DocxHtmlFragmentParser(HTMLParser):
            """
            Parser HTML básico para volcar contenido de resumen/justificación a docx
            preservando títulos, negritas y listas.
            """
            def __init__(self, document):
                super().__init__(convert_charrefs=False)
                self.document = document
                self.current_paragraph = None
                self.bold_depth = 0
                self.italic_depth = 0
                self.list_stack = []
                self.in_heading_level = None
                self.has_written_content = False

            def _new_paragraph(self, style=None):
                self.current_paragraph = self.document.add_paragraph(style=style) if style else self.document.add_paragraph()
                return self.current_paragraph

            def _new_heading(self, level):
                # Mantener jerarquía visual sin competir con el título principal del documento
                word_level = 3 if level <= 2 else min(level + 1, 4)
                self.current_paragraph = self.document.add_heading("", level=word_level)
                return self.current_paragraph

            def _ensure_paragraph(self):
                if self.current_paragraph is None:
                    self._new_paragraph()
                return self.current_paragraph

            def _add_text(self, text):
                if text is None:
                    return
                decoded = pyhtml.unescape(text)
                if not decoded:
                    return
                if not decoded.strip() and "\n" in decoded:
                    return

                paragraph = self._ensure_paragraph()
                run = paragraph.add_run(decoded)
                if self.bold_depth > 0:
                    run.bold = True
                if self.italic_depth > 0:
                    run.italic = True
                if self.in_heading_level is not None:
                    run.font.color.rgb = brand_lila
                if decoded.strip():
                    self.has_written_content = True

            def handle_starttag(self, tag, attrs):
                tag = (tag or "").lower()
                if tag in ("p", "div"):
                    self._new_paragraph()
                elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    self.in_heading_level = int(tag[1])
                    self._new_heading(self.in_heading_level)
                elif tag == "br":
                    paragraph = self._ensure_paragraph()
                    paragraph.add_run().add_break()
                elif tag in ("strong", "b"):
                    self.bold_depth += 1
                elif tag in ("em", "i"):
                    self.italic_depth += 1
                elif tag == "ul":
                    self.list_stack.append("ul")
                elif tag == "ol":
                    self.list_stack.append("ol")
                elif tag == "li":
                    style = "List Number" if self.list_stack and self.list_stack[-1] == "ol" else "List Bullet"
                    self._new_paragraph(style=style)

            def handle_endtag(self, tag):
                tag = (tag or "").lower()
                if tag in ("strong", "b"):
                    self.bold_depth = max(0, self.bold_depth - 1)
                elif tag in ("em", "i"):
                    self.italic_depth = max(0, self.italic_depth - 1)
                elif tag in ("ul", "ol"):
                    if self.list_stack:
                        self.list_stack.pop()
                    self.current_paragraph = None
                elif tag in ("p", "div", "li"):
                    self.current_paragraph = None
                elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    self.in_heading_level = None
                    self.current_paragraph = None

            def handle_data(self, data):
                self._add_text(data)

            def handle_entityref(self, name):
                self._add_text(f"&{name};")

            def handle_charref(self, name):
                self._add_text(f"&#{name};")

        def add_html_formatted_content(document, raw_html, fallback="No disponible"):
            normalized_html = normalize_summary_html(raw_html or "")
            if not normalized_html.strip():
                add_justified_paragraph(document, fallback)
                return

            parser = _DocxHtmlFragmentParser(document)
            try:
                parser.feed(normalized_html)
                parser.close()
            except Exception as parse_error:
                logger.warning(f"No se pudo parsear HTML para docx; usando texto plano. Error: {parse_error}")
                add_justified_paragraph(document, html_to_text_for_export(normalized_html) or fallback)
                return

            if not parser.has_written_content:
                add_justified_paragraph(document, html_to_text_for_export(normalized_html) or fallback)

        doc = Document()
        add_lila_heading(doc, row[1] or f"Subvención {grant_id}", level=1)

        meta = doc.add_paragraph()
        meta.add_run("Plazo: ").bold = True
        meta.add_run(row[3].strftime("%d/%m/%Y") if hasattr(row[3], "strftime") else str(row[3] or "Sin fecha"))
        meta.add_run("\nFondos disponibles: ").bold = True
        meta.add_run(str(row[2] or "No especificado"))
        meta.add_run("\nRegión: ").bold = True
        meta.add_run(str(row[5] or "No especificado"))
        meta.add_run("\nBeneficiarios: ").bold = True
        meta.add_run(str(row[6] or "No especificado"))
        if row[7] is not None:
            meta.add_run("\nCompatibilidad: ").bold = True
            meta.add_run(f"{round(float(row[7]) * 100)}%")

        add_lila_heading(doc, "Resumen", level=2)
        add_html_formatted_content(doc, row[4] or "", fallback="No disponible")

        add_lila_heading(doc, "Por qué este porcentaje de compatibilidad", level=2)
        add_html_formatted_content(doc, row[8] or "", fallback="No disponible")

        add_lila_heading(doc, "Documentación a aportar", level=2)
        if documentacion_items:
            for item in documentacion_items:
                bullet = doc.add_paragraph(item, style="List Bullet")
        else:
            add_justified_paragraph(
                doc,
                "No disponible todavía. Esta sección se completará cuando exista la columna 'Documentacion' en la tabla grants."
            )

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        safe_name = re.sub(r'[^A-Za-z0-9_-]+', '_', (row[1] or f"grant_{grant_id}"))[:80].strip("_")
        filename = f"{safe_name or f'grant_{grant_id}'}_dossier.docx"

        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        logger.error(f"Error generando pack de subvención {grant_id}: {e}", exc_info=True)
        return jsonify({"error": "No se pudo generar el documento de descarga"}), 500
    finally:
        try:
            cursor.close()
            connection.close()
        except Exception:
            pass


@app.route('/api/grants/<int:grant_id>/concessions', methods=['GET'])
@require_auth
def get_grant_concessions(user_id, grant_id):
    """Lista concesiones de la convocatoria BDNS asociada a la subvención."""
    connection = None
    cursor = None
    try:
        page_size = max(1, min(request.args.get('page_size', 50, type=int) or 50, 200))

        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute("""
            SELECT id, titulo_corto, codigobdns
            FROM grants
            WHERE id = %s
              AND resumen_completo IS NOT NULL
            LIMIT 1
        """, (grant_id,))
        row = cursor.fetchone()

        if not row:
            return jsonify(success=False, message="Subvención no encontrada", concessions=[]), 404

        _, titulo_corto, codigobdns = row
        numero_convocatoria = _normalize_bdns_convocatoria_number(codigobdns)
        if not numero_convocatoria:
            return jsonify(
                success=True,
                concessions=[],
                total=0,
                warning="No hay código BDNS disponible para consultar concesiones de esta convocatoria.",
                grant={"id": str(grant_id), "titulo_corto": titulo_corto or "", "codigobdns": None},
                source="BDNS",
                source_url="https://www.pap.hacienda.gob.es/bdnstrans/GE/es/concesiones/consulta",
            )

        result = _fetch_bdns_concesiones_by_convocatoria(numero_convocatoria, page_size=page_size)

        return jsonify(
            success=True,
            concessions=result.get("items", []),
            total=result.get("total", 0),
            warning=result.get("warning"),
            grant={
                "id": str(grant_id),
                "titulo_corto": titulo_corto or "",
                "codigobdns": numero_convocatoria,
            },
            source="BDNS",
            source_url="https://www.pap.hacienda.gob.es/bdnstrans/GE/es/concesiones/consulta",
        )
    except requests.RequestException as e:
        logger.error(f"Error consultando BDNS concesiones para grant {grant_id}: {e}")
        return jsonify(
            success=False,
            message="No se pudieron consultar las concesiones en este momento",
            concessions=[],
        ), 502
    except Exception as e:
        logger.error(f"Error en get_grant_concessions para grant {grant_id}: {e}", exc_info=True)
        return jsonify(success=False, message=str(e), concessions=[]), 500
    finally:
        try:
            cursor.close()
            connection.close()
        except Exception:
            pass


@app.route('/api/grants/<int:grant_id>/download-boe', methods=['GET'])
@require_auth
def download_grant_boe(user_id, grant_id):
    """
    Descarga la normativa legal (BOE) desde S3 si se encuentra por codigobdns.
    """
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        extra = get_grant_extra_metadata(cursor, grant_id)
        codigobdns = extra.get("codigobdns")

        if not codigobdns:
            return jsonify({"error": "No se encontró el identificador del expediente para localizar la normativa"}), 404

        s3_obj = find_boe_s3_object_for_grant(str(codigobdns))
        if not s3_obj:
            return jsonify({"error": "No se encontró la normativa legal (BOE) en AWS para esta subvención"}), 404

        obj = s3_client.get_object(Bucket=s3_obj["bucket"], Key=s3_obj["key"])
        body = obj["Body"].read()
        filename = os.path.basename(s3_obj["key"]) or f"boe_{grant_id}.pdf"
        content_type = obj.get("ContentType") or "application/octet-stream"

        return send_file(
            io.BytesIO(body),
            as_attachment=True,
            download_name=filename,
            mimetype=content_type
        )

    except Exception as e:
        logger.error(f"Error descargando BOE de grant {grant_id}: {e}", exc_info=True)
        return jsonify({"error": "No se pudo descargar la normativa legal"}), 500
    finally:
        try:
            cursor.close()
            connection.close()
        except Exception:
            pass
    
# =============================================
# CRM ENDPOINTS
# =============================================

@app.route('/api/crm/grants', methods=['GET'])
@require_auth
def get_crm_grants(user_id):
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        entity_id = _resolve_entity_id_for_user(cursor, user_id, request.args.get('entity_id'))
        if not entity_id:
            return jsonify(success=False, message="No se encontró una entidad activa para este usuario"), 404

        grants = _fetch_crm_pipeline_grants(cursor, entity_id)
        return jsonify(
            success=True,
            entity_id=str(entity_id),
            grants=grants,
            count=len(grants),
        ), 200
    except Exception as e:
        logger.error(f"Error getting CRM grants for user_id={user_id}: {e}", exc_info=True)
        return jsonify(success=False, message="No se pudo cargar el CRM", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


@app.route('/api/crm/grants', methods=['POST'])
@require_auth
def add_grant_to_crm(user_id):
    connection = None
    cursor = None
    try:
        data = request.json or {}
        grant_id = data.get('grant_id')
        if not grant_id:
            return jsonify(success=False, message="Se requiere grant_id"), 400

        connection = get_connection()
        cursor = connection.cursor()
        _ensure_crm_pipeline_tables(cursor)

        entity_id = _resolve_entity_id_for_user(cursor, user_id, data.get('entity_id'))
        if not entity_id:
            return jsonify(success=False, message="No se encontró una entidad activa para este usuario"), 404

        cursor.execute("SELECT id FROM public.grants WHERE id = %s LIMIT 1", (grant_id,))
        if cursor.fetchone() is None:
            return jsonify(success=False, message="La subvención indicada no existe"), 404

        cursor.execute("""
            SELECT id, status
            FROM public.entity_grant_pipeline
            WHERE entity_id = %s
              AND grant_id = %s
            LIMIT 1
        """, (entity_id, grant_id))
        existing_row = cursor.fetchone()
        if existing_row:
            return jsonify(
                success=True,
                created=False,
                pipeline_id=existing_row[0],
                status=existing_row[1],
                message="La subvención ya estaba en el CRM",
            ), 200

        cursor.execute("""
            INSERT INTO public.entity_grant_pipeline (
                entity_id,
                grant_id,
                status,
                created_by,
                created_at,
                updated_at
            )
            VALUES (%s, %s, 'detectada', %s, NOW(), NOW())
            RETURNING id, status
        """, (entity_id, grant_id, user_id))
        inserted_row = cursor.fetchone()
        pipeline_id = inserted_row[0]
        _record_pipeline_history(cursor, pipeline_id, None, 'detectada', user_id)
        connection.commit()

        return jsonify(
            success=True,
            created=True,
            pipeline_id=pipeline_id,
            status=inserted_row[1],
            message="Subvención añadida al CRM",
        ), 201
    except Exception as e:
        if connection:
            try:
                connection.rollback()
            except Exception:
                pass
        logger.error(f"Error adding grant to CRM for user_id={user_id}: {e}", exc_info=True)
        return jsonify(success=False, message="No se pudo añadir la subvención al CRM", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


@app.route('/api/crm/grants/status/<int:grant_id>', methods=['GET'])
@require_auth
def get_crm_grant_status(user_id, grant_id):
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        _ensure_crm_pipeline_tables(cursor)

        entity_id = _resolve_entity_id_for_user(cursor, user_id, request.args.get('entity_id'))
        if not entity_id:
            return jsonify(success=True, is_in_crm=False, status=None, pipeline_id=None), 200

        cursor.execute("""
            SELECT id, status
            FROM public.entity_grant_pipeline
            WHERE entity_id = %s
              AND grant_id = %s
            LIMIT 1
        """, (entity_id, grant_id))
        row = cursor.fetchone()

        return jsonify(
            success=True,
            is_in_crm=row is not None,
            pipeline_id=row[0] if row else None,
            status=row[1] if row else None,
        ), 200
    except Exception as e:
        logger.error(f"Error checking CRM status for grant_id={grant_id}: {e}", exc_info=True)
        return jsonify(success=False, message="No se pudo comprobar el estado del CRM", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


@app.route('/api/crm/grants/<int:pipeline_id>/status', methods=['PATCH'])
@require_auth
def update_crm_grant_status(user_id, pipeline_id):
    connection = None
    cursor = None
    try:
        data = request.json or {}
        next_status = str(data.get('status') or '').strip().lower()
        notes = data.get('notes')

        if next_status not in CRM_PIPELINE_ALLOWED_STATUSES:
            return jsonify(success=False, message="Estado de CRM no válido"), 400

        connection = get_connection()
        cursor = connection.cursor()
        _ensure_crm_pipeline_tables(cursor)

        entity_id = _resolve_entity_id_for_user(cursor, user_id, data.get('entity_id'))
        if not entity_id:
            return jsonify(success=False, message="No se encontró una entidad activa para este usuario"), 404

        cursor.execute("""
            SELECT p.status
            FROM public.entity_grant_pipeline p
            WHERE p.id = %s
              AND p.entity_id = %s
            LIMIT 1
        """, (pipeline_id, entity_id))
        row = cursor.fetchone()
        if row is None:
            return jsonify(success=False, message="La subvención no existe en el CRM de esta entidad"), 404

        previous_status = row[0]
        if previous_status == next_status:
            return jsonify(success=True, status=next_status, message="Estado sin cambios"), 200

        cursor.execute("""
            UPDATE public.entity_grant_pipeline
            SET status = %s,
                notes = COALESCE(%s, notes),
                updated_at = NOW()
            WHERE id = %s
        """, (next_status, notes, pipeline_id))
        _record_pipeline_history(cursor, pipeline_id, previous_status, next_status, user_id, notes)
        connection.commit()

        return jsonify(success=True, status=next_status, previous_status=previous_status), 200
    except Exception as e:
        if connection:
            try:
                connection.rollback()
            except Exception:
                pass
        logger.error(f"Error updating CRM status pipeline_id={pipeline_id}: {e}", exc_info=True)
        return jsonify(success=False, message="No se pudo actualizar el estado del CRM", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


@app.route('/api/crm/grants/<int:pipeline_id>', methods=['DELETE'])
@require_auth
def delete_crm_grant(user_id, pipeline_id):
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        _ensure_crm_pipeline_tables(cursor)

        entity_id = _resolve_entity_id_for_user(cursor, user_id, request.args.get('entity_id'))
        if not entity_id:
            return jsonify(success=False, message="No se encontró una entidad activa para este usuario"), 404

        cursor.execute("""
            DELETE FROM public.entity_grant_pipeline
            WHERE id = %s
              AND entity_id = %s
            RETURNING id
        """, (pipeline_id, entity_id))
        deleted_row = cursor.fetchone()
        if deleted_row is None:
            return jsonify(success=False, message="La subvención no existe en el CRM de esta entidad"), 404

        connection.commit()
        return jsonify(success=True, message="Subvención eliminada del CRM"), 200
    except Exception as e:
        if connection:
            try:
                connection.rollback()
            except Exception:
                pass
        logger.error(f"Error deleting CRM grant pipeline_id={pipeline_id}: {e}", exc_info=True)
        return jsonify(success=False, message="No se pudo eliminar la subvención del CRM", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


@app.route('/api/crm/corporate-documents', methods=['GET'])
@require_auth
def get_corporate_documents(user_id):
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        _ensure_corporate_document_tables(cursor)
        _sync_legacy_corporate_documents(cursor)

        entity_id = _resolve_entity_id_for_user(cursor, user_id, request.args.get('entity_id'))
        if not entity_id:
            return jsonify(success=False, message="No se encontró una entidad activa para este usuario"), 404

        _backfill_inferred_entity_document_types(cursor, entity_id)
        connection.commit()

        cursor.execute("""
            SELECT
                c.code,
                c.label,
                c.description,
                c.display_order,
                d.id,
                d.original_filename,
                d.s3_key,
                d.s3_bucket,
                d.file_size,
                d.mime_type,
                d.status,
                d.upload_date
            FROM public.corporate_document_catalog c
            LEFT JOIN LATERAL (
                SELECT
                    id,
                    original_filename,
                    s3_key,
                    s3_bucket,
                    file_size,
                    mime_type,
                    status,
                    upload_date
                FROM public.entity_documents
                WHERE entity_id = %s
                  AND document_type_code = c.code
                  AND COALESCE(is_current, TRUE) IS TRUE
                ORDER BY upload_date DESC
                LIMIT 1
            ) d ON TRUE
            WHERE c.is_required IS TRUE
            ORDER BY c.display_order ASC, c.label ASC
        """, (entity_id,))

        items = []
        uploaded_count = 0
        for row in cursor.fetchall():
            has_file = row[4] is not None
            if has_file:
                uploaded_count += 1

            items.append({
                "document_type_code": row[0],
                "label": row[1],
                "description": row[2],
                "display_order": row[3],
                "has_file": has_file,
                "document": {
                    "id": row[4],
                    "filename": row[5],
                    "s3_key": row[6],
                    "s3_bucket": row[7],
                    "file_size": row[8],
                    "mime_type": row[9],
                    "status": row[10],
                    "upload_date": _serialize_datetimeish(row[11]),
                } if has_file else None,
            })

        cursor.execute("""
            SELECT
                d.id,
                d.original_filename,
                d.file_size,
                d.mime_type,
                d.status,
                d.upload_date
            FROM public.entity_documents d
            WHERE d.entity_id = %s
              AND d.document_type_code = %s
              AND COALESCE(d.is_current, TRUE) IS TRUE
            ORDER BY d.upload_date DESC, d.id DESC
        """, (entity_id, EXTRA_CORPORATE_DOCUMENT_CODE))
        extra_documents = [
            {
                "id": row[0],
                "filename": row[1],
                "file_size": row[2],
                "mime_type": row[3],
                "status": row[4],
                "upload_date": _serialize_datetimeish(row[5]),
                "document_type_code": EXTRA_CORPORATE_DOCUMENT_CODE,
                "document_type_label": EXTRA_CORPORATE_DOCUMENT_LABEL,
            }
            for row in cursor.fetchall()
        ]

        cursor.execute("""
            SELECT COUNT(*)
            FROM public.entity_documents
            WHERE entity_id = %s
              AND COALESCE(is_current, TRUE) IS TRUE
        """, (entity_id,))
        total_current_documents = int(cursor.fetchone()[0] or 0)

        return jsonify(
            success=True,
            entity_id=str(entity_id),
            items=items,
            extra_documents=extra_documents,
            extra_documents_count=len(extra_documents),
            uploaded_count=uploaded_count,
            missing_count=max(0, len(items) - uploaded_count),
            total_required=len(items),
            total_current_documents=total_current_documents,
        ), 200
    except Exception as e:
        logger.error(f"Error fetching corporate documents for user_id={user_id}: {e}", exc_info=True)
        return jsonify(success=False, message="No se pudo cargar la documentación corporativa", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


@app.route('/api/crm/corporate-documents/upload', methods=['POST'])
@require_auth
def upload_corporate_document(user_id):
    connection = None
    cursor = None
    try:
        document_type_code = (request.form.get('document_type_code') or '').strip()
        uploaded_file = request.files.get('file')

        if not document_type_code:
            return jsonify(success=False, message="Se requiere document_type_code"), 400
        if not uploaded_file or uploaded_file.filename == '':
            return jsonify(success=False, message="Debes adjuntar un archivo"), 400

        connection = get_connection()
        cursor = connection.cursor()
        _ensure_corporate_document_tables(cursor)

        entity_id = _resolve_entity_id_for_user(cursor, user_id, request.form.get('entity_id'))
        if not entity_id:
            return jsonify(success=False, message="No se encontró una entidad activa para este usuario"), 404

        cursor.execute("""
            SELECT label
            FROM public.corporate_document_catalog
            WHERE code = %s
            LIMIT 1
        """, (document_type_code,))
        catalog_row = cursor.fetchone()
        if catalog_row is None:
            return jsonify(success=False, message="Tipo documental no válido"), 400

        stored_document = _store_entity_document(
            cursor,
            entity_id=entity_id,
            uploaded_file=uploaded_file,
            uploaded_by=user_id,
            document_type_code=document_type_code,
        )
        connection.commit()

        return jsonify(
            success=True,
            message=f"{catalog_row[0]} subido correctamente",
            document=stored_document,
        ), 201
    except Exception as e:
        if connection:
            try:
                connection.rollback()
            except Exception:
                pass
        logger.error(f"Error uploading corporate document for user_id={user_id}: {e}", exc_info=True)
        return jsonify(success=False, message="No se pudo subir el documento corporativo", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


@app.route('/api/crm/corporate-documents/export', methods=['GET'])
@require_auth
def export_corporate_documents_zip(user_id):
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        _ensure_corporate_document_tables(cursor)
        _sync_legacy_corporate_documents(cursor)

        entity_id = _resolve_entity_id_for_user(cursor, user_id, request.args.get('entity_id'))
        if not entity_id:
            return jsonify(success=False, message="No se encontró una entidad activa para este usuario"), 404

        _backfill_inferred_entity_document_types(cursor, entity_id)
        connection.commit()

        cursor.execute("""
            SELECT
                d.original_filename,
                d.s3_key,
                d.s3_bucket,
                COALESCE(
                    c.label,
                    CASE
                        WHEN d.document_type_code = %s THEN %s
                        ELSE NULL
                    END
                ) AS document_type_label
            FROM public.entity_documents d
            LEFT JOIN public.corporate_document_catalog c
                ON c.code = d.document_type_code
            WHERE d.entity_id = %s
              AND COALESCE(d.is_current, TRUE) IS TRUE
              AND EXISTS (
                  SELECT 1
                  FROM public.user_entities ue
                  WHERE ue.entity_id = d.entity_id
                    AND ue.user_id = %s
              )
            ORDER BY d.upload_date DESC, d.id DESC
        """, (
            EXTRA_CORPORATE_DOCUMENT_CODE,
            EXTRA_CORPORATE_DOCUMENT_LABEL,
            entity_id,
            user_id,
        ))
        rows = cursor.fetchall()

        if not rows:
            return jsonify(success=False, message="No hay documentos para exportar todavía"), 404

        zip_buffer = io.BytesIO()
        used_names = set()
        written_files = 0

        with zipfile.ZipFile(zip_buffer, mode='w', compression=zipfile.ZIP_DEFLATED) as zip_file:
            for original_filename, s3_key, s3_bucket, document_type_label in rows:
                try:
                    obj = s3_client.get_object(Bucket=s3_bucket, Key=s3_key)
                    body = obj["Body"].read()
                except Exception as exc:
                    logger.warning(
                        "No se pudo incluir el documento %s en la exportacion ZIP de la entidad %s: %s",
                        original_filename,
                        entity_id,
                        exc,
                    )
                    continue

                entry_name = _build_unique_zip_entry_name(
                    document_type_label,
                    original_filename,
                    used_names,
                )
                zip_file.writestr(entry_name, body)
                written_files += 1

        if written_files == 0:
            return jsonify(success=False, message="No se pudo preparar el ZIP de documentos"), 500

        zip_bytes = zip_buffer.getvalue()
        export_date = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"documentacion_entidad_{entity_id}_{export_date}.zip"
        response = make_response(zip_bytes)
        response.status_code = 200
        response.headers["Content-Type"] = "application/zip"
        response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
        response.headers["Content-Length"] = str(len(zip_bytes))
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0, no-transform"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["Access-Control-Expose-Headers"] = "Content-Disposition, Content-Length, Content-Type"
        return response
    except Exception as e:
        logger.error(
            "Error exporting corporate documents ZIP for user_id=%s: %s",
            user_id,
            e,
            exc_info=True,
        )
        return jsonify(success=False, message="No se pudo exportar la documentación", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


@app.route('/api/crm/corporate-documents/<int:document_id>/download', methods=['GET'])
@require_auth
def download_corporate_document(user_id, document_id):
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        _ensure_corporate_document_tables(cursor)
        _sync_legacy_corporate_documents(cursor)
        connection.commit()

        cursor.execute("""
            SELECT
                d.original_filename,
                d.s3_key,
                d.s3_bucket,
                d.mime_type
            FROM public.entity_documents d
            JOIN public.user_entities ue ON ue.entity_id = d.entity_id
            WHERE d.id = %s
              AND ue.user_id = %s
              AND d.document_type_code IS NOT NULL
            LIMIT 1
        """, (document_id, user_id))
        row = cursor.fetchone()

        if row is None:
            cursor.execute("""
                SELECT
                    d.original_filename,
                    d.s3_key,
                    d.s3_bucket,
                    d.mime_type
                FROM public.entity_corporate_documents d
                JOIN public.user_entities ue ON ue.entity_id = d.entity_id
                WHERE d.id = %s
                  AND ue.user_id = %s
                LIMIT 1
            """, (document_id, user_id))
            row = cursor.fetchone()

        if row is None:
            return jsonify(success=False, message="Documento no encontrado"), 404

        obj = s3_client.get_object(Bucket=row[2], Key=row[1])
        body = obj["Body"].read()
        content_type = row[3] or obj.get("ContentType") or "application/octet-stream"

        return send_file(
            io.BytesIO(body),
            as_attachment=True,
            download_name=row[0] or f"documento_{document_id}",
            mimetype=content_type,
        )
    except Exception as e:
        logger.error(f"Error downloading corporate document document_id={document_id}: {e}", exc_info=True)
        return jsonify(success=False, message="No se pudo descargar el documento", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


@app.route('/api/crm/corporate-documents/<int:document_id>', methods=['DELETE'])
@require_auth
def delete_corporate_document(user_id, document_id):
    connection = None
    cursor = None
    try:
        connection = get_connection()
        cursor = connection.cursor()
        _ensure_corporate_document_tables(cursor)
        _sync_legacy_corporate_documents(cursor)
        connection.commit()

        cursor.execute("""
            SELECT
                d.id,
                d.original_filename,
                d.s3_key,
                d.s3_bucket,
                d.document_type_code
            FROM public.entity_documents d
            WHERE d.id = %s
              AND COALESCE(d.is_current, TRUE) IS TRUE
              AND EXISTS (
                  SELECT 1
                  FROM public.user_entities ue
                  WHERE ue.entity_id = d.entity_id
                    AND ue.user_id = %s
              )
            LIMIT 1
        """, (document_id, user_id))
        row = cursor.fetchone()

        if row is None:
            return jsonify(success=False, message="Documento no encontrado"), 404

        cursor.execute("""
            DELETE FROM public.entity_documents
            WHERE id = %s
        """, (document_id,))
        connection.commit()

        try:
            s3_client.delete_object(Bucket=row[3], Key=row[2])
        except Exception as storage_error:
            logger.warning(
                "Documento %s eliminado de BD pero no de S3 (user_id=%s): %s",
                document_id,
                user_id,
                storage_error,
            )

        logger.info(
            "Documento eliminado correctamente. user_id=%s document_id=%s filename=%s type=%s",
            user_id,
            document_id,
            row[1],
            row[4],
        )
        return jsonify(
            success=True,
            message="Documento eliminado correctamente",
            document_id=document_id,
            filename=row[1],
            document_type_code=row[4],
        ), 200
    except Exception as e:
        if connection:
            try:
                connection.rollback()
            except Exception:
                pass
        logger.error(
            "Error deleting corporate document for user_id=%s document_id=%s: %s",
            user_id,
            document_id,
            e,
            exc_info=True,
        )
        return jsonify(success=False, message="No se pudo eliminar el documento", error=str(e)), 500
    finally:
        if connection:
            try:
                cursor.close()
                connection.close()
            except Exception:
                pass


# =============================================
# FAVORITES ENDPOINTS
# =============================================

@app.route('/api/favorites/add', methods=['POST'])
@require_auth
def add_favorite(user_id):
    """Add a grant to user's favorites"""
    try:
        data = request.json
        grant_id = data.get('grant_id')

        if not grant_id:
            return jsonify({"error": "Se requiere grant_id"}), 400

        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute("""
            INSERT INTO user_favorites (user_id, grant_id)
            VALUES (%s, %s)
            ON CONFLICT (user_id, grant_id) DO NOTHING
        """, (user_id, int(grant_id)))

        connection.commit()
        return jsonify({"status": "success", "message": "Favorito guardado"}), 200

    except Exception as e:
        logger.error(f"Error adding favorite: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()


@app.route('/api/favorites/remove', methods=['POST'])
@require_auth
def remove_favorite(user_id):
    """Remove a grant from user's favorites"""
    try:
        data = request.json
        grant_id = data.get('grant_id')

        if not grant_id:
            return jsonify({"error": "Se requiere grant_id"}), 400

        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute("""
            DELETE FROM user_favorites
            WHERE user_id = %s AND grant_id = %s
        """, (user_id, int(grant_id)))

        connection.commit()
        return jsonify({"status": "success", "message": "Favorito eliminado"}), 200

    except Exception as e:
        logger.error(f"Error removing favorite: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()


@app.route('/api/favorites', methods=['GET'])
@require_auth
def get_favorites(user_id):
    """Get all user's favorite grants with details"""
    try:
        connection = get_connection()
        cursor = connection.cursor()
        selected_entity_id = None

        try:
            cursor.execute("""
                SELECT ue.entity_id
                FROM user_entities ue
                WHERE ue.user_id = %s
                  AND ue.is_selected IS TRUE
                LIMIT 1
            """, (user_id,))
            selected_row = cursor.fetchone()
            if selected_row:
                selected_entity_id = selected_row[0]
        except Exception as entity_err:
            logger.warning(f"No se pudo obtener entidad seleccionada para favoritos user_id={user_id}: {entity_err}")
            selected_entity_id = None

        def _format_favorite_budget(value):
            if value in (None, ""):
                return "No especificado"
            # Si ya viene con símbolo de euro o texto, lo respetamos
            if isinstance(value, str) and "€" in value:
                return value

            numeric = _safe_float(value)
            if numeric is None:
                return str(value)

            formatted = f"{numeric:,.2f}"
            formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
            if formatted.endswith(",00"):
                formatted = formatted[:-3]
            return f"{formatted}€"

        cursor.execute("""
            SELECT
                g.id,
                g.titulo_corto,
                g.presupuesto,
                g.fecha_finalizacion,
                g.fecha_de_publicacion,
                g.resumen_completo,
                string_agg(DISTINCT b.value, ', ') AS beneficiarios,
                g.region_impacto,
                uf.created_at,
                m.numero_match
            FROM user_favorites uf
            JOIN grants g ON uf.grant_id = g.id
            LEFT JOIN LATERAL (
                SELECT numero_match
                FROM matches
                WHERE matches.grant_id = g.id
                  AND matches.entity_id = %s
                ORDER BY numero_match DESC
                LIMIT 1
            ) m ON TRUE
            LEFT JOIN LATERAL (
                SELECT value
                FROM jsonb_array_elements_text(
                    CASE
                        WHEN jsonb_typeof(g."Beneficiarios_Short"->'categorias') = 'array'
                        THEN g."Beneficiarios_Short"->'categorias'
                        ELSE '[]'::jsonb
                    END
                )
            ) AS b(value) ON TRUE
            WHERE uf.user_id = %s
              AND g.resumen_completo IS NOT NULL
            GROUP BY g.id, g.titulo_corto, g.presupuesto, g.fecha_finalizacion, g.fecha_de_publicacion,
                     g.resumen_completo, g.region_impacto, uf.created_at, m.numero_match
            ORDER BY uf.created_at DESC
        """, (selected_entity_id, user_id))

        rows = cursor.fetchall()
        favorites = []
        for row in rows:
            deadline = _format_deadline_for_ui(row[3], row[4])

            favorites.append({
                "grant_id": row[0],
                "titulo_corto": row[1],
                "presupuesto": _format_favorite_budget(row[2]),
                "fecha_limite": deadline,
                "resumen": row[5] if row[5] else "Sin descripcion",
                "beneficiarios": row[6] if row[6] else "No especificado",
                "region_impacto": row[7] if row[7] else "No especificado",
                "favorited_at": row[8].isoformat() if row[8] else None,
                "numero_match": (
                    int(round(float(row[9]) * 100)) if row[9] is not None and float(row[9]) <= 1
                    else int(round(float(row[9]))) if row[9] is not None
                    else None
                ),
            })

        return jsonify({"favorites": favorites, "count": len(favorites)}), 200

    except Exception as e:
        logger.error(f"Error getting favorites: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()


@app.route('/api/favorites/status/<int:grant_id>', methods=['GET'])
@require_auth
def get_favorite_status(user_id, grant_id):
    """Check if a grant is in user's favorites"""
    try:
        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute("""
            SELECT 1
            FROM user_favorites
            WHERE user_id = %s AND grant_id = %s
            LIMIT 1
        """, (user_id, grant_id))

        is_favorite = cursor.fetchone() is not None
        return jsonify({"is_favorite": is_favorite}), 200

    except Exception as e:
        logger.error(f"Error getting favorite status: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()


# =============================================
# ALERTS ENDPOINTS
# =============================================

@app.route('/api/alerts/create', methods=['POST'])
@require_auth
def create_alert(user_id):
    """Save current filter preferences as an alert"""
    try:
        data = request.json
        alert_name = data.get('alert_name', 'Mi alerta')
        filters = data.get('filters', {})

        if not isinstance(filters, dict):
            return jsonify({"error": "Formato de filtros invalido"}), 400

        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute("""
            INSERT INTO user_alerts (user_id, alert_name, filters)
            VALUES (%s, %s, %s)
            RETURNING id, created_at
        """, (user_id, alert_name, json.dumps(filters)))

        result = cursor.fetchone()
        connection.commit()

        return jsonify({
            "status": "success",
            "alert": {
                "id": result[0],
                "alert_name": alert_name,
                "filters": filters,
                "created_at": result[1].isoformat() if result[1] else None,
                "is_active": True
            }
        }), 201

    except Exception as e:
        logger.error(f"Error creating alert: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()


@app.route('/api/alerts', methods=['GET'])
@require_auth
def get_alerts(user_id):
    """Get all user's saved alerts"""
    try:
        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute("""
            SELECT id, alert_name, filters, created_at, is_active
            FROM user_alerts
            WHERE user_id = %s
            ORDER BY created_at DESC
        """, (user_id,))

        rows = cursor.fetchall()
        alerts = []
        for row in rows:
            alerts.append({
                "id": row[0],
                "alert_name": row[1],
                "filters": row[2] if row[2] else {},
                "created_at": row[3].isoformat() if row[3] else None,
                "is_active": row[4]
            })

        return jsonify({"alerts": alerts, "count": len(alerts)}), 200

    except Exception as e:
        logger.error(f"Error getting alerts: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()


@app.route('/api/alerts/<int:alert_id>', methods=['DELETE'])
@require_auth
def delete_alert(user_id, alert_id):
    """Delete a user alert"""
    try:
        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute("""
            DELETE FROM user_alerts
            WHERE id = %s AND user_id = %s
        """, (alert_id, user_id))

        if cursor.rowcount == 0:
            return jsonify({"error": "Alerta no encontrada"}), 404

        connection.commit()
        return jsonify({"status": "success", "message": "Alerta eliminada"}), 200

    except Exception as e:
        logger.error(f"Error deleting alert: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()


@app.route('/api/alerts/<int:alert_id>/toggle', methods=['PATCH'])
@require_auth
def toggle_alert(user_id, alert_id):
    """Toggle an alert's active status"""
    try:
        connection = get_connection()
        cursor = connection.cursor()

        cursor.execute("""
            UPDATE user_alerts
            SET is_active = NOT is_active
            WHERE id = %s AND user_id = %s
            RETURNING is_active
        """, (alert_id, user_id))

        result = cursor.fetchone()
        if not result:
            return jsonify({"error": "Alerta no encontrada"}), 404

        connection.commit()
        return jsonify({"status": "success", "is_active": result[0]}), 200

    except Exception as e:
        logger.error(f"Error toggling alert: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        if 'connection' in locals() and connection:
            cursor.close()
            connection.close()


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
